#!/usr/bin/env python3
"""Quiz Master Pro - Backend Server"""

import sys
import os

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
import sqlite3
import hashlib
import secrets
import json
import os
import re
import io
import tempfile
import threading
from datetime import datetime, timedelta
from functools import wraps

app = Flask(__name__, static_folder=os.path.join(BASE_DIR, 'static'))
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16 MB upload limit
CORS(app)

# Ensure all HTTP error responses come back as JSON, not HTML.
# Flask's default handlers return HTML pages which the frontend cannot parse.
@app.errorhandler(400)
def bad_request_handler(e):
    return jsonify({'error': str(e)}), 400

@app.errorhandler(413)
def too_large(e):
    return jsonify({'error': 'File too large. Maximum upload size is 16MB.'}), 413

DATABASE = os.path.join(BASE_DIR, 'quiz_master.db')

# Admin token for protected admin endpoints (set this to a secret value in production)
ADMIN_TOKEN = os.environ.get('QUIZ_ADMIN_TOKEN', 'change-me-in-production')

# DB write lock for SQLite thread safety
_db_write_lock = threading.Lock()

# Rate limiting (Flask-Limiter)
try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
    limiter = Limiter(
        get_remote_address,
        app=app,
        default_limits=[],
        storage_uri="memory://"
    )
    LIMITER_AVAILABLE = True
except ImportError:
    LIMITER_AVAILABLE = False
    print("Warning: flask-limiter not installed. Rate limiting disabled. Install with: pip install flask-limiter")

# === Static Routes ===

@app.route('/')
def index():
    return send_from_directory(BASE_DIR, 'index.html')

@app.route('/app')
def app_page():
    return send_from_directory(BASE_DIR, 'index.html')

@app.route('/static/<path:filename>')
def serve_static(filename):
    return send_from_directory(os.path.join(BASE_DIR, 'static'), filename)

# === Database ===

def get_db():
    conn = sqlite3.connect(DATABASE, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    conn.execute('PRAGMA foreign_keys = ON')
    return conn

def init_db():
    conn = get_db()
    c = conn.cursor()
    
    c.execute('''CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        email TEXT UNIQUE NOT NULL,
        password_hash TEXT NOT NULL,
        salt TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        last_login TIMESTAMP,
        is_active BOOLEAN DEFAULT 1
    )''')
    
    # Email verification columns (added in Phase 3 — idempotent ALTER TABLE)
    for _col, _def in [
        ('email_verified',      'BOOLEAN DEFAULT 0'),
        ('email_token',         'TEXT'),
        ('email_token_expires', 'TIMESTAMP'),
    ]:
        try:
            c.execute(f'ALTER TABLE users ADD COLUMN {_col} {_def}')
        except sqlite3.OperationalError:
            pass  # column already exists

    c.execute('''CREATE TABLE IF NOT EXISTS sessions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        token TEXT UNIQUE NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        expires_at TIMESTAMP NOT NULL,
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS quizzes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        title TEXT NOT NULL,
        description TEXT,
        questions TEXT NOT NULL,
        color TEXT DEFAULT '#6366f1',
        is_public BOOLEAN DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        last_modified TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
    )''')
    
    c.execute('''CREATE TABLE IF NOT EXISTS attempts (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        quiz_id INTEGER NOT NULL,
        user_id INTEGER NOT NULL,
        score INTEGER NOT NULL,
        total INTEGER NOT NULL,
        percentage INTEGER NOT NULL,
        answers TEXT NOT NULL,
        study_mode BOOLEAN DEFAULT 0,
        timed BOOLEAN DEFAULT 0,
        max_streak INTEGER DEFAULT 0,
        time_taken INTEGER,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (quiz_id) REFERENCES quizzes(id) ON DELETE CASCADE,
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
    )''')
    
    # NEW: User profiles for gamification data sync
    c.execute('''CREATE TABLE IF NOT EXISTS user_profiles (
        user_id INTEGER PRIMARY KEY,
        xp INTEGER DEFAULT 0,
        level INTEGER DEFAULT 1,
        gems INTEGER DEFAULT 0,
        daily_streak INTEGER DEFAULT 0,
        last_active_date TEXT,
        achievements TEXT DEFAULT '[]',
        total_answered INTEGER DEFAULT 0,
        total_correct INTEGER DEFAULT 0,
        quizzes_completed INTEGER DEFAULT 0,
        perfect_scores INTEGER DEFAULT 0,
        settings TEXT DEFAULT '{"soundEnabled":true,"animationsEnabled":true}',
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
    )''')
    
    # NEW: In-progress quiz state for cross-device sync
    c.execute('''CREATE TABLE IF NOT EXISTS quiz_progress (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        quiz_id INTEGER NOT NULL,
        question_index INTEGER DEFAULT 0,
        answers TEXT DEFAULT '[]',
        flagged TEXT DEFAULT '[]',
        study_mode BOOLEAN DEFAULT 1,
        randomize_options BOOLEAN DEFAULT 0,
        option_shuffles TEXT DEFAULT '{}',
        quiz_streak INTEGER DEFAULT 0,
        max_quiz_streak INTEGER DEFAULT 0,
        timer_enabled BOOLEAN DEFAULT 0,
        time_remaining INTEGER,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        UNIQUE(user_id, quiz_id),
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY (quiz_id) REFERENCES quizzes(id) ON DELETE CASCADE
    )''')
    
    c.execute('CREATE INDEX IF NOT EXISTS idx_sessions_token ON sessions(token)')
    c.execute('CREATE INDEX IF NOT EXISTS idx_quizzes_user ON quizzes(user_id)')
    c.execute('CREATE INDEX IF NOT EXISTS idx_progress_user_quiz ON quiz_progress(user_id, quiz_id)')

    # === Phase 1: Normalized question storage & certification support ===

    c.execute('''CREATE TABLE IF NOT EXISTS certifications (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        code TEXT UNIQUE NOT NULL,
        name TEXT NOT NULL,
        vendor TEXT NOT NULL,
        description TEXT,
        passing_score INTEGER,
        passing_scale TEXT DEFAULT 'percentage',
        exam_duration_minutes INTEGER,
        total_questions INTEGER,
        is_active BOOLEAN DEFAULT 1,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS domains (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        certification_id INTEGER,
        name TEXT NOT NULL,
        code TEXT,
        weight REAL,
        parent_domain_id INTEGER,
        sort_order INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (certification_id) REFERENCES certifications(id) ON DELETE CASCADE,
        FOREIGN KEY (parent_domain_id) REFERENCES domains(id) ON DELETE SET NULL
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS questions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        quiz_id INTEGER NOT NULL,
        question_index INTEGER NOT NULL,
        question_text TEXT NOT NULL,
        type TEXT NOT NULL DEFAULT 'choice',
        options TEXT,
        correct TEXT,
        pairs TEXT,
        code TEXT,
        code_language TEXT,
        image TEXT,
        image_alt TEXT,
        explanation TEXT,
        difficulty INTEGER DEFAULT 0,
        is_active BOOLEAN DEFAULT 1,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (quiz_id) REFERENCES quizzes(id) ON DELETE CASCADE
    )''')

    # question_domains table: links questions to certification domains
    # Used by certification question bank (managed by admin), NOT by user quizzes
    c.execute('''CREATE TABLE IF NOT EXISTS question_domains (
        question_id INTEGER NOT NULL,
        domain_id INTEGER NOT NULL,
        PRIMARY KEY (question_id, domain_id),
        FOREIGN KEY (question_id) REFERENCES questions(id) ON DELETE CASCADE,
        FOREIGN KEY (domain_id) REFERENCES domains(id) ON DELETE CASCADE
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS question_performance (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        question_id INTEGER NOT NULL,
        times_seen INTEGER DEFAULT 0,
        times_correct INTEGER DEFAULT 0,
        times_incorrect INTEGER DEFAULT 0,
        last_seen_at TIMESTAMP,
        last_correct_at TIMESTAMP,
        average_time_ms INTEGER,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        UNIQUE(user_id, question_id),
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY (question_id) REFERENCES questions(id) ON DELETE CASCADE
    )''')

    c.execute('CREATE INDEX IF NOT EXISTS idx_questions_quiz ON questions(quiz_id)')
    c.execute('CREATE INDEX IF NOT EXISTS idx_questions_quiz_index ON questions(quiz_id, question_index)')
    c.execute('CREATE INDEX IF NOT EXISTS idx_qd_domain ON question_domains(domain_id)')
    c.execute('CREATE INDEX IF NOT EXISTS idx_qp_user ON question_performance(user_id)')

    # === Phase 2: Certification tracking & exam simulation ===

    c.execute('''CREATE TABLE IF NOT EXISTS user_certifications (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        certification_id INTEGER NOT NULL,
        target_date DATE,
        status TEXT DEFAULT 'studying',
        started_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        completed_at TIMESTAMP,
        UNIQUE(user_id, certification_id),
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY (certification_id) REFERENCES certifications(id) ON DELETE CASCADE
    )''')

    c.execute('''CREATE TABLE IF NOT EXISTS exam_simulations (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        certification_id INTEGER NOT NULL,
        score INTEGER NOT NULL,
        total INTEGER NOT NULL,
        percentage REAL NOT NULL,
        passed BOOLEAN NOT NULL,
        time_taken INTEGER,
        time_limit INTEGER,
        domain_scores TEXT NOT NULL,
        answers TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY (certification_id) REFERENCES certifications(id) ON DELETE CASCADE
    )''')

    c.execute('CREATE INDEX IF NOT EXISTS idx_sim_user_cert ON exam_simulations(user_id, certification_id)')

    # Study sessions for analytics
    c.execute('''CREATE TABLE IF NOT EXISTS study_sessions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        session_type TEXT NOT NULL DEFAULT 'quiz',
        quiz_id INTEGER,
        certification_id INTEGER,
        started_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        ended_at TIMESTAMP,
        questions_reviewed INTEGER DEFAULT 0,
        questions_correct INTEGER DEFAULT 0,
        duration_seconds INTEGER DEFAULT 0,
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY (quiz_id) REFERENCES quizzes(id) ON DELETE SET NULL
    )''')
    c.execute('CREATE INDEX IF NOT EXISTS idx_ss_user ON study_sessions(user_id, started_at)')

    # Spaced Repetition System (SRS)
    c.execute('''CREATE TABLE IF NOT EXISTS srs_cards (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        question_id INTEGER NOT NULL,
        ease_factor REAL DEFAULT 2.5,
        interval_days INTEGER DEFAULT 0,
        repetitions INTEGER DEFAULT 0,
        next_review_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        last_reviewed_at TIMESTAMP,
        status TEXT DEFAULT 'new',
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        UNIQUE(user_id, question_id),
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY (question_id) REFERENCES questions(id) ON DELETE CASCADE
    )''')
    c.execute('CREATE INDEX IF NOT EXISTS idx_srs_user_next ON srs_cards(user_id, next_review_at)')

    # Bookmarks
    c.execute('''CREATE TABLE IF NOT EXISTS bookmarks (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        question_id INTEGER NOT NULL,
        note TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        UNIQUE(user_id, question_id),
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY (question_id) REFERENCES questions(id) ON DELETE CASCADE
    )''')
    c.execute('CREATE INDEX IF NOT EXISTS idx_bookmarks_user ON bookmarks(user_id)')

    # Objective confidence self-assessment
    c.execute('''CREATE TABLE IF NOT EXISTS objective_confidence (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        domain_id INTEGER NOT NULL,
        confidence INTEGER NOT NULL DEFAULT 0,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        UNIQUE(user_id, domain_id),
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE,
        FOREIGN KEY (domain_id) REFERENCES domains(id) ON DELETE CASCADE
    )''')
    c.execute('CREATE INDEX IF NOT EXISTS idx_obj_conf_user ON objective_confidence(user_id)')

    # Study resources linked to certifications
    c.execute('''CREATE TABLE IF NOT EXISTS study_resources (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        certification_id INTEGER NOT NULL,
        domain_id INTEGER,
        title TEXT NOT NULL,
        url TEXT NOT NULL,
        resource_type TEXT NOT NULL DEFAULT 'article',
        provider TEXT,
        is_free BOOLEAN DEFAULT 1,
        sort_order INTEGER DEFAULT 0,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (certification_id) REFERENCES certifications(id) ON DELETE CASCADE,
        FOREIGN KEY (domain_id) REFERENCES domains(id) ON DELETE SET NULL
    )''')
    c.execute('CREATE INDEX IF NOT EXISTS idx_study_res_cert ON study_resources(certification_id)')

    # AI quiz generation usage tracking & rate limiting
    c.execute('''CREATE TABLE IF NOT EXISTS ai_usage (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        input_tokens INTEGER DEFAULT 0,
        output_tokens INTEGER DEFAULT 0,
        question_count INTEGER DEFAULT 0,
        model TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE CASCADE
    )''')
    c.execute('CREATE INDEX IF NOT EXISTS idx_ai_usage_user_created ON ai_usage(user_id, created_at)')

    # Phase 7.4 - Usage analytics event log
    c.execute('''CREATE TABLE IF NOT EXISTS events (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER,
        event TEXT NOT NULL,
        metadata TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users(id) ON DELETE SET NULL
    )''')
    c.execute('CREATE INDEX IF NOT EXISTS idx_events_user ON events(user_id)')
    c.execute('CREATE INDEX IF NOT EXISTS idx_events_created ON events(created_at)')

    # Phase 7.2 - Additional performance indexes
    c.execute('CREATE INDEX IF NOT EXISTS idx_questions_quiz_id ON questions(quiz_id)')
    c.execute('CREATE INDEX IF NOT EXISTS idx_question_performance_user ON question_performance(user_id, question_id)')
    c.execute('CREATE INDEX IF NOT EXISTS idx_question_domains_domain ON question_domains(domain_id)')

    # Enable WAL mode for better concurrent read performance
    conn.execute('PRAGMA journal_mode=WAL')

    # Add new columns to quizzes table (safe to run multiple times)
    # NOTE: certification_id column intentionally removed — quizzes and certifications are independent systems
    try:
        c.execute('ALTER TABLE quizzes ADD COLUMN is_migrated BOOLEAN DEFAULT 0')
    except sqlite3.OperationalError:
        pass  # Column already exists
    try:
        c.execute('ALTER TABLE questions ADD COLUMN option_explanations TEXT')
    except sqlite3.OperationalError:
        pass  # Column already exists

    conn.commit()
    conn.close()

# === Question Normalization Helpers ===

def _insert_questions_for_quiz(c, quiz_id, questions_list):
    """Insert parsed question objects into the normalized questions table."""
    for idx, q in enumerate(questions_list):
        c.execute('''INSERT INTO questions
            (quiz_id, question_index, question_text, type, options, correct, pairs,
             code, code_language, image, image_alt, explanation, difficulty, option_explanations)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
            (quiz_id, idx,
             q.get('question', ''),
             q.get('type', 'choice'),
             json.dumps(q.get('options')) if q.get('options') is not None else None,
             json.dumps(q.get('correct')) if q.get('correct') is not None else None,
             json.dumps(q.get('pairs')) if q.get('pairs') is not None else None,
             q.get('code'),
             q.get('codeLanguage') or q.get('code_language'),
             q.get('image'),
             q.get('imageAlt') or q.get('image_alt'),
             q.get('explanation'),
             q.get('difficulty', 0),
             json.dumps(q.get('optionExplanations')) if q.get('optionExplanations') else None))

def _read_questions_for_quiz(c, quiz_id):
    """Read questions from normalized table and return as list of dicts matching frontend format."""
    c.execute('SELECT * FROM questions WHERE quiz_id = ? AND is_active = 1 ORDER BY question_index', (quiz_id,))
    rows = c.fetchall()
    questions = []
    for row in rows:
        q = {
            'id': row['id'],
            'question': row['question_text'],
            'type': row['type'],
        }
        if row['options'] is not None:
            q['options'] = json.loads(row['options'])
        if row['correct'] is not None:
            q['correct'] = json.loads(row['correct'])
        if row['pairs'] is not None:
            q['pairs'] = json.loads(row['pairs'])
        if row['code']:
            q['code'] = row['code']
        if row['code_language']:
            q['codeLanguage'] = row['code_language']
        if row['image']:
            q['image'] = row['image']
        if row['image_alt']:
            q['imageAlt'] = row['image_alt']
        if row['explanation']:
            q['explanation'] = row['explanation']
        if row['option_explanations']:
            q['optionExplanations'] = json.loads(row['option_explanations'])
        questions.append(q)
    return questions


def _migrate_quiz(c, quiz_id, questions_json_str):
    """Migrate a single quiz from JSON blob to normalized questions table."""
    try:
        questions_list = json.loads(questions_json_str)
    except (json.JSONDecodeError, TypeError):
        return False
    if not questions_list:
        return False
    _insert_questions_for_quiz(c, quiz_id, questions_list)
    c.execute('UPDATE quizzes SET is_migrated = 1 WHERE id = ?', (quiz_id,))
    return True

def _update_question_performance(c, user_id, quiz_id, answers_data, question_times=None):
    """Update question_performance table after an attempt.
    answers_data: dict mapping question_index (str) -> user's answer
    question_times: optional dict mapping question_index (str) -> time_ms
    """
    # Get question IDs for this quiz
    c.execute('SELECT id, question_index, correct, type, pairs FROM questions WHERE quiz_id = ? ORDER BY question_index', (quiz_id,))
    question_rows = c.fetchall()
    if not question_rows:
        return

    for qrow in question_rows:
        q_idx = str(qrow['question_index'])
        if q_idx not in answers_data:
            continue

        user_answer = answers_data[q_idx]
        correct_data = json.loads(qrow['correct']) if qrow['correct'] else []
        q_type = qrow['type']

        # Determine if answer is correct
        is_correct = _check_answer_correct(user_answer, correct_data, q_type, qrow)

        time_ms = question_times.get(q_idx) if question_times else None
        now = datetime.now()

        c.execute('''INSERT INTO question_performance
            (user_id, question_id, times_seen, times_correct, times_incorrect,
             last_seen_at, last_correct_at, average_time_ms)
            VALUES (?, ?, 1, ?, ?, ?, ?, ?)
            ON CONFLICT(user_id, question_id) DO UPDATE SET
            times_seen = times_seen + 1,
            times_correct = times_correct + ?,
            times_incorrect = times_incorrect + ?,
            last_seen_at = ?,
            last_correct_at = CASE WHEN ? = 1 THEN ? ELSE last_correct_at END,
            average_time_ms = CASE
                WHEN ? IS NOT NULL AND average_time_ms IS NOT NULL
                THEN (average_time_ms * (times_seen - 1) + ?) / times_seen
                WHEN ? IS NOT NULL THEN ?
                ELSE average_time_ms END,
            updated_at = ?''',
            (user_id, qrow['id'],
             1 if is_correct else 0,
             0 if is_correct else 1,
             now,
             now if is_correct else None,
             time_ms,
             # ON CONFLICT params
             1 if is_correct else 0,
             0 if is_correct else 1,
             now,
             1 if is_correct else 0, now,
             time_ms, time_ms,
             time_ms, time_ms,
             now))

def _check_answer_correct(user_answer, correct_data, q_type, qrow):
    """Check if a user's answer is correct for a given question."""
    if user_answer is None:
        return False
    if q_type == 'truefalse':
        # correct_data is [0] for True, [1] for False (0 = True index, 1 = False index)
        # user_answer is a Python bool (True/False) sent from the frontend
        if isinstance(correct_data, list) and len(correct_data) > 0:
            correct_bool = correct_data[0] == 0  # 0 means True is correct
            user_bool = bool(user_answer)
            return user_bool == correct_bool
        return False
    if q_type == 'choice':
        if isinstance(correct_data, list) and len(correct_data) > 0:
            if isinstance(user_answer, list):
                return set(user_answer) == set(correct_data)
            return user_answer == correct_data[0]
        return user_answer == correct_data
    elif q_type == 'matching':
        # Matching is correct if all pairs are matched
        if isinstance(user_answer, dict):
            pairs = json.loads(qrow['pairs']) if qrow['pairs'] else []
            return len(user_answer) == len(pairs)
        return False
    elif q_type == 'ordering':
        if isinstance(user_answer, list):
            expected = list(range(len(user_answer)))
            actual = [item.get('origIndex', i) if isinstance(item, dict) else item for i, item in enumerate(user_answer)]
            return actual == expected
        return False
    return False

# === Auth Helpers ===

# Try to use bcrypt for secure password hashing
try:
    import bcrypt
    BCRYPT_AVAILABLE = True
except ImportError:
    BCRYPT_AVAILABLE = False
    print("Warning: bcrypt not installed. Using PBKDF2 fallback. Install with: pip install bcrypt")

def hash_password(password, salt=None):
    """
    Hash password securely.
    - Uses bcrypt if available (recommended)
    - Falls back to PBKDF2 with SHA256 (still secure)
    """
    if BCRYPT_AVAILABLE:
        # bcrypt handles salt internally
        hashed = bcrypt.hashpw(password.encode(), bcrypt.gensalt(rounds=12))
        return hashed.decode(), 'bcrypt'
    else:
        # PBKDF2 fallback - much better than plain SHA256
        if salt is None:
            salt = secrets.token_hex(32)
        # 100,000 iterations of PBKDF2
        h = hashlib.pbkdf2_hmac('sha256', password.encode(), salt.encode(), 100000).hex()
        return h, salt

def verify_password(password, stored_hash, salt):
    """
    Verify password against stored hash.
    Supports bcrypt, PBKDF2, and legacy SHA256 (for existing users).
    """
    # Check for bcrypt first
    if salt == 'bcrypt' and BCRYPT_AVAILABLE:
        try:
            return bcrypt.checkpw(password.encode(), stored_hash.encode())
        except:
            return False
    
    # Try legacy SHA256 first (for existing users)
    # This was the original method: sha256(password + salt)
    legacy_hash = hashlib.sha256((password + salt).encode()).hexdigest()
    if secrets.compare_digest(legacy_hash, stored_hash):
        return True
    
    # Try PBKDF2 (for users created after the upgrade)
    if len(stored_hash) == 64 and len(salt) == 64:
        try:
            pbkdf2_hash = hashlib.pbkdf2_hmac('sha256', password.encode(), salt.encode(), 100000).hex()
            if secrets.compare_digest(pbkdf2_hash, stored_hash):
                return True
        except:
            pass
    
    return False

def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('Authorization', '')
        if token.startswith('Bearer '):
            token = token[7:]
        if not token:
            return jsonify({'error': 'Token missing'}), 401
        
        conn = get_db()
        c = conn.cursor()
        c.execute('''SELECT s.user_id, u.username, u.email FROM sessions s
            JOIN users u ON s.user_id = u.id
            WHERE s.token = ? AND s.expires_at > ? AND u.is_active = 1''',
            (token, datetime.now()))
        session = c.fetchone()
        conn.close()
        
        if not session:
            return jsonify({'error': 'Invalid token'}), 401
        
        request.user_id = session['user_id']
        request.username = session['username']
        return f(*args, **kwargs)
    return decorated

# === Auth Routes ===

def rate_limit(limit_str):
    """Decorator that applies rate limiting if flask-limiter is available."""
    def decorator(f):
        if LIMITER_AVAILABLE:
            return limiter.limit(limit_str)(f)
        return f
    return decorator

@app.route('/api/auth/register', methods=['POST'])
@rate_limit("5 per hour")
def register():
    data = request.get_json()
    username = data.get('username', '').strip()
    email = data.get('email', '').strip().lower()
    password = data.get('password', '')
    
    if not username or len(username) < 3:
        return jsonify({'error': 'Username must be 3+ characters'}), 400
    if not password or len(password) < 8:
        return jsonify({'error': 'Password must be 8+ characters'}), 400
    if not email:
        return jsonify({'error': 'Email address is required'}), 400

    h, salt = hash_password(password)
    conn = get_db()
    c = conn.cursor()

    try:
        verify_token = secrets.token_urlsafe(32)
        verify_expires = datetime.now() + timedelta(hours=24)
        c.execute(
            '''INSERT INTO users (username, email, password_hash, salt,
                                  email_verified, email_token, email_token_expires)
               VALUES (?, ?, ?, ?, 0, ?, ?)''',
            (username, email, h, salt, verify_token, verify_expires))
        user_id = c.lastrowid

        token = secrets.token_hex(32)
        c.execute('INSERT INTO sessions (user_id, token, expires_at) VALUES (?, ?, ?)',
            (user_id, token, datetime.now() + timedelta(days=7)))

        conn.commit()
        print(f"[DEV] Email verify token for {username}: {verify_token}")
        return jsonify({
            'token': token,
            'user': {'id': user_id, 'username': username, 'email_verified': False}
        }), 201
    except sqlite3.IntegrityError:
        return jsonify({'error': 'Username or email already taken'}), 409
    finally:
        conn.close()

@app.route('/api/auth/login', methods=['POST'])
@rate_limit("10 per minute")
def login():
    data = request.get_json()
    username = data.get('username', '').strip()
    password = data.get('password', '')
    
    conn = get_db()
    c = conn.cursor()
    c.execute(
        '''SELECT id, username, password_hash, salt, is_active,
                  COALESCE(email_verified, 0) as email_verified
           FROM users WHERE username = ?''', (username,))
    user = c.fetchone()

    if not user or not user['is_active'] or not verify_password(password, user['password_hash'], user['salt']):
        conn.close()
        return jsonify({'error': 'Invalid credentials'}), 401

    # Prune expired sessions for this user
    c.execute('DELETE FROM sessions WHERE user_id = ? AND expires_at < ?',
              (user['id'], datetime.now()))

    token = secrets.token_hex(32)
    c.execute('INSERT INTO sessions (user_id, token, expires_at) VALUES (?, ?, ?)',
        (user['id'], token, datetime.now() + timedelta(days=7)))
    c.execute('UPDATE users SET last_login = ? WHERE id = ?', (datetime.now(), user['id']))
    conn.commit()
    conn.close()

    return jsonify({
        'token': token,
        'user': {
            'id': user['id'],
            'username': user['username'],
            'email_verified': bool(user['email_verified'])
        }
    })

@app.route('/api/auth/resend-verification', methods=['POST'])
@token_required
def resend_verification():
    """Generate a new verification token and (dev) print it to the server log."""
    conn = get_db()
    c = conn.cursor()
    # Check if already verified
    c.execute('SELECT COALESCE(email_verified, 0) FROM users WHERE id = ?', (request.user_id,))
    row = c.fetchone()
    if row and row[0]:
        conn.close()
        return jsonify({'message': 'Email already verified'}), 200

    verify_token = secrets.token_urlsafe(32)
    verify_expires = datetime.now() + timedelta(hours=24)
    with _db_write_lock:
        c.execute(
            'UPDATE users SET email_token = ?, email_token_expires = ? WHERE id = ?',
            (verify_token, verify_expires, request.user_id))
        conn.commit()
    conn.close()
    print(f"[DEV] Email verify token for user {request.user_id}: {verify_token}")
    return jsonify({'message': 'Verification token generated (dev: check server log)'})


@app.route('/api/auth/verify-email', methods=['GET'])
def verify_email():
    """Mark an account as verified given a valid token (no auth required)."""
    token = request.args.get('token', '').strip()
    if not token:
        return jsonify({'error': 'Token required'}), 400

    conn = get_db()
    c = conn.cursor()
    c.execute(
        '''SELECT id FROM users
           WHERE email_token = ? AND email_token_expires > ?''',
        (token, datetime.now()))
    row = c.fetchone()
    if not row:
        conn.close()
        return jsonify({'error': 'Invalid or expired token'}), 400

    with _db_write_lock:
        c.execute(
            'UPDATE users SET email_verified = 1, email_token = NULL, email_token_expires = NULL WHERE id = ?',
            (row['id'],))
        conn.commit()
    conn.close()
    return jsonify({'message': 'Email verified successfully'})


# === Quiz Routes ===

@app.route('/api/quizzes', methods=['GET'])
@token_required
def get_quizzes():
    conn = get_db()
    c = conn.cursor()
    # Return owned quizzes + all public quizzes, owned ones first
    c.execute('''
        SELECT *, (user_id = ?) as is_owned
        FROM quizzes
        WHERE user_id = ? OR is_public = 1
        ORDER BY is_owned DESC, last_modified DESC
    ''', (request.user_id, request.user_id))
    quizzes = [dict(r) for r in c.fetchall()]
    for q in quizzes:
        # Dual-path: read from normalized table if migrated
        if q.get('is_migrated'):
            q['questions'] = _read_questions_for_quiz(c, q['id'])
        else:
            q['questions'] = json.loads(q['questions'])
    conn.close()
    return jsonify({'quizzes': quizzes})

@app.route('/api/quizzes', methods=['POST'])
@token_required
def create_quiz():
    data = request.get_json()
    title = data.get('title', '').strip()
    if not title:
        return jsonify({'error': 'Title required'}), 400

    questions_list = data.get('questions', [])
    conn = get_db()
    c = conn.cursor()
    # Dual-write: JSON blob + normalized questions table
    c.execute('INSERT INTO quizzes (user_id, title, description, questions, color, is_migrated) VALUES (?, ?, ?, ?, ?, 1)',
        (request.user_id, title, data.get('description', ''), json.dumps(questions_list),
         data.get('color', '#6366f1')))
    quiz_id = c.lastrowid
    _insert_questions_for_quiz(c, quiz_id, questions_list)
    conn.commit()
    conn.close()
    return jsonify({'message': 'Created', 'quiz_id': quiz_id}), 201


# === File Upload — Extract text from uploaded documents ===

UPLOAD_MAX_SIZE = 10 * 1024 * 1024  # 10 MB
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt', 'md', 'rtf', 'csv'}

def _extract_text_from_file(file_storage):
    """Extract text from an uploaded file.  Returns (text, error)."""
    filename = (file_storage.filename or '').lower()
    ext = filename.rsplit('.', 1)[-1] if '.' in filename else ''

    if ext not in ALLOWED_EXTENSIONS:
        return None, f'Unsupported file type: .{ext}. Supported: {", ".join(sorted(ALLOWED_EXTENSIONS))}'

    try:
        raw = file_storage.read()
    except Exception as e:
        return None, f'Could not read uploaded file stream: {str(e)}'

    if len(raw) == 0:
        return None, 'Uploaded file is empty.'
    if len(raw) > UPLOAD_MAX_SIZE:
        return None, f'File too large ({len(raw) // (1024*1024)}MB). Maximum is 10MB.'

    # Plain text variants
    if ext in ('txt', 'md', 'csv', 'rtf'):
        try:
            text = raw.decode('utf-8')
        except UnicodeDecodeError:
            try:
                text = raw.decode('latin-1')
            except Exception:
                return None, 'Could not decode text file.'
        return text.strip(), None

    # PDF
    if ext == 'pdf':
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(raw))
            pages = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
            if not pages:
                return None, 'Could not extract text from PDF. The file may be scanned/image-based.'
            return '\n\n'.join(pages).strip(), None
        except ImportError:
            return None, 'PDF support is not installed on the server.'
        except Exception as e:
            return None, f'Failed to read PDF: {str(e)}'

    # DOCX
    if ext in ('docx', 'doc'):
        try:
            from docx import Document
            doc = Document(io.BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            if not paragraphs:
                return None, 'No text found in document.'
            return '\n\n'.join(paragraphs).strip(), None
        except ImportError:
            return None, 'DOCX support is not installed on the server.'
        except Exception as e:
            return None, f'Failed to read document: {str(e)}'

    return None, 'Unsupported format.'


@app.route('/api/upload-material', methods=['POST'])
@token_required
def upload_material():
    """Extract text from uploaded file and return it for use in the wizard."""
    try:
        files = request.files
    except Exception as e:
        # Werkzeug raises BadRequest (HTTP 400) when multipart parsing fails,
        # which Flask would normally return as an HTML page.  Catch it here so
        # the frontend always receives JSON it can display to the user.
        app.logger.error(f'upload_material: failed to parse multipart body: {e}')
        return jsonify({
            'error': f'Failed to parse upload request: {str(e)}',
            'detail': (
                f'Content-Type received: {request.environ.get("CONTENT_TYPE", "not set")} — '
                'make sure the request is multipart/form-data with a file field named "file".'
            )
        }), 400

    if 'file' not in files:
        content_type = request.content_type or 'not set'
        form_keys = list(request.form.keys())
        app.logger.warning(
            f'upload_material: "file" field missing. '
            f'Content-Type={content_type!r}, form keys={form_keys}'
        )
        return jsonify({
            'error': 'No file received — the "file" field was not found in the request.',
            'detail': (
                f'Content-Type was {content_type!r}. '
                'This usually means the request was not sent as multipart/form-data, '
                'or the field name does not match "file".'
            )
        }), 400

    f = files['file']
    if not f.filename:
        return jsonify({'error': 'No filename — the file appears to have no name. Try re-selecting it.'}), 400

    try:
        text, error = _extract_text_from_file(f)
    except Exception as e:
        app.logger.error(f'upload_material: unexpected error in _extract_text_from_file: {e}', exc_info=True)
        return jsonify({'error': f'Unexpected error reading file: {str(e)}'}), 400

    if error:
        return jsonify({'error': error}), 400

    # Truncate very large texts to avoid overwhelming the AI (keep first ~80k chars)
    max_chars = 80_000
    truncated = len(text) > max_chars
    if truncated:
        text = text[:max_chars]

    return jsonify({
        'text': text,
        'char_count': len(text),
        'truncated': truncated,
        'filename': f.filename
    })


# === AI Quiz Generation ===

# Rate limit: max generations per user per hour
AI_RATE_LIMIT = 10
AI_MAX_QUESTIONS_PER_REQUEST = 300
# gpt-4.1-nano output cap is 32 768 tokens; ~150 tokens/question → ~200 max.
# 50 per batch is conservative and leaves headroom for verbose questions.
AI_BATCH_SIZE = 50
AI_MIN_MATERIAL_LENGTH = 100

def _check_ai_rate_limit(user_id):
    """Return (allowed, remaining, retry_after_seconds)."""
    conn = get_db()
    c = conn.cursor()
    one_hour_ago = datetime.now() - timedelta(hours=1)
    c.execute('SELECT COUNT(*) as cnt FROM ai_usage WHERE user_id = ? AND created_at > ?',
              (user_id, one_hour_ago))
    count = c.fetchone()['cnt']
    conn.close()
    if count >= AI_RATE_LIMIT:
        # Find oldest record in window to calculate retry-after
        conn = get_db()
        c = conn.cursor()
        c.execute('SELECT MIN(created_at) as oldest FROM ai_usage WHERE user_id = ? AND created_at > ?',
                  (user_id, one_hour_ago))
        oldest = c.fetchone()['oldest']
        conn.close()
        if oldest:
            oldest_dt = datetime.fromisoformat(oldest)
            retry_after = int((oldest_dt + timedelta(hours=1) - datetime.now()).total_seconds()) + 1
            return False, 0, max(retry_after, 1)
        return False, 0, 3600
    return True, AI_RATE_LIMIT - count, 0

def _log_ai_usage(user_id, input_tokens, output_tokens, question_count, model):
    """Record an AI generation in the usage table."""
    conn = get_db()
    c = conn.cursor()
    with _db_write_lock:
        c.execute('INSERT INTO ai_usage (user_id, input_tokens, output_tokens, question_count, model) VALUES (?, ?, ?, ?, ?)',
                  (user_id, input_tokens, output_tokens, question_count, model))
        conn.commit()
    conn.close()

def _build_ai_system_prompt():
    """Return the system prompt for quiz generation. Cached across requests."""
    return """You are a quiz question generator for educational study material.

DISTRACTOR DESIGN RULES (critical):
- Every wrong answer must be a real concept from the same domain as the correct answer
- Wrong answers should represent common misconceptions, adjacent facts, or plausible confusions
- If the correct answer is a specific term, wrong answers should be related terms from the same field — never absurd or obviously wrong options
- For numerical answers, wrong options should be plausible numbers (neighboring values, common confusions, off-by-one errors)
- Never include "All of the above" or "None of the above" as an option
- Each wrong answer should be wrong for a DIFFERENT reason — do not repeat the same type of mistake across distractors

EXPLANATION RULES:
- Every question MUST have an explanation
- The explanation should teach the underlying concept, not just restate the answer
- Explain WHY the correct answer is right
- Briefly note why each distractor is wrong

QUESTION QUALITY RULES:
- Derive questions strictly from the provided study material — do not invent facts
- Write clear, unambiguous question stems
- Ensure there is exactly one defensibly correct answer for single-choice questions
- Test understanding and application, not just recall
- Vary question difficulty naturally based on the complexity of the source material

You must output valid JSON matching the provided schema exactly. Do not include any text outside the JSON object."""

def _build_ai_user_prompt(study_material, question_count, question_types, category, include_code):
    """Build the user prompt with study material and preferences."""
    type_descriptions = {
        'choice': 'Multiple Choice (one correct answer from 4 options, "type": "choice")',
        'multiselect': 'Multi-Select (2+ correct answers from 4-5 options, "type": "choice" with multiple indices in "correct")',
        'truefalse': 'True/False (statement with True or False answer, "type": "truefalse")',
        'matching': 'Matching (4 term-definition pairs, "type": "matching")',
        'ordering': 'Ordering (4 items in correct sequence, "type": "ordering")',
    }
    requested_types = [type_descriptions[t] for t in question_types if t in type_descriptions]
    types_str = '\n'.join(f'- {t}' for t in requested_types)

    code_instruction = ""
    if include_code:
        code_instruction = """
When relevant to the material, include code snippets directly in the question text.
Format code questions like: "What does the following code output?" followed by the code in the question field.
Place the code in the "code" field and specify the language in "codeLanguage"."""

    return f"""Generate exactly {question_count} quiz questions from the study material below.

QUESTION TYPES TO USE (distribute roughly evenly, but prioritize types that fit the material):
{types_str}
{code_instruction}
{f'Subject area: {category}' if category else ''}

STUDY MATERIAL:
---
{study_material}
---"""

def _build_ai_json_schema(question_types):
    """Build the JSON schema for OpenAI Structured Outputs based on requested question types."""
    question_schema = {
        "type": "object",
        "properties": {
            "question": {
                "type": "string",
                "description": "The question text"
            },
            "type": {
                "type": "string",
                "enum": ["choice", "truefalse", "matching", "ordering"],
                "description": "Question type"
            },
            "options": {
                "type": "array",
                "items": {"type": "string"},
                "description": "Answer options. For truefalse: ['True', 'False']. For choice: 4 options. For ordering: items in correct order."
            },
            "correct": {
                "type": "array",
                "items": {"type": "integer"},
                "description": "Indices of correct options (0-based). Single-choice: one index. Multi-select: multiple. True/False: [0] for True, [1] for False. Ordering: sequential [0,1,2,3]."
            },
            "explanation": {
                "type": "string",
                "description": "Explanation of the correct answer and why distractors are wrong"
            },
            "pairs": {
                "type": ["array", "null"],
                "items": {
                    "type": "object",
                    "properties": {
                        "left": {"type": "string"},
                        "right": {"type": "string"}
                    },
                    "required": ["left", "right"],
                    "additionalProperties": False
                },
                "description": "Only for matching type: array of {left, right} pairs. null for other types."
            },
            "code": {
                "type": ["string", "null"],
                "description": "Code snippet if the question involves code. null otherwise."
            },
            "codeLanguage": {
                "type": ["string", "null"],
                "description": "Programming language of the code snippet (python, javascript, sql, etc). null if no code."
            }
        },
        "required": ["question", "type", "options", "correct", "explanation", "pairs", "code", "codeLanguage"],
        "additionalProperties": False
    }

    return {
        "name": "quiz_questions",
        "strict": True,
        "schema": {
            "type": "object",
            "properties": {
                "questions": {
                    "type": "array",
                    "items": question_schema
                }
            },
            "required": ["questions"],
            "additionalProperties": False
        }
    }

def _validate_generated_questions(questions):
    """Validate AI-generated questions, returning (valid_questions, warnings)."""
    valid = []
    warnings = []

    for i, q in enumerate(questions):
        q_num = i + 1
        q_type = q.get('type', 'choice')
        q_text = q.get('question', '').strip()

        if not q_text:
            warnings.append(f"Question {q_num}: empty question text, skipped")
            continue

        if q_type == 'matching':
            pairs = q.get('pairs')
            if not pairs or len(pairs) < 2:
                warnings.append(f"Question {q_num}: matching question needs at least 2 pairs, skipped")
                continue
            # Ensure options and correct are set from pairs
            q['options'] = [p['right'] for p in pairs]
            q['correct'] = list(range(len(pairs)))

        elif q_type == 'truefalse':
            q['options'] = ['True', 'False']
            correct = q.get('correct', [0])
            if not correct or correct[0] not in (0, 1):
                q['correct'] = [0]

        elif q_type == 'ordering':
            options = q.get('options', [])
            if len(options) < 2:
                warnings.append(f"Question {q_num}: ordering question needs at least 2 items, skipped")
                continue
            q['correct'] = list(range(len(options)))

        else:  # choice / multiselect
            options = q.get('options', [])
            correct = q.get('correct', [])
            if len(options) < 2:
                warnings.append(f"Question {q_num}: needs at least 2 options, skipped")
                continue
            if not correct:
                warnings.append(f"Question {q_num}: no correct answer marked, defaulting to first option")
                q['correct'] = [0]
            # Validate correct indices are in range
            q['correct'] = [c for c in correct if 0 <= c < len(options)]
            if not q['correct']:
                q['correct'] = [0]

        # Clean up null fields to match frontend expectations
        if not q.get('pairs'):
            q.pop('pairs', None)
        if not q.get('code'):
            q.pop('code', None)
            q.pop('codeLanguage', None)

        valid.append(q)

    return valid, warnings


@app.route('/api/generate-quiz', methods=['POST'])
@token_required
def generate_quiz_ai():
    """Generate quiz questions from study material using AI."""

    # Check if OpenAI API key is configured
    api_key = os.environ.get('OPENAI_API_KEY')
    if not api_key:
        return jsonify({'error': 'AI generation is not configured. Set the OPENAI_API_KEY environment variable.'}), 503

    # Check rate limit
    allowed, remaining, retry_after = _check_ai_rate_limit(request.user_id)
    if not allowed:
        return jsonify({
            'error': f'Rate limit exceeded. You can generate up to {AI_RATE_LIMIT} quizzes per hour.',
            'retry_after': retry_after
        }), 429

    data = request.get_json()
    if not data:
        return jsonify({'error': 'Request body required'}), 400

    study_material = data.get('study_material', '').strip()
    question_count = data.get('question_count', 15)
    question_types = data.get('question_types', ['choice'])
    category = data.get('category', '').strip()
    include_code = data.get('include_code', False)

    # Validate study material
    if len(study_material) < AI_MIN_MATERIAL_LENGTH:
        return jsonify({
            'error': f'Study material is too short. Please provide at least {AI_MIN_MATERIAL_LENGTH} characters of content.'
        }), 400

    # Clamp question count
    question_count = max(1, min(question_count, AI_MAX_QUESTIONS_PER_REQUEST))

    # Filter to valid question types
    valid_types = {'choice', 'multiselect', 'truefalse', 'matching', 'ordering'}
    question_types = [t for t in question_types if t in valid_types]
    if not question_types:
        question_types = ['choice']

    # Build shared prompts/schema
    system_prompt = _build_ai_system_prompt()
    json_schema = _build_ai_json_schema(question_types)

    try:
        import openai
    except ImportError:
        return jsonify({'error': 'AI generation is not configured. The openai package is not installed.'}), 503

    model = 'gpt-4.1-nano'
    # max_retries=0: don't let the openai client retry 429/5xx internally —
    # that holds the WSGI worker open until PythonAnywhere's harakiri kills it (502).
    client = openai.OpenAI(api_key=api_key, max_retries=0)

    # ---- Batch generation ------------------------------------------------
    # gpt-4.1-nano's output cap is ~32 768 tokens. At ~150 tokens/question
    # we can fit ~200 questions, but we use AI_BATCH_SIZE=50 to be safe and
    # to give the model focused context for each batch.
    all_questions = []
    total_input_tokens = 0
    total_output_tokens = 0
    truncated = False  # set True if a mid-run rate-limit cut batches short

    remaining_questions = question_count
    while remaining_questions > 0:
        batch_size = min(remaining_questions, AI_BATCH_SIZE)
        # ~160 tokens per question + small buffer; stay well under the model cap
        batch_max_tokens = min(batch_size * 160 + 500, 16000)
        user_prompt = _build_ai_user_prompt(
            study_material, batch_size, question_types, category, include_code
        )

        try:
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                response_format={"type": "json_schema", "json_schema": json_schema},
                max_tokens=batch_max_tokens,
                temperature=0.7,
            )
        except openai.AuthenticationError:
            return jsonify({'error': 'AI service authentication failed. Check your API key configuration.'}), 503
        except openai.RateLimitError:
            if all_questions:
                # Return however many questions we managed to generate
                truncated = True
                break
            return jsonify({'error': 'AI service is temporarily busy. Please try again in a moment.'}), 503
        except openai.BadRequestError as e:
            msg = str(e)
            print(f"[AI] OpenAI bad request for user {request.user_id}: {msg}", flush=True)
            if all_questions:
                truncated = True
                break
            return jsonify({'error': 'AI rejected the request. Try fewer questions or shorter study material.'}), 400
        except openai.APITimeoutError:
            if all_questions:
                truncated = True
                break
            return jsonify({'error': 'AI generation timed out. Try fewer questions or shorter study material.'}), 504
        except openai.APIConnectionError:
            return jsonify({'error': 'Could not connect to AI service. Please try again.'}), 503
        except openai.APIError as e:
            print(f"[AI] OpenAI API error for user {request.user_id}: {e}", flush=True)
            if all_questions:
                truncated = True
                break
            return jsonify({'error': 'AI service encountered an error. Please try again.'}), 503

        usage = response.usage
        total_input_tokens += usage.prompt_tokens if usage else 0
        total_output_tokens += usage.completion_tokens if usage else 0

        try:
            result = json.loads(response.choices[0].message.content)
            batch_questions = result.get('questions', [])
        except (json.JSONDecodeError, TypeError, AttributeError):
            batch_questions = []

        all_questions.extend(batch_questions)
        remaining_questions -= batch_size
    # ---- End batch loop --------------------------------------------------

    input_tokens = total_input_tokens
    output_tokens = total_output_tokens

    # Validate all collected questions
    valid_questions, warnings = _validate_generated_questions(all_questions)

    if not valid_questions:
        _log_ai_usage(request.user_id, input_tokens, output_tokens, 0, model)
        return jsonify({
            'error': 'AI could not generate valid questions from this material. Try adding more detailed notes or definitions.',
            'warnings': warnings
        }), 422

    if truncated:
        warnings.append(
            f'Generation was cut short due to an API limit. '
            f'Returning {len(valid_questions)} of {question_count} requested questions.'
        )

    # Log successful usage
    _log_ai_usage(request.user_id, input_tokens, output_tokens, len(valid_questions), model)

    response_data = {
        'questions': valid_questions,
        'count': len(valid_questions),
        'requested_count': question_count,
        'model': model,
        'usage': {
            'input_tokens': input_tokens,
            'output_tokens': output_tokens
        },
        'rate_limit': {
            'remaining': remaining - 1,
            'limit': AI_RATE_LIMIT,
        }
    }

    if warnings:
        response_data['warnings'] = warnings

    return jsonify(response_data), 200


@app.route('/api/quizzes/<int:id>', methods=['GET'])
@token_required
def get_quiz(id):
    conn = get_db()
    c = conn.cursor()
    # Allow fetching owned quizzes OR public quizzes (for community study)
    c.execute('SELECT *, (user_id = ?) as is_owned FROM quizzes WHERE id = ? AND (user_id = ? OR is_public = 1)',
              (request.user_id, id, request.user_id))
    quiz = c.fetchone()

    if not quiz:
        conn.close()
        return jsonify({'error': 'Not found'}), 404

    q = dict(quiz)
    # Dual-path: read from normalized table if migrated, else JSON blob
    if q.get('is_migrated'):
        q['questions'] = _read_questions_for_quiz(c, id)
    else:
        q['questions'] = json.loads(q['questions'])
    conn.close()
    return jsonify({'quiz': q})

@app.route('/api/quizzes/<int:id>', methods=['PUT'])
@token_required
def update_quiz(id):
    data = request.get_json()
    questions_list = data.get('questions', [])
    conn = get_db()
    c = conn.cursor()
    # Dual-write: update JSON blob AND normalized questions
    c.execute('UPDATE quizzes SET title=?, description=?, questions=?, color=?, is_public=?, is_migrated=1, last_modified=? WHERE id=? AND user_id=?',
        (data.get('title'), data.get('description', ''), json.dumps(questions_list),
         data.get('color', '#6366f1'),
         1 if data.get('is_public') else 0,
         datetime.now(), id, request.user_id))
    # Replace normalized questions: delete old, insert new
    c.execute('DELETE FROM questions WHERE quiz_id = ?', (id,))
    _insert_questions_for_quiz(c, id, questions_list)
    conn.commit()
    conn.close()
    return jsonify({'message': 'Updated'})

@app.route('/api/quizzes/<int:id>/settings', methods=['PATCH'])
@token_required
def update_quiz_settings(id):
    """Lightweight update: is_public only."""
    data = request.get_json() or {}
    conn = get_db()
    c = conn.cursor()
    with _db_write_lock:
        c.execute('''UPDATE quizzes
                     SET is_public=?, last_modified=?
                     WHERE id=? AND user_id=?''',
                  (1 if data.get('is_public') else 0,
                   datetime.now(), id, request.user_id))
        if c.rowcount == 0:
            conn.close()
            return jsonify({'error': 'Not found or not authorized'}), 404
        conn.commit()
    conn.close()
    return jsonify({'message': 'Settings updated'})

@app.route('/api/quizzes/<int:id>', methods=['DELETE'])
@token_required
def delete_quiz(id):
    conn = get_db()
    c = conn.cursor()
    c.execute('DELETE FROM quizzes WHERE id = ? AND user_id = ?', (id, request.user_id))
    conn.commit()
    conn.close()
    return jsonify({'message': 'Deleted'})

@app.route('/api/quizzes/<int:id>/attempts', methods=['POST'])
@token_required
def record_attempt(id):
    data = request.get_json()
    conn = get_db()
    c = conn.cursor()
    c.execute('''INSERT INTO attempts (quiz_id, user_id, score, total, percentage, answers, study_mode, timed, max_streak, time_taken)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
        (id, request.user_id, data.get('score', 0), data.get('total', 0), data.get('percentage', 0),
         json.dumps(data.get('answers', {})), data.get('study_mode', False), data.get('timed', False),
         data.get('max_streak', 0), data.get('time_taken')))

    # Update per-question performance tracking
    try:
        answers_data = data.get('answers', {})
        question_times = data.get('question_times', {})
        if isinstance(answers_data, dict):
            _update_question_performance(c, request.user_id, id, answers_data, question_times or None)
    except Exception as e:
        print(f"Warning: question_performance update failed: {e}")

    conn.commit()
    conn.close()
    return jsonify({'message': 'Recorded'}), 201

# === Profile & Stats (Synced to Database) ===

@app.route('/api/profile', methods=['GET'])
@token_required
def get_profile():
    conn = get_db()
    c = conn.cursor()
    
    # Get user info
    c.execute('SELECT id, username, email, created_at FROM users WHERE id = ?', (request.user_id,))
    user = dict(c.fetchone())
    
    # Get quiz count
    c.execute('SELECT COUNT(*) as count FROM quizzes WHERE user_id = ?', (request.user_id,))
    user['quiz_count'] = c.fetchone()['count']
    
    # Try to get or create profile (handle missing table gracefully)
    try:
        c.execute('SELECT * FROM user_profiles WHERE user_id = ?', (request.user_id,))
        profile_row = c.fetchone()
        
        if profile_row:
            profile = dict(profile_row)
            profile['achievements'] = json.loads(profile['achievements'] or '[]')
            profile['settings'] = json.loads(profile['settings'] or '{}')
        else:
            # Create default profile
            c.execute('''INSERT INTO user_profiles (user_id) VALUES (?)''', (request.user_id,))
            conn.commit()
            profile = {
                'xp': 0, 'level': 1, 'gems': 0, 'daily_streak': 0,
                'last_active_date': None, 'achievements': [],
                'total_answered': 0, 'total_correct': 0,
                'quizzes_completed': 0, 'perfect_scores': 0,
                'settings': {'soundEnabled': True, 'animationsEnabled': True}
            }
    except sqlite3.OperationalError as e:
        # Table doesn't exist - return defaults
        print(f"Profile table not found, returning defaults: {e}")
        profile = {
            'xp': 0, 'level': 1, 'gems': 0, 'daily_streak': 0,
            'last_active_date': None, 'achievements': [],
            'total_answered': 0, 'total_correct': 0,
            'quizzes_completed': 0, 'perfect_scores': 0,
            'settings': {'soundEnabled': True, 'animationsEnabled': True}
        }
    
    conn.close()
    
    return jsonify({
        'user': user,
        'profile': profile
    })

@app.route('/api/profile', methods=['PUT'])
@token_required
def update_profile():
    data = request.get_json()
    conn = get_db()
    c = conn.cursor()
    
    # Upsert profile
    c.execute('''INSERT INTO user_profiles (user_id, xp, level, gems, daily_streak, last_active_date,
                 achievements, total_answered, total_correct, quizzes_completed, perfect_scores, settings, updated_at)
                 VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                 ON CONFLICT(user_id) DO UPDATE SET
                 xp=excluded.xp, level=excluded.level, gems=excluded.gems,
                 daily_streak=excluded.daily_streak, last_active_date=excluded.last_active_date,
                 achievements=excluded.achievements, total_answered=excluded.total_answered,
                 total_correct=excluded.total_correct, quizzes_completed=excluded.quizzes_completed,
                 perfect_scores=excluded.perfect_scores, settings=excluded.settings, updated_at=excluded.updated_at''',
        (request.user_id,
         data.get('xp', 0),
         data.get('level', 1),
         data.get('gems', 0),
         data.get('daily_streak', 0),
         data.get('last_active_date'),
         json.dumps(data.get('achievements', [])),
         data.get('total_answered', 0),
         data.get('total_correct', 0),
         data.get('quizzes_completed', 0),
         data.get('perfect_scores', 0),
         json.dumps(data.get('settings', {})),
         datetime.now()))
    
    conn.commit()
    conn.close()
    return jsonify({'message': 'Profile updated'})

@app.route('/api/stats', methods=['GET'])
@token_required
def get_stats():
    conn = get_db()
    c = conn.cursor()
    
    # Recent attempts (last 10)
    c.execute('''
        SELECT a.*, q.title as quiz_title 
        FROM attempts a 
        JOIN quizzes q ON a.quiz_id = q.id 
        WHERE a.user_id = ? 
        ORDER BY a.created_at DESC 
        LIMIT 10
    ''', (request.user_id,))
    recent = [dict(r) for r in c.fetchall()]
    
    # Best scores per quiz
    c.execute('''
        SELECT q.id, q.title, MAX(a.percentage) as best_score, COUNT(a.id) as attempts
        FROM quizzes q
        LEFT JOIN attempts a ON q.id = a.quiz_id
        WHERE q.user_id = ?
        GROUP BY q.id
        ORDER BY q.last_modified DESC
    ''', (request.user_id,))
    quizzes = [dict(r) for r in c.fetchall()]
    
    conn.close()
    return jsonify({
        'recent_attempts': recent,
        'quiz_stats': quizzes
    })

# === Quiz Progress Sync ===

@app.route('/api/progress/<int:quiz_id>', methods=['GET'])
@token_required
def get_quiz_progress(quiz_id):
    conn = get_db()
    c = conn.cursor()
    
    c.execute('SELECT * FROM quiz_progress WHERE user_id = ? AND quiz_id = ?', 
              (request.user_id, quiz_id))
    row = c.fetchone()
    conn.close()
    
    if not row:
        return jsonify({'progress': None})
    
    progress = dict(row)
    progress['answers'] = json.loads(progress['answers'] or '[]')
    progress['flagged'] = json.loads(progress['flagged'] or '[]')
    progress['option_shuffles'] = json.loads(progress['option_shuffles'] or '{}')
    
    return jsonify({'progress': progress})

@app.route('/api/progress/<int:quiz_id>', methods=['PUT'])
@token_required
def save_quiz_progress(quiz_id):
    data = request.get_json()
    conn = get_db()
    c = conn.cursor()
    
    c.execute('''INSERT INTO quiz_progress 
                 (user_id, quiz_id, question_index, answers, flagged, study_mode, 
                  randomize_options, option_shuffles, quiz_streak, max_quiz_streak,
                  timer_enabled, time_remaining, updated_at)
                 VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                 ON CONFLICT(user_id, quiz_id) DO UPDATE SET
                 question_index=excluded.question_index, answers=excluded.answers,
                 flagged=excluded.flagged, study_mode=excluded.study_mode,
                 randomize_options=excluded.randomize_options, option_shuffles=excluded.option_shuffles,
                 quiz_streak=excluded.quiz_streak, max_quiz_streak=excluded.max_quiz_streak,
                 timer_enabled=excluded.timer_enabled, time_remaining=excluded.time_remaining,
                 updated_at=excluded.updated_at''',
        (request.user_id, quiz_id,
         data.get('question_index', 0),
         json.dumps(data.get('answers', [])),
         json.dumps(data.get('flagged', [])),
         data.get('study_mode', True),
         data.get('randomize_options', False),
         json.dumps(data.get('option_shuffles', {})),
         data.get('quiz_streak', 0),
         data.get('max_quiz_streak', 0),
         data.get('timer_enabled', False),
         data.get('time_remaining'),
         datetime.now()))
    
    conn.commit()
    conn.close()
    return jsonify({'message': 'Progress saved'})

@app.route('/api/progress/<int:quiz_id>', methods=['DELETE'])
@token_required
def clear_quiz_progress(quiz_id):
    conn = get_db()
    c = conn.cursor()
    c.execute('DELETE FROM quiz_progress WHERE user_id = ? AND quiz_id = ?',
              (request.user_id, quiz_id))
    conn.commit()
    conn.close()
    return jsonify({'message': 'Progress cleared'})

@app.route('/api/progress', methods=['GET'])
@token_required
def get_all_progress():
    """Get all in-progress quizzes for the library display."""
    conn = get_db()
    c = conn.cursor()
    
    try:
        c.execute('''
            SELECT p.*, q.title as quiz_title,
                   CASE WHEN q.is_migrated = 1
                        THEN (SELECT COUNT(*) FROM questions WHERE quiz_id = q.id AND is_active = 1)
                        ELSE json_array_length(q.questions)
                   END as total_questions
            FROM quiz_progress p
            JOIN quizzes q ON p.quiz_id = q.id
            WHERE p.user_id = ?
            ORDER BY p.updated_at DESC
        ''', (request.user_id,))
        
        rows = c.fetchall()
        conn.close()
        
        progress_list = []
        for row in rows:
            p = dict(row)
            p['answers'] = json.loads(p['answers'] or '[]')
            p['flagged'] = json.loads(p['flagged'] or '[]')
            progress_list.append(p)
        
        return jsonify({'progress': progress_list})
    except sqlite3.OperationalError as e:
        # Table doesn't exist yet
        conn.close()
        print(f"Progress table not found: {e}")
        return jsonify({'progress': []})

# === Migration & Certification Routes ===

@app.route('/api/migrate', methods=['POST'])
@token_required
def migrate_quizzes():
    """Migrate all non-migrated quizzes from JSON blob to normalized questions table."""
    conn = get_db()
    c = conn.cursor()
    c.execute('SELECT id, questions FROM quizzes WHERE user_id = ? AND is_migrated = 0', (request.user_id,))
    rows = c.fetchall()
    migrated = 0
    for row in rows:
        if _migrate_quiz(c, row['id'], row['questions']):
            migrated += 1
    conn.commit()
    conn.close()
    return jsonify({'message': f'Migrated {migrated} quizzes', 'migrated': migrated})

@app.route('/api/certifications', methods=['GET'])
def get_certifications():
    conn = get_db()
    c = conn.cursor()
    try:
        c.execute('SELECT * FROM certifications WHERE is_active = 1 ORDER BY vendor, name')
        certs = [dict(r) for r in c.fetchall()]
    except sqlite3.OperationalError:
        certs = []
    conn.close()
    return jsonify({'certifications': certs})

@app.route('/api/certifications/<int:id>', methods=['GET'])
def get_certification(id):
    conn = get_db()
    c = conn.cursor()
    try:
        c.execute('SELECT * FROM certifications WHERE id = ?', (id,))
        cert = c.fetchone()
        if not cert:
            conn.close()
            return jsonify({'error': 'Not found'}), 404
        cert = dict(cert)
        c.execute('SELECT * FROM domains WHERE certification_id = ? ORDER BY sort_order, code', (id,))
        cert['domains'] = [dict(r) for r in c.fetchall()]
    except sqlite3.OperationalError:
        conn.close()
        return jsonify({'error': 'Not found'}), 404
    conn.close()
    return jsonify({'certification': cert})

@app.route('/api/certifications', methods=['POST'])
@token_required
def create_certification():
    data = request.get_json()
    conn = get_db()
    c = conn.cursor()
    try:
        c.execute('''INSERT INTO certifications (code, name, vendor, description, passing_score, passing_scale, exam_duration_minutes, total_questions)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
            (data['code'], data['name'], data['vendor'], data.get('description'),
             data.get('passing_score'), data.get('passing_scale', 'percentage'),
             data.get('exam_duration_minutes'), data.get('total_questions')))
        cert_id = c.lastrowid

        # Insert domains if provided
        for i, domain in enumerate(data.get('domains', [])):
            c.execute('''INSERT INTO domains (certification_id, name, code, weight, sort_order)
                VALUES (?, ?, ?, ?, ?)''',
                (cert_id, domain['name'], domain.get('code'), domain.get('weight'), i))

        conn.commit()
        return jsonify({'message': 'Created', 'certification_id': cert_id}), 201
    except sqlite3.IntegrityError:
        return jsonify({'error': 'Certification code already exists'}), 409
    finally:
        conn.close()

# === Phase 2: User Certifications, Domain Performance, Exam Simulation ===

@app.route('/api/user-certifications', methods=['GET'])
@token_required
def get_user_certifications():
    conn = get_db()
    c = conn.cursor()
    try:
        c.execute('''SELECT uc.*, c.name, c.vendor, c.code, c.passing_score, c.passing_scale,
                      c.exam_duration_minutes, c.total_questions
                FROM user_certifications uc
                JOIN certifications c ON uc.certification_id = c.id
                WHERE uc.user_id = ?
                ORDER BY uc.started_at DESC''', (request.user_id,))
        rows = [dict(r) for r in c.fetchall()]
    except sqlite3.OperationalError:
        rows = []
    conn.close()
    return jsonify({'user_certifications': rows})

@app.route('/api/user-certifications', methods=['POST'])
@token_required
def enroll_certification():
    data = request.get_json()
    conn = get_db()
    c = conn.cursor()
    try:
        c.execute('''INSERT INTO user_certifications (user_id, certification_id, target_date)
            VALUES (?, ?, ?)
            ON CONFLICT(user_id, certification_id) DO UPDATE SET
            target_date = excluded.target_date, status = 'studying' ''',
            (request.user_id, data['certification_id'], data.get('target_date')))
        conn.commit()
        return jsonify({'message': 'Enrolled'}), 201
    except Exception as e:
        return jsonify({'error': str(e)}), 400
    finally:
        conn.close()

@app.route('/api/user-certifications/<int:cert_id>', methods=['DELETE'])
@token_required
def unenroll_certification(cert_id):
    conn = get_db()
    c = conn.cursor()
    c.execute('DELETE FROM user_certifications WHERE user_id = ? AND certification_id = ?',
              (request.user_id, cert_id))
    conn.commit()
    conn.close()
    return jsonify({'message': 'Unenrolled'})


@app.route('/api/certifications/<int:cert_id>/performance', methods=['GET'])
@token_required
def get_cert_performance(cert_id):
    """Get domain-level performance breakdown for a certification."""
    conn = get_db()
    c = conn.cursor()

    # Get top-level domains only (exclude sub-domains)
    c.execute('SELECT * FROM domains WHERE certification_id = ? AND parent_domain_id IS NULL ORDER BY sort_order', (cert_id,))
    domains = [dict(r) for r in c.fetchall()]

    for domain in domains:
        c.execute('''SELECT
                COALESCE(SUM(qp.times_seen), 0) as total_seen,
                COALESCE(SUM(qp.times_correct), 0) as total_correct,
                COALESCE(SUM(qp.times_incorrect), 0) as total_incorrect,
                MAX(qp.last_seen_at) as last_studied,
                COUNT(DISTINCT qp.question_id) as unique_questions
            FROM question_performance qp
            JOIN question_domains qd ON qp.question_id = qd.question_id
            WHERE qp.user_id = ? AND qd.domain_id = ?''',
            (request.user_id, domain['id']))
        stats = dict(c.fetchone())
        domain['total_seen'] = stats['total_seen']
        domain['total_correct'] = stats['total_correct']
        domain['total_incorrect'] = stats['total_incorrect']
        domain['accuracy'] = round(stats['total_correct'] / stats['total_seen'] * 100, 1) if stats['total_seen'] > 0 else 0
        domain['last_studied'] = stats['last_studied']
        domain['unique_questions'] = stats['unique_questions']

    conn.close()
    return jsonify({'domains': domains})

@app.route('/api/certifications/<int:cert_id>/trends', methods=['GET'])
@token_required
def get_cert_trends(cert_id):
    """Get score trends over time for exam simulations."""
    conn = get_db()
    c = conn.cursor()

    c.execute('''SELECT id, score, total, percentage, passed, time_taken, domain_scores, created_at
        FROM exam_simulations
        WHERE user_id = ? AND certification_id = ?
        ORDER BY created_at DESC
        LIMIT 20''', (request.user_id, cert_id))
    sims = []
    for row in c.fetchall():
        s = dict(row)
        s['domain_scores'] = json.loads(s['domain_scores'])
        sims.append(s)

    conn.close()
    return jsonify({'simulations': sims})

@app.route('/api/questions/weak', methods=['GET'])
@token_required
def get_weak_questions():
    """Get questions the user struggles with most."""
    limit = request.args.get('limit', 20, type=int)
    cert_id = request.args.get('certification_id', type=int)
    conn = get_db()
    c = conn.cursor()

    query = '''SELECT qp.*, q.question_text, q.type, q.options, q.explanation, q.quiz_id,
                      qz.title as quiz_title
        FROM question_performance qp
        JOIN questions q ON qp.question_id = q.id
        JOIN quizzes qz ON q.quiz_id = qz.id
        WHERE qp.user_id = ? AND qp.times_seen >= 2 AND qp.times_incorrect > 0'''
    params = [request.user_id]

    if cert_id:
        query += ''' AND EXISTS (
            SELECT 1 FROM question_domains qd
            JOIN domains d ON qd.domain_id = d.id
            WHERE qd.question_id = qp.question_id AND d.certification_id = ?)'''
        params.append(cert_id)

    query += ' ORDER BY CAST(qp.times_incorrect AS REAL) / qp.times_seen DESC LIMIT ?'
    params.append(limit)

    c.execute(query, params)
    questions = []
    for row in c.fetchall():
        q = dict(row)
        if q.get('options'):
            q['options'] = json.loads(q['options'])
        questions.append(q)

    conn.close()
    return jsonify({'questions': questions})

@app.route('/api/certifications/<int:cert_id>/simulate', methods=['POST'])
@token_required
def start_exam_simulation(cert_id):
    """Generate an exam simulation with domain-weighted question selection."""
    conn = get_db()
    c = conn.cursor()

    # Get cert info
    c.execute('SELECT * FROM certifications WHERE id = ?', (cert_id,))
    cert = c.fetchone()
    if not cert:
        conn.close()
        return jsonify({'error': 'Certification not found'}), 404
    cert = dict(cert)

    # Get domains with weights
    c.execute('SELECT * FROM domains WHERE certification_id = ? ORDER BY sort_order', (cert_id,))
    domains = [dict(r) for r in c.fetchall()]

    total_questions = cert.get('total_questions') or 60
    data = request.get_json() or {}
    requested_count = data.get('question_count') or total_questions

    import random
    selected_questions = []

    for domain in domains:
        weight = domain.get('weight') or (1.0 / len(domains))
        domain_count = max(1, round(requested_count * weight))

        # Get questions tagged with this domain
        c.execute('''SELECT q.* FROM questions q
            JOIN question_domains qd ON q.id = qd.question_id
            WHERE qd.domain_id = ? AND q.is_active = 1''', (domain['id'],))
        domain_questions = [dict(r) for r in c.fetchall()]

        if domain_questions:
            chosen = random.sample(domain_questions, min(domain_count, len(domain_questions)))
            for q in chosen:
                q['domain_id'] = domain['id']
                q['domain_name'] = domain['name']
                q['domain_code'] = domain['code']
            selected_questions.extend(chosen)

    random.shuffle(selected_questions)

    # Convert to frontend format
    sim_questions = []
    for row in selected_questions:
        q = {
            'id': row['id'],
            'question': row['question_text'],
            'type': row['type'],
        }
        if row.get('options'):
            q['options'] = json.loads(row['options'])
        if row.get('correct'):
            q['correct'] = json.loads(row['correct'])
        if row.get('pairs'):
            q['pairs'] = json.loads(row['pairs'])
        if row.get('code'):
            q['code'] = row['code']
        if row.get('code_language'):
            q['codeLanguage'] = row['code_language']
        if row.get('explanation'):
            q['explanation'] = row['explanation']
        if row.get('domain_name'):
            q['domainName'] = row['domain_name']
            q['domainCode'] = row.get('domain_code')
        sim_questions.append(q)

    conn.close()

    return jsonify({
        'simulation': {
            'certification': cert,
            'questions': sim_questions,
            'time_limit': (cert.get('exam_duration_minutes') or 90) * 60,
            'passing_score': cert.get('passing_score'),
            'passing_scale': cert.get('passing_scale'),
            'total_questions': len(sim_questions),
            'domains': domains
        }
    })

@app.route('/api/exam-simulations', methods=['POST'])
@token_required
def record_exam_simulation():
    """Record a completed exam simulation."""
    data = request.get_json()
    conn = get_db()
    c = conn.cursor()
    c.execute('''INSERT INTO exam_simulations
        (user_id, certification_id, score, total, percentage, passed, time_taken, time_limit, domain_scores, answers)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''',
        (request.user_id, data['certification_id'], data['score'], data['total'],
         data['percentage'], data['passed'], data.get('time_taken'),
         data.get('time_limit'), json.dumps(data.get('domain_scores', {})),
         json.dumps(data.get('answers', []))))

    # Also update question_performance for each answered question
    try:
        answers_data = data.get('answers_detail', {})
        question_times = data.get('question_times', {})
        if isinstance(answers_data, dict):
            # For simulations, we need to map by question_id not index
            for q_id_str, user_answer in answers_data.items():
                q_id = int(q_id_str)
                c.execute('SELECT correct, type, pairs FROM questions WHERE id = ?', (q_id,))
                qrow = c.fetchone()
                if not qrow:
                    continue
                correct_data = json.loads(qrow['correct']) if qrow['correct'] else []
                is_correct = _check_answer_correct(user_answer, correct_data, qrow['type'], qrow)
                time_ms = question_times.get(q_id_str)
                now = datetime.now()
                c.execute('''INSERT INTO question_performance
                    (user_id, question_id, times_seen, times_correct, times_incorrect, last_seen_at, last_correct_at, average_time_ms)
                    VALUES (?, ?, 1, ?, ?, ?, ?, ?)
                    ON CONFLICT(user_id, question_id) DO UPDATE SET
                    times_seen = times_seen + 1,
                    times_correct = times_correct + ?,
                    times_incorrect = times_incorrect + ?,
                    last_seen_at = ?,
                    last_correct_at = CASE WHEN ? = 1 THEN ? ELSE last_correct_at END,
                    updated_at = ?''',
                    (request.user_id, q_id,
                     1 if is_correct else 0, 0 if is_correct else 1,
                     now, now if is_correct else None, time_ms,
                     1 if is_correct else 0, 0 if is_correct else 1,
                     now, 1 if is_correct else 0, now, now))
    except Exception as e:
        print(f"Warning: sim question_performance update failed: {e}")

    conn.commit()
    conn.close()
    return jsonify({'message': 'Simulation recorded'}), 201

# === Study Sessions ===

@app.route('/api/study-sessions', methods=['POST'])
@token_required
def create_study_session():
    """Start a new study session."""
    data = request.get_json() or {}
    conn = get_db()
    c = conn.cursor()
    c.execute('''INSERT INTO study_sessions (user_id, session_type, quiz_id, certification_id)
        VALUES (?, ?, ?, ?)''',
        (request.user_id, data.get('session_type', 'quiz'),
         data.get('quiz_id'), data.get('certification_id') if data.get('session_type') != 'quiz' else None))
    session_id = c.lastrowid
    conn.commit()
    conn.close()
    return jsonify({'session_id': session_id}), 201

@app.route('/api/study-sessions/<int:session_id>', methods=['PUT'])
@token_required
def end_study_session(session_id):
    """End a study session with stats."""
    data = request.get_json() or {}
    conn = get_db()
    c = conn.cursor()
    c.execute('''UPDATE study_sessions SET ended_at = ?, questions_reviewed = ?,
        questions_correct = ?, duration_seconds = ? WHERE id = ? AND user_id = ?''',
        (datetime.now(), data.get('questions_reviewed', 0),
         data.get('questions_correct', 0), data.get('duration_seconds', 0),
         session_id, request.user_id))
    conn.commit()
    conn.close()
    return jsonify({'message': 'Session ended'})

@app.route('/api/study-sessions/stats', methods=['GET'])
@token_required
def get_study_stats():
    """Get aggregated study statistics."""
    period = request.args.get('period', 'week')
    conn = get_db()
    c = conn.cursor()

    if period == 'week':
        cutoff = datetime.now() - timedelta(days=7)
    elif period == 'month':
        cutoff = datetime.now() - timedelta(days=30)
    else:
        cutoff = datetime(2000, 1, 1)

    # Aggregate stats
    c.execute('''SELECT
        COUNT(*) as total_sessions,
        COALESCE(SUM(duration_seconds), 0) as total_seconds,
        COALESCE(SUM(questions_reviewed), 0) as total_questions,
        COALESCE(SUM(questions_correct), 0) as total_correct
        FROM study_sessions
        WHERE user_id = ? AND started_at >= ?''',
        (request.user_id, cutoff))
    agg = dict(c.fetchone())

    # Daily breakdown for the period
    c.execute('''SELECT
        DATE(started_at) as day,
        SUM(duration_seconds) as seconds,
        SUM(questions_reviewed) as questions
        FROM study_sessions
        WHERE user_id = ? AND started_at >= ?
        GROUP BY DATE(started_at)
        ORDER BY day''',
        (request.user_id, cutoff))
    daily = [dict(r) for r in c.fetchall()]

    conn.close()
    return jsonify({
        'stats': agg,
        'daily': daily
    })

# === Study Guide Builder ===

import html as html_module
from io import BytesIO

# Try to import document parsing libraries
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class StudyGuideBuilder:
    ICONS = ['🔐', '🧮', '📊', '#️⃣', '⚛️', '🔒', '📝', '💡', '🎯', '🔍', '📚', '🌐', '⚙️', '🔧', '📡']
    
    # Common words to ignore as key terms
    STOP_WORDS = {
        # Articles, prepositions, conjunctions
        'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with',
        'by', 'from', 'as', 'is', 'was', 'are', 'were', 'been', 'be', 'have', 'has', 'had',
        # Modal/auxiliary verbs
        'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must',
        'shall', 'can', 'need', 'dare', 'ought', 'used', 'being', 'having',
        # Pronouns
        'it', 'its', 'this', 'that', 'these', 'those', 'i', 'you', 'he', 'she', 'we', 'they',
        'what', 'which', 'who', 'whom', 'whose', 'where', 'when', 'why', 'how', 'me', 'him', 'her',
        # Quantifiers
        'all', 'each', 'every', 'both', 'few', 'more', 'most', 'other', 'some', 'such',
        'any', 'none', 'either', 'neither', 'enough', 'less', 'least', 'little',
        # Negation and emphasis
        'no', 'nor', 'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 'just',
        'also', 'now', 'here', 'there', 'then', 'once', 'still', 'already', 'ever',
        # Conjunctions and transitions  
        'if', 'unless', 'until', 'while', 'although', 'because', 'since', 'about',
        'into', 'through', 'during', 'before', 'after', 'above', 'below', 'between', 'under',
        'again', 'further', 'always', 'never', 'often', 'sometimes', 'usually',
        # Common generic words
        'example', 'examples', 'way', 'ways', 'much', 'many', 'several', 'eve', 'even',
        'first', 'second', 'third', 'one', 'two', 'three', 'new', 'old', 'good', 'best',
        'like', 'well', 'back', 'over', 'such', 'only', 'come', 'make', 'made', 'know',
        'take', 'get', 'got', 'see', 'saw', 'look', 'think', 'thought', 'want', 'give',
        'use', 'find', 'tell', 'ask', 'work', 'seem', 'feel', 'try', 'leave', 'call',
        'keep', 'let', 'begin', 'began', 'begun', 'show', 'hear', 'play', 'run', 'move',
        'live', 'believe', 'hold', 'bring', 'happen', 'write', 'provide', 'sit', 'stand',
        'lose', 'pay', 'meet', 'include', 'continue', 'set', 'learn', 'change', 'lead',
        'understand', 'watch', 'follow', 'stop', 'create', 'speak', 'read', 'allow',
        'add', 'spend', 'grow', 'open', 'walk', 'win', 'offer', 'remember', 'love',
        'consider', 'appear', 'buy', 'wait', 'serve', 'die', 'send', 'expect', 'build',
        'stay', 'fall', 'cut', 'reach', 'kill', 'remain', 'using', 'another', 'goes',
        # Common phrases that get picked up
        'however', 'therefore', 'thus', 'hence', 'otherwise', 'moreover', 'furthermore',
        'nonetheless', 'nevertheless', 'meanwhile', 'instead', 'indeed', 'certainly',
        'simply', 'actually', 'basically', 'generally', 'usually', 'typically', 'probably',
        'possibly', 'perhaps', 'maybe', 'likely', 'unlikely', 'clearly', 'obviously',
        # Filler/connector words
        'really', 'quite', 'rather', 'somewhat', 'almost', 'nearly', 'completely',
        'exactly', 'especially', 'particularly', 'specifically', 'mainly', 'mostly',
        'kind', 'sort', 'type', 'thing', 'things', 'stuff', 'part', 'parts', 'point',
        'fact', 'case', 'issue', 'problem', 'problems', 'question', 'answer', 'result',
        'able', 'sure', 'true', 'false', 'real', 'right', 'wrong', 'different', 'similar',
        # Generic action/technical words that aren't terms
        'divide', 'solve', 'grouping', 'regular', 'process', 'system', 'systems',
        'input', 'output', 'inputs', 'outputs', 'message', 'messages', 'data',
        'value', 'values', 'number', 'numbers', 'function', 'functions', 'method',
        'easy', 'hard', 'simple', 'complex', 'basic', 'advanced', 'standard',
        'alice', 'bob', 'eve',  # Cryptography placeholder names
        'homework', 'review', 'note', 'notes', 'section', 'chapter', 'appendix'
    }
    
    def __init__(self):
        self.key_terms = set()
        
    def _is_valid_term(self, term):
        """Check if a term is worth highlighting."""
        term = term.strip()
        
        # Too short or too long
        if len(term) < 5 or len(term) > 45:
            return False
        
        # Skip things that look like math/calculations
        if any(x in term for x in ['=', '+', '×', '÷', '/', '*', ' x ', '(mod', 'mod ']):
            return False
        
        # Skip things starting with numbers
        if term[0].isdigit():
            return False
        
        # Skip things with colons, quotes (embedded quotes are messy)
        if any(x in term for x in [':', '"', '"', '"', "'"]):
            return False
        
        # Skip things that look like hex/hashes or random alphanumeric
        if re.match(r'^[A-Z0-9]{6,}$', term):  # Like OEG20021111S0036
            return False
        if re.match(r'^[a-f0-9]{8,}$', term.lower()):
            return False
        
        # Skip URLs and URL fragments
        if any(x in term.lower() for x in ['http', 'www.', '.com', '.org', '.pdf', 'youtube', '.html']):
            return False
        
        # Single word checks
        words = term.lower().split()
        if len(words) == 1:
            # Single word must be at least 5 chars and not a stop word
            if term.lower() in self.STOP_WORDS:
                return False
            # Must contain letters (not just numbers/symbols)
            if not re.search(r'[a-zA-Z]{3,}', term):
                return False
            # Skip ALL CAPS words that are likely emphasis or section titles
            if term.isupper():
                # Allow known acronyms but skip generic caps words
                allowed_caps = {'SHA-256', 'SSL-TLS', 'SHA256', 'MD5', 'AES', 'RSA', 'PKI', 'PFS', 
                               'TLS', 'SSL', 'QEC', 'DHM', 'ECDH', 'HMAC', 'SHA1', 'SHA3', 'SAM'}
                if term not in allowed_caps and len(term) < 10:
                    return False
            return True
        
        # Multi-word phrase checks
        # Max 5 words (avoid sentence fragments)
        if len(words) > 5:
            return False
        
        # Filter out phrases that are mostly stop words
        non_stop_words = [w for w in words if w not in self.STOP_WORDS and len(w) > 2]
        if len(non_stop_words) < 2:  # Need at least 2 meaningful words
            return False
        
        # Skip phrases starting with bad words
        bad_starts = ['perform', 'divide', 'thus', 'note', 'determine', 'review', 'read', 
                      'make', 'take', 'when', 'what', 'where', 'which', 'this', 'that',
                      'there', 'these', 'those', 'here', 'they', 'your', 'first', 'most',
                      'above', 'below', 'essentially', 'input', 'output', 'easy', 'hard',
                      'can', 'cannot', 'should', 'would', 'could', 'will', 'shall']
        if words[0] in bad_starts:
            return False
        
        # Skip phrases ending with generic words
        bad_ends = ['way', 'ways', 'used', 'them', 'this', 'that', 'here', 'there',
                    'states', 'related', 'needed', 'required', 'done', 'help']
        if words[-1] in bad_ends:
            return False
            
        return True
    
    def _clean_term(self, term):
        """Clean up a term for consistency."""
        term = term.strip()
        # Remove trailing punctuation
        term = term.rstrip('.,;:!?')
        # Remove leading/trailing quotes
        term = term.strip('"\'""''')
        return term
    
    def _extract_key_terms_from_paragraph(self, para):
        """Extract key terms from bold/highlighted runs, combining adjacent runs."""
        terms = []
        current_term_parts = []
        
        for run in para.runs:
            is_key = run.bold or run.font.highlight_color
            text = run.text
            
            if is_key and text.strip():
                current_term_parts.append(text)
            else:
                # End of a key term sequence
                if current_term_parts:
                    full_term = ''.join(current_term_parts)
                    full_term = self._clean_term(full_term)
                    if self._is_valid_term(full_term):
                        terms.append(full_term)
                    current_term_parts = []
        
        # Don't forget the last term
        if current_term_parts:
            full_term = ''.join(current_term_parts)
            full_term = self._clean_term(full_term)
            if self._is_valid_term(full_term):
                terms.append(full_term)
        
        return terms
    
    def _extract_acronyms(self, text):
        """Extract acronyms like PKI, PFS, D-H, SHA-256, etc."""
        # Pattern for acronyms: 2+ uppercase letters, optionally with numbers/hyphens
        acronym_pattern = r'\b[A-Z][A-Z0-9][-A-Z0-9]*[A-Z0-9]\b'
        matches = re.findall(acronym_pattern, text)
        return [m for m in matches if len(m) >= 2 and m not in {'II', 'III', 'IV'}]
    
    def parse_docx(self, file_stream):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx required")
        
        doc = Document(file_stream)
        content = {'title': '', 'subtitle': '', 'sections': [], 'key_terms': []}
        current_section = None
        current_subsection = None
        all_text = []  # Collect all text for acronym extraction
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            all_text.append(text)
            style = para.style.name if para.style else 'Normal'
            
            # Extract key terms from bold/highlighted runs
            para_terms = self._extract_key_terms_from_paragraph(para)
            for term in para_terms:
                self.key_terms.add(term)
            
            if style == 'Title':
                content['title'] = text
            elif style == 'Subtitle':
                content['subtitle'] = text
            elif 'Heading 1' in style:
                current_section = {'title': text, 'level': 1, 'content': [], 'subsections': []}
                content['sections'].append(current_section)
                current_subsection = None
            elif 'Heading 2' in style and current_section:
                current_subsection = {'title': text, 'level': 2, 'content': []}
                current_section['subsections'].append(current_subsection)
            elif 'Heading 3' in style:
                target = current_subsection or current_section
                if target:
                    target['content'].append({'type': 'subheading', 'text': text})
            elif 'List' in style:
                target = current_subsection or current_section
                if target:
                    if target['content'] and target['content'][-1].get('type') == 'list':
                        target['content'][-1]['items'].append(text)  # Store raw, process later
                    else:
                        target['content'].append({'type': 'list', 'items': [text]})
            else:
                target = current_subsection or current_section
                if target:
                    target['content'].append({'type': 'paragraph', 'text': text})
        
        # Extract acronyms from all text (but be selective)
        full_text = ' '.join(all_text)
        acronyms = self._extract_acronyms(full_text)
        # Common acronyms/caps words to skip (not technical terms)
        skip_acronyms = {
            # General
            'DNA', 'RNA', 'USA', 'UK', 'EU', 'UN', 'TV', 'PC', 'IT', 'ID', 
            'OK', 'AM', 'PM', 'VS', 'IE', 'EG', 'AKA', 'FAQ', 'DIY', 'CEO',
            'NOT', 'AND', 'THE', 'FOR', 'BUT', 'ARE', 'WAS', 'HAS', 'HAD',
            # Document/section markers
            'HOMEWORK', 'NOTE', 'NOTES', 'TODO', 'TBD', 'NB', 'PS', 'FYI',
            # Generic technical words
            'SYSTEM', 'SOLVE', 'SALT', 'UNIX', 'SAM', 'NSA', 'CIA', 'FBI',
            # Action words that get caps'd
            'READ', 'WRITE', 'SEND', 'RECEIVE', 'GET', 'SET', 'PUT', 'POST',
            'TRUE', 'FALSE', 'NULL', 'VOID', 'INT', 'CHAR', 'BOOL'
        }
        for acr in acronyms:
            if len(acr) >= 3 and acr not in skip_acronyms:
                self.key_terms.add(acr)
        
        # Deduplicate case-insensitively (keep first occurrence's casing)
        seen_lower = {}
        for term in self.key_terms:
            lower = term.lower()
            if lower not in seen_lower:
                seen_lower[lower] = term
        self.key_terms = set(seen_lower.values())
        
        # Now process all text with final key terms
        self._process_all_content(content)
        
        content['key_terms'] = sorted(list(self.key_terms), key=lambda x: (-len(x), x.lower()))
        return content
    
    def _process_all_content(self, content):
        """Process all content to add key term highlighting."""
        for section in content.get('sections', []):
            new_content = []
            for item in section.get('content', []):
                if item['type'] == 'paragraph':
                    item['text'] = self._process_text(item['text'])
                elif item['type'] == 'list':
                    item['items'] = [self._process_text(i) for i in item['items']]
                new_content.append(item)
            section['content'] = new_content
            
            for sub in section.get('subsections', []):
                new_sub_content = []
                for item in sub.get('content', []):
                    if item['type'] == 'paragraph':
                        item['text'] = self._process_text(item['text'])
                    elif item['type'] == 'list':
                        item['items'] = [self._process_text(i) for i in item['items']]
                    new_sub_content.append(item)
                sub['content'] = new_sub_content
    
    def parse_pdf(self, file_stream):
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 required")
        
        reader = PyPDF2.PdfReader(file_stream)
        content = {'title': 'Imported PDF', 'sections': [], 'key_terms': []}
        
        full_text = '\n'.join(page.extract_text() or '' for page in reader.pages)
        
        # Extract acronyms
        acronyms = self._extract_acronyms(full_text)
        for acr in acronyms:
            self.key_terms.add(acr)
        
        lines = full_text.split('\n')
        current_section = None
        current_content = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if self._looks_like_heading(line):
                if current_section:
                    current_section['content'] = current_content
                    content['sections'].append(current_section)
                current_section = {'title': line.title() if line.isupper() else line, 'level': 1, 'content': [], 'subsections': []}
                current_content = []
            elif current_section:
                if line.startswith(('•', '-', '*', '·')):
                    text = line[1:].strip()
                    if current_content and current_content[-1].get('type') == 'list':
                        current_content[-1]['items'].append(self._process_text(text))
                    else:
                        current_content.append({'type': 'list', 'items': [self._process_text(text)]})
                else:
                    current_content.append({'type': 'paragraph', 'text': self._process_text(line)})
        
        if current_section:
            current_section['content'] = current_content
            content['sections'].append(current_section)
        
        if not content['sections'] and full_text:
            content['sections'].append({'title': 'Document Content', 'level': 1, 'content': [{'type': 'paragraph', 'text': self._process_text(full_text[:5000])}], 'subsections': []})
        
        if content['sections']:
            content['title'] = content['sections'][0]['title']
        
        content['key_terms'] = sorted(list(self.key_terms), key=lambda x: (-len(x), x.lower()))
        return content
    
    def _process_text(self, text):
        """Add key term highlighting to text."""
        text = html_module.escape(text)
        
        # Track which positions are already highlighted to avoid overlaps
        highlighted = set()
        
        # Sort terms by length (longest first) to avoid partial matches
        sorted_terms = sorted(self.key_terms, key=len, reverse=True)
        
        for term in sorted_terms:
            escaped_term = html_module.escape(term)
            # Use word boundaries to avoid partial matches
            pattern = re.compile(r'\b' + re.escape(escaped_term) + r'\b', re.IGNORECASE)
            
            # Find all matches
            for match in pattern.finditer(text):
                start, end = match.start(), match.end()
                
                # Check if this region is already highlighted
                if any(start < h_end and end > h_start for h_start, h_end in highlighted):
                    continue
                
                # Only highlight first occurrence
                original = match.group()
                replacement = f'<span class="key-term">{original}</span>'
                text = text[:start] + replacement + text[end:]
                
                # Update highlighted positions (accounting for added HTML)
                added_len = len(replacement) - len(original)
                highlighted = {(s + added_len if s >= end else s, e + added_len if e >= end else e) for s, e in highlighted}
                highlighted.add((start, start + len(replacement)))
                
                break  # Only highlight first occurrence per term
        
        # Linkify URLs
        url_pattern = r'(https?://[^\s<>"\']+)'
        text = re.sub(url_pattern, r'<a href="\1" target="_blank" rel="noopener" class="link">\1</a>', text)
        
        return text
    
    def _looks_like_heading(self, line):
        if not line:
            return False
        if line.isupper() and 3 < len(line) < 60:
            return True
        if re.match(r'^(\d+\.)+\s+[A-Z]', line):
            return True
        return len(line) < 50 and not line.endswith('.') and line[0].isupper()
    
    def _slugify(self, text):
        text = re.sub(r'[^\w\s-]', '', text.lower())
        return re.sub(r'[\s_]+', '-', text)[:50]
    
    def _truncate_title(self, title, max_len=35):
        """Truncate title at word boundary."""
        if len(title) <= max_len:
            return title
        # Find last space before max_len
        truncated = title[:max_len]
        last_space = truncated.rfind(' ')
        if last_space > max_len * 0.5:  # Only truncate at space if reasonable
            return truncated[:last_space] + '…'
        return truncated + '…'
    
    def generate_html(self, content):
        title = content.get('title', 'Study Guide')
        subtitle = content.get('subtitle', '')
        sections = content.get('sections', [])
        key_terms = content.get('key_terms', [])
        
        total_items = sum(len(s.get('content', [])) + sum(len(sub.get('content', [])) for sub in s.get('subsections', [])) for s in sections)
        read_time = max(5, total_items // 3)
        
        nav = '\n'.join(f'<a href="#{self._slugify(s["title"])}" class="nav-link">{self.ICONS[i % len(self.ICONS)]} {html_module.escape(self._truncate_title(s["title"]))}</a>' for i, s in enumerate(sections))
        nav += '\n<a href="#terms" class="nav-link">📋 Key Terms</a>'
        
        sections_html = self._render_sections(sections)
        terms_html = self._render_terms(key_terms)
        
        return f'''<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{html_module.escape(title)} | Study Guide</title>
<link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
<style>{self._get_css()}</style>
</head>
<body>
<header class="hero"><div class="hero-content"><div class="course-tag">📚 Study Guide</div><h1>{html_module.escape(title)}</h1>{f'<p class="hero-subtitle">{html_module.escape(subtitle)}</p>' if subtitle else ''}<div class="hero-stats"><div class="stat"><div class="stat-value">{len(sections)}</div><div class="stat-label">Sections</div></div><div class="stat"><div class="stat-value">{len(key_terms)}</div><div class="stat-label">Key Terms</div></div><div class="stat"><div class="stat-value">~{read_time}</div><div class="stat-label">Min Read</div></div></div></div></header>
<nav class="nav-container"><div class="nav-scroll"><div class="nav-links">{nav}</div></div></nav>
<main class="main-content">{sections_html}{terms_html}</main>
<div class="progress-tracker"><svg class="progress-ring" viewBox="0 0 44 44"><circle class="bg" cx="22" cy="22" r="18"/><circle class="progress" cx="22" cy="22" r="18" stroke-dasharray="113" stroke-dashoffset="113"/></svg><span class="progress-percent">0%</span></div>
<button class="print-btn" onclick="window.print()" title="Print">🖨️</button>
<script>{self._get_js()}</script>
</body>
</html>'''

    def _render_sections(self, sections):
        parts = []
        for i, s in enumerate(sections):
            sid = self._slugify(s['title'])
            icon = self.ICONS[i % len(self.ICONS)]
            content = self._render_content(s.get('content', []))
            subs = ''.join(f'<div id="{self._slugify(sub["title"])}" class="subsection"><h3 class="subsection-title">{html_module.escape(sub["title"])}</h3>{self._render_content(sub.get("content", []))}</div>' for sub in s.get('subsections', []))
            parts.append(f'<section id="{sid}" class="section"><div class="section-header"><div class="section-icon">{icon}</div><div class="section-meta"><div class="section-number">Section {i+1:02d}</div><h2 class="section-title">{html_module.escape(s["title"])}</h2></div></div>{content}{subs}</section>')
        return '\n'.join(parts)
    
    def _render_content(self, content):
        parts = []
        for b in content:
            t = b.get('type', 'paragraph')
            if t == 'paragraph':
                parts.append(f'<p>{b["text"]}</p>')
            elif t == 'subheading':
                parts.append(f'<h4 class="content-subheading">{html_module.escape(b["text"])}</h4>')
            elif t == 'list':
                items = ''.join(f'<li>{item}</li>' for item in b['items'])
                parts.append(f'<ul class="bullet-list">{items}</ul>')
            elif t == 'formula':
                parts.append(f'<div class="formula"><code>{html_module.escape(b["text"])}</code></div>')
        return '\n'.join(parts)
    
    def _render_terms(self, terms):
        if not terms:
            return ''
        chips = ''.join(f'<div class="term-chip">{html_module.escape(t)}</div>' for t in sorted(set(terms))[:40])
        return f'<section id="terms" class="quick-ref"><h2>📋 Key Terms Reference</h2><div class="terms-cloud">{chips}</div></section>'
    
    def _get_css(self):
        return ''':root{--bg-dark:#0f0f14;--bg-card:#16161e;--bg-elevated:#1e1e28;--text-primary:#e4e4e7;--text-secondary:#a1a1aa;--text-muted:#71717a;--primary:#a78bfa;--primary-soft:#c4b5fd;--accent:#67e8f9;--accent-soft:#a5f3fc;--success:#34d399;--border:rgba(255,255,255,0.06);--glow:rgba(167,139,250,0.15)}*{margin:0;padding:0;box-sizing:border-box}html{scroll-behavior:smooth}body{font-family:'Inter',-apple-system,BlinkMacSystemFont,sans-serif;background:var(--bg-dark);color:var(--text-primary);line-height:1.8;font-size:16px;-webkit-font-smoothing:antialiased}.hero{position:relative;padding:4rem 2rem 3rem;background:linear-gradient(180deg,rgba(167,139,250,0.08) 0%,transparent 100%);border-bottom:1px solid var(--border)}.hero-content{position:relative;max-width:800px;margin:0 auto}.course-tag{display:inline-block;padding:0.4rem 1rem;background:rgba(167,139,250,0.15);border:1px solid rgba(167,139,250,0.2);border-radius:2rem;font-size:0.75rem;font-weight:600;color:var(--primary);text-transform:uppercase;letter-spacing:0.05em;margin-bottom:1.25rem}.hero h1{font-size:2.25rem;font-weight:700;line-height:1.2;margin-bottom:1rem;color:#fff}.hero-subtitle{font-size:1.05rem;color:var(--text-secondary);margin-bottom:2rem;line-height:1.6}.hero-stats{display:flex;gap:2.5rem}.stat{text-align:left}.stat-value{font-size:1.75rem;font-weight:700;color:#fff;margin-bottom:0.125rem}.stat-label{font-size:0.7rem;color:var(--text-muted);text-transform:uppercase;letter-spacing:0.05em}.nav-container{position:sticky;top:0;z-index:100;background:rgba(15,15,20,0.92);backdrop-filter:blur(12px);border-bottom:1px solid var(--border)}.nav-scroll{max-width:1000px;margin:0 auto;padding:0 1.5rem;overflow-x:auto;scrollbar-width:none}.nav-scroll::-webkit-scrollbar{display:none}.nav-links{display:flex;gap:0.25rem;padding:0.625rem 0}.nav-link{padding:0.5rem 0.875rem;color:var(--text-muted);text-decoration:none;font-size:0.8rem;font-weight:500;border-radius:6px;transition:all 0.2s;white-space:nowrap}.nav-link:hover{color:var(--text-primary);background:var(--bg-elevated)}.nav-link.active{color:var(--primary);background:rgba(167,139,250,0.1)}.main-content{max-width:760px;margin:0 auto;padding:3rem 1.5rem 5rem}.section{margin-bottom:4rem;scroll-margin-top:70px}.section-header{display:flex;align-items:flex-start;gap:1rem;margin-bottom:1.75rem}.section-icon{width:48px;height:48px;display:flex;align-items:center;justify-content:center;background:linear-gradient(135deg,rgba(167,139,250,0.2),rgba(103,232,249,0.15));border:1px solid rgba(167,139,250,0.15);border-radius:12px;font-size:1.375rem;flex-shrink:0}.section-meta{flex:1;padding-top:0.25rem}.section-number{font-size:0.65rem;font-weight:600;text-transform:uppercase;letter-spacing:0.1em;color:var(--primary);margin-bottom:0.375rem;opacity:0.8}.section-title{font-size:1.5rem;font-weight:600;line-height:1.3;color:#fff}.subsection{margin:2.5rem 0;padding-left:1.25rem;border-left:2px solid var(--border)}.subsection-title{color:var(--accent);font-size:1.05rem;font-weight:600;margin-bottom:1rem}.content-subheading{color:var(--primary-soft);font-size:1rem;font-weight:600;margin:2rem 0 0.75rem}p{margin:1.25rem 0;color:var(--text-secondary);font-size:0.9375rem}.key-term{color:var(--text-primary);font-weight:500;border-bottom:1px dotted rgba(167,139,250,0.4)}.link{color:var(--accent);text-decoration:none;word-break:break-all}.link:hover{text-decoration:underline}.bullet-list{list-style:none;margin:1.5rem 0;padding:0}.bullet-list li{position:relative;padding:0.75rem 0 0.75rem 1.5rem;color:var(--text-secondary);font-size:0.9375rem;border-bottom:1px solid var(--border)}.bullet-list li:last-child{border-bottom:none}.bullet-list li::before{content:'';position:absolute;left:0;top:1.1rem;width:6px;height:6px;background:var(--primary);border-radius:50%;opacity:0.6}.formula{font-family:'JetBrains Mono',monospace;font-size:0.875rem;background:var(--bg-elevated);border:1px solid var(--border);border-radius:8px;padding:1rem 1.25rem;margin:1.25rem 0;overflow-x:auto;color:var(--accent-soft)}.quick-ref{margin-top:4rem;padding-top:2.5rem;border-top:1px solid var(--border)}.quick-ref h2{font-size:1.25rem;font-weight:600;margin-bottom:1.25rem;color:#fff}.terms-cloud{display:flex;flex-wrap:wrap;gap:0.5rem}.term-chip{padding:0.5rem 1rem;background:var(--bg-card);border:1px solid var(--border);border-radius:2rem;font-size:0.8rem;color:var(--text-secondary);transition:all 0.2s}.term-chip:hover{border-color:rgba(167,139,250,0.3);color:var(--primary-soft)}.progress-tracker{position:fixed;bottom:1.5rem;right:1.5rem;background:var(--bg-card);border:1px solid var(--border);border-radius:12px;padding:0.75rem 1rem;display:flex;align-items:center;gap:0.75rem;box-shadow:0 8px 32px rgba(0,0,0,0.4);z-index:100}.progress-ring{width:40px;height:40px;transform:rotate(-90deg)}.progress-ring circle{fill:none;stroke-width:3}.progress-ring .bg{stroke:var(--border)}.progress-ring .progress{stroke:var(--primary);stroke-linecap:round;transition:stroke-dashoffset 0.5s}.progress-percent{font-size:1rem;font-weight:600;color:var(--primary)}.print-btn{position:fixed;bottom:1.5rem;left:1.5rem;width:42px;height:42px;background:var(--bg-card);border:1px solid var(--border);border-radius:50%;display:flex;align-items:center;justify-content:center;font-size:1.125rem;cursor:pointer;transition:all 0.2s;z-index:100}.print-btn:hover{background:var(--bg-elevated);border-color:var(--primary)}@media(max-width:768px){.hero{padding:2rem 1.25rem}.hero h1{font-size:1.75rem}.section-header{flex-direction:column;gap:0.625rem}.section-icon{width:42px;height:42px;font-size:1.25rem}.section-title{font-size:1.25rem}.progress-tracker{bottom:1rem;right:1rem;padding:0.625rem 0.875rem}.print-btn{bottom:1rem;left:1rem}}@media print{.nav-container,.progress-tracker,.print-btn{display:none!important}body{background:#fff;color:#222;font-size:11pt}.hero{background:none;padding:1rem 0}.hero h1{color:#000;font-size:18pt}.section{page-break-inside:avoid}.key-term{color:#000;font-weight:600;border-bottom:none}.link{color:#333;text-decoration:underline}}'''
    
    def _get_js(self):
        return '''const sections=document.querySelectorAll('.section'),navLinks=document.querySelectorAll('.nav-link'),progressCircle=document.querySelector('.progress-ring .progress'),progressText=document.querySelector('.progress-percent'),circumference=113;function update(){const scrollHeight=document.documentElement.scrollHeight-window.innerHeight,progress=scrollHeight>0?(window.scrollY/scrollHeight)*100:0;progressCircle.style.strokeDashoffset=circumference-(progress/100)*circumference;progressText.textContent=Math.round(progress)+'%';let current='';sections.forEach(s=>{if(window.scrollY>=s.offsetTop-100)current=s.id});navLinks.forEach(l=>{l.classList.remove('active');if(l.getAttribute('href')==='#'+current)l.classList.add('active')})}window.addEventListener('scroll',update);update();navLinks.forEach(l=>l.addEventListener('click',e=>{e.preventDefault();document.querySelector(l.getAttribute('href'))?.scrollIntoView({behavior:'smooth'})}))'''


from werkzeug.utils import secure_filename

ALLOWED_EXTENSIONS = {'docx', 'pdf'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/api/study-guide/upload', methods=['POST'])
@token_required
def upload_study_guide():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if not file.filename or not allowed_file(file.filename):
        return jsonify({'error': 'File type not supported. Use DOCX or PDF.'}), 400
    
    try:
        builder = StudyGuideBuilder()
        ext = secure_filename(file.filename).rsplit('.', 1)[1].lower()
        
        file_stream = BytesIO(file.read())
        content = builder.parse_docx(file_stream) if ext == 'docx' else builder.parse_pdf(file_stream)
        html_content = builder.generate_html(content)
        
        return jsonify({
            'success': True,
            'title': content.get('title', 'Study Guide'),
            'sections': len(content.get('sections', [])),
            'key_terms': len(content.get('key_terms', [])),
            'html': html_content
        })
    except ImportError as e:
        return jsonify({'error': str(e) + '. Install with: pip install python-docx PyPDF2'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/study-guide/preview', methods=['POST'])
@token_required
def preview_study_guide():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if not file.filename or not allowed_file(file.filename):
        return jsonify({'error': 'Unsupported file type'}), 400
    
    try:
        builder = StudyGuideBuilder()
        ext = secure_filename(file.filename).rsplit('.', 1)[1].lower()
        
        file_stream = BytesIO(file.read())
        content = builder.parse_docx(file_stream) if ext == 'docx' else builder.parse_pdf(file_stream)
        
        return jsonify({
            'success': True,
            'title': content.get('title', 'Study Guide'),
            'sections': [{'title': s['title'], 'items': len(s.get('content', []))} for s in content.get('sections', [])],
            'key_terms': content.get('key_terms', [])[:20]
        })
    except ImportError as e:
        return jsonify({'error': str(e) + '. Install with: pip install python-docx PyPDF2'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# === Spaced Repetition System (SRS) Routes ===

@app.route('/api/srs/due', methods=['GET'])
@token_required
def get_due_cards():
    """Get SRS cards due for review."""
    limit = request.args.get('limit', 20, type=int)
    conn = get_db()
    c = conn.cursor()
    c.execute('''SELECT sc.*, q.question_text, q.type, q.options, q.correct, q.pairs,
        q.code, q.code_language, q.image, q.image_alt, q.explanation, q.quiz_id
        FROM srs_cards sc
        JOIN questions q ON sc.question_id = q.id
        WHERE sc.user_id = ? AND sc.next_review_at <= datetime('now')
        AND sc.status != 'graduated'
        ORDER BY sc.next_review_at ASC LIMIT ?''',
        (request.user_id, limit))
    rows = c.fetchall()
    cards = []
    for row in rows:
        card = dict(row)
        # Parse JSON fields
        if card.get('options'):
            card['options'] = json.loads(card['options'])
        if card.get('correct'):
            card['correct'] = json.loads(card['correct'])
        if card.get('pairs'):
            card['pairs'] = json.loads(card['pairs'])
        cards.append(card)
    conn.close()
    return jsonify({'cards': cards})

@app.route('/api/srs/cards', methods=['POST'])
@token_required
def add_srs_cards():
    """Add questions to the SRS deck."""
    data = request.get_json() or {}
    question_ids = data.get('questionIds', [])
    if not question_ids:
        return jsonify({'error': 'No question IDs provided'}), 400

    conn = get_db()
    c = conn.cursor()
    added = 0
    for qid in question_ids:
        try:
            c.execute('''INSERT OR IGNORE INTO srs_cards (user_id, question_id)
                VALUES (?, ?)''', (request.user_id, qid))
            if c.rowcount > 0:
                added += 1
        except sqlite3.IntegrityError:
            pass  # Question doesn't exist or already added
    conn.commit()
    conn.close()
    return jsonify({'added': added, 'total_requested': len(question_ids)})

@app.route('/api/srs/review', methods=['POST'])
@token_required
def submit_srs_review():
    """Submit a review result using the SM-2 algorithm."""
    data = request.get_json() or {}
    card_id = data.get('cardId')
    quality = data.get('quality')

    if card_id is None or quality is None:
        return jsonify({'error': 'cardId and quality are required'}), 400
    if not isinstance(quality, int) or quality < 0 or quality > 5:
        return jsonify({'error': 'quality must be an integer 0-5'}), 400

    conn = get_db()
    c = conn.cursor()

    # Fetch the card
    c.execute('SELECT * FROM srs_cards WHERE id = ? AND user_id = ?', (card_id, request.user_id))
    card = c.fetchone()
    if not card:
        conn.close()
        return jsonify({'error': 'Card not found'}), 404

    repetitions = card['repetitions']
    interval = card['interval_days']
    ease_factor = card['ease_factor']

    # SM-2 algorithm
    if quality >= 3:  # correct response
        if repetitions == 0:
            interval = 1
        elif repetitions == 1:
            interval = 6
        else:
            interval = round(interval * ease_factor)
        repetitions += 1
    else:  # incorrect response
        repetitions = 0
        interval = 1

    # Update ease factor
    ease_factor = max(1.3, ease_factor + (0.1 - (5 - quality) * (0.08 + (5 - quality) * 0.02)))

    # Determine status
    now = datetime.now()
    next_review = now + timedelta(days=interval)
    status = 'review' if repetitions > 0 else 'learning'
    if interval >= 21:
        status = 'graduated'

    c.execute('''UPDATE srs_cards SET
        ease_factor = ?, interval_days = ?, repetitions = ?,
        next_review_at = ?, last_reviewed_at = ?, status = ?, updated_at = ?
        WHERE id = ? AND user_id = ?''',
        (ease_factor, interval, repetitions, next_review, now, status, now,
         card_id, request.user_id))

    # Also update question_performance
    is_correct = quality >= 3
    c.execute('''INSERT INTO question_performance
        (user_id, question_id, times_seen, times_correct, times_incorrect, last_seen_at, last_correct_at)
        VALUES (?, ?, 1, ?, ?, ?, ?)
        ON CONFLICT(user_id, question_id) DO UPDATE SET
        times_seen = times_seen + 1,
        times_correct = times_correct + ?,
        times_incorrect = times_incorrect + ?,
        last_seen_at = ?,
        last_correct_at = CASE WHEN ? = 1 THEN ? ELSE last_correct_at END,
        updated_at = ?''',
        (request.user_id, card['question_id'],
         1 if is_correct else 0,
         0 if is_correct else 1,
         now,
         now if is_correct else None,
         # ON CONFLICT params
         1 if is_correct else 0,
         0 if is_correct else 1,
         now,
         1 if is_correct else 0, now,
         now))

    conn.commit()
    conn.close()

    return jsonify({
        'card_id': card_id,
        'new_interval': interval,
        'new_ease_factor': round(ease_factor, 2),
        'new_status': status,
        'next_review_at': next_review.isoformat()
    })

@app.route('/api/srs/stats', methods=['GET'])
@token_required
def get_srs_stats():
    """Get SRS statistics for the current user."""
    conn = get_db()
    c = conn.cursor()

    # Total cards
    c.execute('SELECT COUNT(*) as cnt FROM srs_cards WHERE user_id = ?', (request.user_id,))
    total_cards = c.fetchone()['cnt']

    # Due today
    c.execute('''SELECT COUNT(*) as cnt FROM srs_cards
        WHERE user_id = ? AND next_review_at <= datetime('now') AND status != 'graduated' ''',
        (request.user_id,))
    due_today = c.fetchone()['cnt']

    # Status counts
    c.execute('''SELECT status, COUNT(*) as cnt FROM srs_cards
        WHERE user_id = ? GROUP BY status''', (request.user_id,))
    status_counts = {row['status']: row['cnt'] for row in c.fetchall()}

    # Streak: consecutive days with at least one review
    c.execute('''SELECT DISTINCT DATE(last_reviewed_at) as review_date
        FROM srs_cards WHERE user_id = ? AND last_reviewed_at IS NOT NULL
        ORDER BY review_date DESC''', (request.user_id,))
    review_dates = [row['review_date'] for row in c.fetchall()]
    streak = 0
    if review_dates:
        from datetime import date
        today = date.today()
        expected = today
        for rd in review_dates:
            if rd is None:
                break
            review_d = datetime.strptime(rd, '%Y-%m-%d').date() if isinstance(rd, str) else rd
            if review_d == expected:
                streak += 1
                expected = expected - timedelta(days=1)
            elif review_d < expected:
                break

    conn.close()
    return jsonify({
        'total_cards': total_cards,
        'due_today': due_today,
        'new_cards': status_counts.get('new', 0),
        'learning': status_counts.get('learning', 0),
        'review': status_counts.get('review', 0),
        'graduated': status_counts.get('graduated', 0),
        'streak': streak
    })

# ==================== Session Plan (Immersive Experience) ====================

@app.route('/api/session/plan', methods=['GET'])
@token_required
def get_session_plan():
    """Compute a personalized study session plan based on user's certs, weak areas, SRS, and exam date."""
    conn = get_db()
    c = conn.cursor()

    blocks = []

    # 1. Get user certifications with target dates
    c.execute('''SELECT uc.certification_id, uc.target_date, uc.status,
        cert.code, cert.name, cert.passing_score
        FROM user_certifications uc
        JOIN certifications cert ON uc.certification_id = cert.id
        WHERE uc.user_id = ?
        ORDER BY uc.started_at DESC''', (request.user_id,))
    user_certs = [dict(r) for r in c.fetchall()]

    primary_cert = user_certs[0] if user_certs else None

    # Calculate days remaining
    days_remaining = None
    if primary_cert and primary_cert.get('target_date'):
        try:
            target = datetime.strptime(primary_cert['target_date'], '%Y-%m-%d').date()
            days_remaining = (target - datetime.now().date()).days
            if days_remaining < 0:
                days_remaining = 0
        except (ValueError, TypeError):
            pass

    # 2. SRS cards due
    c.execute('''SELECT COUNT(*) as cnt FROM srs_cards
        WHERE user_id = ? AND next_review_at <= datetime('now') AND status != 'graduated' ''',
        (request.user_id,))
    srs_due = c.fetchone()['cnt']

    c.execute('''SELECT COUNT(*) as cnt FROM srs_cards WHERE user_id = ?''', (request.user_id,))
    srs_total = c.fetchone()['cnt']

    c.execute('''SELECT COUNT(*) as cnt FROM srs_cards
        WHERE user_id = ? AND status = 'graduated' ''', (request.user_id,))
    srs_graduated = c.fetchone()['cnt']

    if srs_due > 0:
        est_minutes = max(2, round(srs_due * 0.5))
        blocks.append({
            'type': 'srs_review',
            'priority': 1,
            'title': f'{srs_due} cards due for review',
            'subtitle': f'{srs_total} total \u00b7 {srs_graduated} mastered',
            'estimate_minutes': est_minutes,
            'action': 'startSrsReview',
            'count': srs_due,
        })

    # 3. Weak domain blocks (if user has a primary cert)
    # NOTE: In-progress personal quizzes intentionally excluded — quizzes and certifications are independent systems
    weak_domains = []
    if primary_cert:
        cert_id = primary_cert['certification_id']
        c.execute('''SELECT d.id, d.name, d.code, d.weight,
            COALESCE(SUM(qp.times_correct), 0) as correct,
            COALESCE(SUM(qp.times_seen), 0) as seen
        FROM domains d
        LEFT JOIN question_domains qd ON qd.domain_id = d.id
        LEFT JOIN question_performance qp ON qp.question_id = qd.question_id AND qp.user_id = ?
        WHERE d.certification_id = ? AND d.parent_domain_id IS NULL
        GROUP BY d.id
        ORDER BY CASE
            WHEN COALESCE(SUM(qp.times_seen), 0) = 0 THEN 1.0
            ELSE 1.0 - (CAST(COALESCE(SUM(qp.times_correct), 0) AS REAL) / MAX(COALESCE(SUM(qp.times_seen), 0), 1))
        END DESC''', (request.user_id, cert_id))

        for row in c.fetchall():
            d = dict(row)
            seen = d['seen']
            correct = d['correct']
            score = round(correct / seen * 100) if seen > 0 else 0
            if seen == 0:
                status = 'unseen'
            elif score >= 80:
                status = 'strong'
            elif score >= 60:
                status = 'moderate'
            else:
                status = 'weak'
            weak_domains.append({
                'domain_id': d['id'],
                'name': d['name'],
                'code': d['code'],
                'score': score,
                'seen': seen,
                'status': status,
                'weight': d['weight'] or 0,
            })

        # Build targeted quiz blocks for the 2 weakest non-strong domains
        block_priority = 10
        for domain in weak_domains[:3]:
            if domain['status'] == 'strong':
                continue
            # Count available questions for this domain
            c.execute('''SELECT COUNT(*) as cnt FROM questions q
                JOIN question_domains qd ON q.id = qd.question_id
                WHERE qd.domain_id = ? AND q.is_active = 1''', (domain['domain_id'],))
            available = c.fetchone()['cnt']
            if available == 0:
                continue
            question_count = min(available, 15)
            est_minutes = max(5, round(question_count * 1.2))

            rationale = ''
            if domain['status'] == 'unseen':
                rationale = 'You haven\'t covered this domain yet'
            elif domain['status'] == 'weak':
                rationale = f'{domain["score"]}% accuracy \u2014 needs focused practice'
            else:
                rationale = f'{domain["score"]}% accuracy \u2014 room to improve'

            blocks.append({
                'type': 'domain_quiz',
                'priority': block_priority,
                'title': domain['name'],
                'subtitle': rationale,
                'domain_id': domain['domain_id'],
                'question_count': question_count,
                'estimate_minutes': est_minutes,
                'action': 'startDomainQuiz',
                'action_data': {
                    'certId': cert_id,
                    'domainId': domain['domain_id'],
                    'count': question_count,
                },
                'domain_score': domain['score'],
                'domain_status': domain['status'],
            })
            block_priority += 1

    # 5. Check if user is ready for a simulation
    if primary_cert:
        cert_id = primary_cert['certification_id']
        covered = sum(1 for d in weak_domains if d['seen'] >= 5)
        total_domains = len(weak_domains) if weak_domains else 1
        coverage = covered / total_domains

        # Check when user last completed a simulation (or if ever)
        c.execute('''SELECT MAX(created_at) as last_sim FROM exam_simulations
            WHERE user_id = ? AND certification_id = ?''', (request.user_id, cert_id))
        last_sim_row = c.fetchone()
        last_sim = last_sim_row['last_sim'] if last_sim_row else None
        days_since_sim = None
        if last_sim:
            try:
                last_sim_dt = datetime.strptime(last_sim[:19], '%Y-%m-%d %H:%M:%S')
                days_since_sim = (datetime.now() - last_sim_dt).days
            except (ValueError, TypeError):
                pass

        # Always offer diagnostic if user has never taken a simulation,
        # otherwise offer when they've covered 60%+ of domains and it's been 3+ days
        offer_sim = False
        if not last_sim:
            offer_sim = True
        elif coverage >= 0.6 and days_since_sim is not None and days_since_sim >= 3:
            offer_sim = True

        if offer_sim:
            overall = sum(d['score'] * d['weight'] for d in weak_domains) / max(sum(d['weight'] for d in weak_domains), 1) if weak_domains else 0
            if not last_sim:
                sim_subtitle = 'Take your first diagnostic to identify weak areas and build your study plan.'
            else:
                sim_subtitle = f'You\'ve covered {round(coverage * 100)}% of domains. Time to test under exam conditions.'
            blocks.append({
                'type': 'simulation_prompt',
                'priority': 50,
                'title': f'Practice Exam: {primary_cert["name"]}',
                'subtitle': sim_subtitle,
                'certification_id': cert_id,
                'action': 'startSimulation',
                'action_data': {'certId': cert_id},
                'readiness_pct': round(overall),
            })

    # 6. Overall readiness for header context
    overall_readiness = 0
    if weak_domains:
        total_weight = sum(d['weight'] for d in weak_domains)
        if total_weight > 0:
            overall_readiness = round(sum(d['score'] * d['weight'] for d in weak_domains) / total_weight)

    # 7. Recent study stats
    c.execute('''SELECT COALESCE(SUM(duration_seconds), 0) as total_secs
        FROM study_sessions WHERE user_id = ? AND started_at >= datetime('now', '-30 days')''',
        (request.user_id,))
    study_seconds = c.fetchone()['total_secs']
    study_hours = round(study_seconds / 3600, 1) if study_seconds else 0

    # Sort blocks by priority
    blocks.sort(key=lambda b: b['priority'])

    conn.close()

    return jsonify({
        'session': {
            'blocks': blocks,
            'context': {
                'certification': {
                    'id': primary_cert['certification_id'],
                    'code': primary_cert['code'],
                    'name': primary_cert['name'],
                } if primary_cert else None,
                'days_remaining': days_remaining,
                'overall_readiness': overall_readiness,
                'study_hours_30d': study_hours,
                'domains': weak_domains[:6] if weak_domains else [],
            },
        }
    })


@app.route('/api/session/domain-quiz', methods=['POST'])
@token_required
def start_domain_quiz():
    """Generate a targeted quiz for a specific domain."""
    data = request.get_json() or {}
    domain_id = data.get('domain_id')
    count = data.get('count', 15)

    if not domain_id:
        return jsonify({'error': 'domain_id required'}), 400

    conn = get_db()
    c = conn.cursor()

    # Get domain info
    c.execute('SELECT * FROM domains WHERE id = ?', (domain_id,))
    domain = c.fetchone()
    if not domain:
        conn.close()
        return jsonify({'error': 'Domain not found'}), 404
    domain = dict(domain)

    # Get questions for this domain
    c.execute('''SELECT q.* FROM questions q
        JOIN question_domains qd ON q.id = qd.question_id
        WHERE qd.domain_id = ? AND q.is_active = 1''', (domain_id,))
    all_questions = [dict(r) for r in c.fetchall()]

    if not all_questions:
        conn.close()
        return jsonify({'error': 'No questions available for this domain'}), 404

    import random
    selected = random.sample(all_questions, min(count, len(all_questions)))

    # Format questions for the quiz component
    questions = []
    for q in selected:
        formatted = {
            'id': q['id'],
            'question': q['question_text'],
            'type': q['type'] or 'choice',
            'explanation': q.get('explanation', ''),
        }
        if q.get('options'):
            formatted['options'] = json.loads(q['options'])
        if q.get('correct') is not None:
            try:
                formatted['correct'] = json.loads(q['correct']) if isinstance(q['correct'], str) else q['correct']
            except (json.JSONDecodeError, TypeError):
                formatted['correct'] = q['correct']
        if q.get('code'):
            formatted['code'] = q['code']
        if q.get('pairs'):
            try:
                formatted['pairs'] = json.loads(q['pairs'])
            except (json.JSONDecodeError, TypeError):
                pass
        if q.get('option_explanations'):
            try:
                formatted['optionExplanations'] = json.loads(q['option_explanations'])
            except (json.JSONDecodeError, TypeError):
                pass
        questions.append(formatted)

    conn.close()

    return jsonify({
        'quiz': {
            'id': f'domain-{domain_id}-{int(datetime.now().timestamp())}',
            'title': domain['name'],
            'questions': questions,
            'is_domain_quiz': True,
            'domain_id': domain_id,
            'certification_id': domain.get('certification_id'),
        }
    })


# ==================== Bookmarks ====================

@app.route('/api/bookmarks', methods=['GET'])
@token_required
def get_bookmarks():
    """Get all bookmarked questions for the current user."""
    conn = get_db()
    c = conn.cursor()
    c.execute('''SELECT b.id, b.question_id, b.note, b.created_at,
        q.question_text, q.type, q.options, q.correct, q.explanation,
        qz.id as quiz_id, qz.title as quiz_title
        FROM bookmarks b
        JOIN questions q ON b.question_id = q.id
        JOIN quizzes qz ON q.quiz_id = qz.id
        WHERE b.user_id = ?
        ORDER BY b.created_at DESC''', (request.user_id,))
    rows = c.fetchall()
    bookmarks = []
    for row in rows:
        bm = dict(row)
        if bm.get('options'):
            try:
                bm['options'] = json.loads(bm['options'])
            except Exception:
                pass
        bookmarks.append(bm)
    conn.close()
    return jsonify({'bookmarks': bookmarks})

@app.route('/api/bookmarks', methods=['POST'])
@token_required
def add_bookmark():
    """Bookmark a question."""
    data = request.get_json() or {}
    question_id = data.get('question_id')
    note = data.get('note')
    if not question_id:
        return jsonify({'error': 'question_id is required'}), 400
    conn = get_db()
    c = conn.cursor()
    try:
        c.execute('INSERT OR IGNORE INTO bookmarks (user_id, question_id, note) VALUES (?, ?, ?)',
                  (request.user_id, question_id, note))
        conn.commit()
        bookmark_id = c.lastrowid
    except sqlite3.IntegrityError:
        conn.close()
        return jsonify({'error': 'Question not found'}), 404
    conn.close()
    return jsonify({'success': True, 'id': bookmark_id})

@app.route('/api/bookmarks/<int:question_id>', methods=['DELETE'])
@token_required
def remove_bookmark(question_id):
    """Remove a bookmark."""
    conn = get_db()
    c = conn.cursor()
    c.execute('DELETE FROM bookmarks WHERE user_id = ? AND question_id = ?',
              (request.user_id, question_id))
    conn.commit()
    conn.close()
    return jsonify({'success': True})

@app.route('/api/certifications/<int:cert_id>/readiness', methods=['GET'])
@token_required
def get_cert_readiness(cert_id):
    """Get certification readiness dashboard data aggregating domains, simulations, study time, and prediction."""
    conn = get_db()
    c = conn.cursor()

    # Get certification info
    c.execute('SELECT * FROM certifications WHERE id = ?', (cert_id,))
    cert_row = c.fetchone()
    if not cert_row:
        conn.close()
        return jsonify({'error': 'Certification not found'}), 404
    cert = dict(cert_row)

    # Domain performance
    c.execute('''SELECT d.id, d.name, d.code, d.weight,
        COALESCE(SUM(qp.times_correct), 0) as correct,
        COALESCE(SUM(qp.times_seen), 0) as seen
    FROM domains d
    LEFT JOIN question_domains qd ON qd.domain_id = d.id
    LEFT JOIN question_performance qp ON qp.question_id = qd.question_id AND qp.user_id = ?
    WHERE d.certification_id = ? AND d.parent_domain_id IS NULL
    GROUP BY d.id
    ORDER BY d.sort_order''', (request.user_id, cert_id))
    domain_rows = c.fetchall()

    domains = []
    weighted_score_sum = 0
    total_weight = 0
    for row in domain_rows:
        d = dict(row)
        seen = d['seen']
        correct = d['correct']
        score = round(correct / seen * 100, 1) if seen > 0 else 0
        if seen == 0:
            status = 'unseen'
        elif score >= 80:
            status = 'strong'
        elif score >= 60:
            status = 'moderate'
        else:
            status = 'weak'
        weight = d['weight'] or 0
        # Always include domain weight — unseen domains score 0, not excluded
        weighted_score_sum += score * weight
        total_weight += weight
        domains.append({
            'name': d['name'],
            'code': d['code'],
            'weight': weight,
            'score': score,
            'seen': seen,
            'correct': correct,
            'status': status
        })

    overall_score = round(weighted_score_sum / total_weight, 1) if total_weight > 0 else 0

    # Recent simulations (last 5)
    c.execute('''SELECT score, total, percentage, passed, time_taken, created_at
        FROM exam_simulations
        WHERE user_id = ? AND certification_id = ?
        ORDER BY created_at DESC LIMIT 5''', (request.user_id, cert_id))
    sim_rows = c.fetchall()
    simulations = []
    for row in sim_rows:
        s = dict(row)
        simulations.append({
            'score': s['score'],
            'total': s['total'],
            'percentage': round(s['percentage'], 1) if s['percentage'] else 0,
            'passed': bool(s['passed']),
            'date': s['created_at']
        })

    # Study time (last 30 days)
    c.execute('''SELECT COALESCE(SUM(duration_seconds), 0) as total_seconds,
        COUNT(*) as session_count
    FROM study_sessions
    WHERE user_id = ? AND certification_id = ?
    AND started_at >= datetime('now', '-30 days')''', (request.user_id, cert_id))
    study_row = dict(c.fetchone())
    total_seconds = study_row['total_seconds']
    study_time = {
        'total_hours': round(total_seconds / 3600, 1) if total_seconds else 0,
        'sessions': study_row['session_count']
    }

    # Prediction logic
    passing_score = cert.get('passing_score', 70)
    passing_scale = cert.get('passing_scale', 'percentage')

    sim_percentages = [s['percentage'] for s in simulations if s['percentage'] is not None]
    sim_avg = sum(sim_percentages) / len(sim_percentages) if sim_percentages else 0

    # Domain coverage: % of domains with >= 10 questions answered
    domains_with_coverage = sum(1 for d in domains if d['seen'] >= 10)
    domain_coverage = domains_with_coverage / len(domains) if domains else 0

    # Trend based on last 3 sims (list is ordered DESC, index 0 = most recent)
    trend = 'stable'
    if len(sim_percentages) >= 3:
        recent_3 = sim_percentages[:3]
        if recent_3[0] > recent_3[2] + 3:
            trend = 'improving'
        elif recent_3[0] < recent_3[2] - 3:
            trend = 'declining'

    # Determine pass likelihood using percentage-based comparison
    passing_pct = passing_score
    if passing_scale == 'scaled':
        passing_pct = 75

    likely_pass = False
    confidence = 'low'
    if len(sim_percentages) >= 2:
        if sim_avg >= passing_pct + 5 and domain_coverage >= 0.7:
            likely_pass = True
            confidence = 'high' if sim_avg >= passing_pct + 15 and domain_coverage >= 0.9 else 'moderate'
        elif sim_avg >= passing_pct and domain_coverage >= 0.5:
            likely_pass = True
            confidence = 'low'
        else:
            likely_pass = False
            confidence = 'moderate' if sim_avg >= passing_pct - 10 else 'high'
    elif len(sim_percentages) == 1:
        likely_pass = sim_avg >= passing_pct
        confidence = 'low'

    prediction = {
        'likely_pass': likely_pass,
        'confidence': confidence,
        'trend': trend
    }

    conn.close()
    return jsonify({
        'certification': {
            'name': cert['name'],
            'passing_score': passing_score,
            'passing_scale': passing_scale
        },
        'overall_score': overall_score,
        'domains': domains,
        'simulations': simulations,
        'study_time': study_time,
        'prediction': prediction
    })

# ==================== Objective Confidence & Study Resources ====================

@app.route('/api/certifications/<int:cert_id>/objectives', methods=['GET'])
@token_required
def get_objectives(cert_id):
    """Get all sub-objectives for a certification with user's confidence ratings."""
    conn = get_db()
    c = conn.cursor()

    # Fetch top-level domains and their sub-objectives
    c.execute('''SELECT d.id, d.name, d.code, d.weight, d.parent_domain_id, d.sort_order
        FROM domains d
        WHERE d.certification_id = ?
        ORDER BY d.parent_domain_id NULLS FIRST, d.sort_order''', (cert_id,))
    all_domains = [dict(r) for r in c.fetchall()]

    # Fetch user's confidence ratings
    domain_ids = [d['id'] for d in all_domains]
    confidence_map = {}
    if domain_ids:
        placeholders = ','.join('?' * len(domain_ids))
        c.execute(f'''SELECT domain_id, confidence FROM objective_confidence
            WHERE user_id = ? AND domain_id IN ({placeholders})''',
            [request.user_id] + domain_ids)
        for row in c.fetchall():
            confidence_map[row['domain_id']] = row['confidence']

    # Build nested structure: top-level domains with children
    top_domains = [d for d in all_domains if d['parent_domain_id'] is None]
    children_map = {}
    for d in all_domains:
        if d['parent_domain_id'] is not None:
            children_map.setdefault(d['parent_domain_id'], []).append(d)

    result = []
    for td in top_domains:
        children = children_map.get(td['id'], [])
        obj_list = []
        for child in children:
            obj_list.append({
                'domain_id': child['id'],
                'code': child['code'],
                'name': child['name'],
                'confidence': confidence_map.get(child['id'], 0),
            })
        result.append({
            'domain_id': td['id'],
            'code': td['code'],
            'name': td['name'],
            'weight': td['weight'],
            'confidence': confidence_map.get(td['id'], 0),
            'objectives': obj_list,
        })

    conn.close()
    return jsonify({'domains': result})


@app.route('/api/certifications/<int:cert_id>/objectives', methods=['PUT'])
@token_required
def update_objectives(cert_id):
    """Update confidence ratings for objectives. Body: {ratings: {domain_id: confidence}}"""
    data = request.get_json()
    ratings = data.get('ratings', {})
    if not ratings:
        return jsonify({'error': 'No ratings provided'}), 400

    conn = get_db()
    c = conn.cursor()
    now = datetime.now()
    with _db_write_lock:
        for domain_id_str, confidence in ratings.items():
            domain_id = int(domain_id_str)
            confidence = max(0, min(3, int(confidence)))  # clamp 0-3
            c.execute('''INSERT INTO objective_confidence (user_id, domain_id, confidence, updated_at)
                VALUES (?, ?, ?, ?)
                ON CONFLICT(user_id, domain_id) DO UPDATE SET
                confidence = excluded.confidence, updated_at = excluded.updated_at''',
                (request.user_id, domain_id, confidence, now))
        conn.commit()
    conn.close()
    return jsonify({'message': 'Updated', 'count': len(ratings)})


@app.route('/api/certifications/<int:cert_id>/resources', methods=['GET'])
@token_required
def get_study_resources(cert_id):
    """Get study resources for a certification, optionally filtered by domain."""
    domain_id = request.args.get('domain_id', type=int)
    conn = get_db()
    c = conn.cursor()
    if domain_id:
        c.execute('''SELECT * FROM study_resources
            WHERE certification_id = ? AND (domain_id = ? OR domain_id IS NULL)
            ORDER BY sort_order''', (cert_id, domain_id))
    else:
        c.execute('''SELECT * FROM study_resources
            WHERE certification_id = ?
            ORDER BY sort_order''', (cert_id,))
    resources = [dict(r) for r in c.fetchall()]
    conn.close()
    return jsonify({'resources': resources})


@app.route('/api/community/quizzes', methods=['GET'])
@token_required
def community_quizzes():
    """Return public quizzes, optionally filtered by search query."""
    q_search = (request.args.get('q') or '').strip()
    limit = min(int(request.args.get('limit', 50)), 100)

    conn = get_db()
    c = conn.cursor()

    sql = '''
        SELECT qz.id, qz.title, qz.description,
               (SELECT COUNT(*) FROM questions WHERE quiz_id = qz.id) AS question_count,
               u.username
        FROM quizzes qz
        LEFT JOIN users u ON u.id = qz.user_id
        WHERE qz.is_public = 1
    '''
    params = []
    if q_search:
        sql += ' AND (qz.title LIKE ? OR qz.description LIKE ?)'
        params.extend([f'%{q_search}%', f'%{q_search}%'])
    sql += ' ORDER BY qz.last_modified DESC LIMIT ?'
    params.append(limit)

    c.execute(sql, params)
    rows = c.fetchall()
    conn.close()

    quizzes = [{
        'id': r['id'],
        'title': r['title'],
        'description': r['description'],
        'question_count': r['question_count'],
        'username': r['username'],
    } for r in rows]
    return jsonify({'quizzes': quizzes})

@app.route('/api/community/quizzes/<int:quiz_id>/copy', methods=['POST'])
@token_required
def copy_community_quiz(quiz_id):
    """Copy a public quiz to the current user's library."""
    conn = get_db()
    c = conn.cursor()

    # Verify the quiz exists and is public
    c.execute('SELECT * FROM quizzes WHERE id = ? AND is_public = 1', (quiz_id,))
    source = c.fetchone()
    if not source:
        conn.close()
        return jsonify({'error': 'Quiz not found or not public'}), 404

    # Don't allow copying your own quiz
    if source['user_id'] == request.user_id:
        conn.close()
        return jsonify({'error': 'This quiz is already in your library'}), 409

    # Read normalized questions from source
    questions_list = _read_questions_for_quiz(c, quiz_id) if source['is_migrated'] else json.loads(source['questions'])

    # Create new quiz for the current user
    title = source['title']
    c.execute('''INSERT INTO quizzes (user_id, title, description, questions, color, is_migrated)
                 VALUES (?, ?, ?, ?, ?, 1)''',
              (request.user_id, title, source['description'] or '',
               json.dumps(questions_list), source['color'] or '#6366f1'))
    new_quiz_id = c.lastrowid
    _insert_questions_for_quiz(c, new_quiz_id, questions_list)

    conn.commit()
    conn.close()
    return jsonify({'message': 'Quiz copied to your library', 'quiz_id': new_quiz_id}), 201

@app.route('/api/sample-quiz', methods=['POST'])
@token_required
def create_sample_quiz():
    """Create a sample quiz so new users can try the app immediately."""
    conn = get_db()
    c = conn.cursor()

    # Only create if user has zero quizzes (true first-timer)
    c.execute('SELECT COUNT(*) as cnt FROM quizzes WHERE user_id = ?', (request.user_id,))
    if c.fetchone()['cnt'] > 0:
        conn.close()
        return jsonify({'message': 'You already have quizzes'}), 200

    sample_questions = [
        {
            'question': 'What does HTTP stand for?',
            'type': 'choice',
            'options': ['HyperText Transfer Protocol', 'High Tech Transfer Protocol',
                        'HyperText Transmission Process', 'Home Tool Transfer Protocol'],
            'correct': 0,
            'explanation': 'HTTP stands for HyperText Transfer Protocol. It is the foundation of data communication on the web.'
        },
        {
            'question': 'Which layer of the OSI model does a router operate at?',
            'type': 'choice',
            'options': ['Layer 1 - Physical', 'Layer 2 - Data Link',
                        'Layer 3 - Network', 'Layer 4 - Transport'],
            'correct': 2,
            'explanation': 'Routers operate at Layer 3 (Network layer) of the OSI model, making forwarding decisions based on IP addresses.'
        },
        {
            'question': 'TCP is a connection-oriented protocol.',
            'type': 'truefalse',
            'options': ['True', 'False'],
            'correct': 0,
            'explanation': 'TCP (Transmission Control Protocol) is connection-oriented. It establishes a connection using a three-way handshake before data transfer.'
        },
        {
            'question': 'Which of these are valid IP address classes? (Select all that apply)',
            'type': 'multiselect',
            'options': ['Class A', 'Class B', 'Class F', 'Class C'],
            'correct': [0, 1, 3],
            'explanation': 'IP addresses are divided into classes A, B, C, D, and E. There is no Class F.'
        },
        {
            'question': 'What is the default port for HTTPS?',
            'type': 'choice',
            'options': ['80', '443', '8080', '22'],
            'correct': 1,
            'explanation': 'HTTPS uses port 443 by default. Port 80 is used by HTTP, 8080 is an alternative HTTP port, and 22 is for SSH.'
        },
    ]

    c.execute('''INSERT INTO quizzes (user_id, title, description, questions, color, is_migrated)
                 VALUES (?, ?, ?, ?, ?, 1)''',
              (request.user_id, 'Networking Basics (Sample Quiz)',
               'A sample quiz to help you explore the app. Feel free to edit or delete it!',
               json.dumps(sample_questions), '#6366f1'))
    quiz_id = c.lastrowid
    _insert_questions_for_quiz(c, quiz_id, sample_questions)
    conn.commit()
    conn.close()
    return jsonify({'message': 'Sample quiz created', 'quiz_id': quiz_id}), 201

@app.route('/api/health')
def health():
    return jsonify({'status': 'ok'})

# === Admin Routes ===

def require_admin_token(f):
    """Decorator requiring X-Admin-Token header matching ADMIN_TOKEN."""
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('X-Admin-Token', '')
        if not secrets.compare_digest(token, ADMIN_TOKEN):
            return jsonify({'error': 'Unauthorized'}), 401
        return f(*args, **kwargs)
    return decorated

@app.route('/api/admin/seed', methods=['POST'])
@require_admin_token
def admin_seed():
    """Manually trigger certification seeding (protected by admin token)."""
    seed_certifications()
    conn = get_db()
    c = conn.cursor()
    c.execute('SELECT COUNT(*) as cnt FROM certifications')
    count = c.fetchone()['cnt']
    conn.close()
    return jsonify({'message': f'Seed complete. {count} certifications in database.'})

@app.route('/api/admin/cleanup-sessions', methods=['POST'])
@require_admin_token
def cleanup_sessions():
    """Delete expired sessions. Run daily via cron."""
    conn = get_db()
    c = conn.cursor()
    with _db_write_lock:
        c.execute('DELETE FROM sessions WHERE expires_at < datetime("now")')
        deleted = c.rowcount
        conn.commit()
    conn.close()
    return jsonify({'message': f'Deleted {deleted} expired sessions.'})

@app.route('/api/admin/events', methods=['GET'])
@require_admin_token
def admin_events():
    """Event log summary: events per user per day, most-used features, drop-off points."""
    conn = get_db()
    c = conn.cursor()
    try:
        # Events per user per day (last 30 days)
        c.execute('''
            SELECT u.username, date(e.created_at) as day, e.event, COUNT(*) as count
            FROM events e
            LEFT JOIN users u ON e.user_id = u.id
            WHERE e.created_at >= datetime("now", "-30 days")
            GROUP BY e.user_id, day, e.event
            ORDER BY day DESC, count DESC
        ''')
        daily = [dict(r) for r in c.fetchall()]

        # Most used features (all time)
        c.execute('''
            SELECT event, COUNT(*) as count
            FROM events
            GROUP BY event
            ORDER BY count DESC
        ''')
        feature_counts = [dict(r) for r in c.fetchall()]

        # Active users (last 7 days)
        c.execute('''
            SELECT COUNT(DISTINCT user_id) as active_users
            FROM events
            WHERE created_at >= datetime("now", "-7 days")
        ''')
        active = c.fetchone()['active_users']

        conn.close()
        return jsonify({
            'active_users_7d': active,
            'feature_counts': feature_counts,
            'daily_breakdown': daily
        })
    except Exception as e:
        conn.close()
        return jsonify({'error': str(e)}), 500

# === Event Logging ===

@app.route('/api/events', methods=['POST'])
@token_required
def log_event():
    """Log a frontend event. Called from the browser on key user actions."""
    data = request.get_json()
    event_name = data.get('event', '').strip()
    if not event_name:
        return jsonify({'error': 'event required'}), 400
    metadata = data.get('metadata')
    conn = get_db()
    c = conn.cursor()
    with _db_write_lock:
        c.execute(
            'INSERT INTO events (user_id, event, metadata) VALUES (?, ?, ?)',
            (request.user_id, event_name, json.dumps(metadata) if metadata else None)
        )
        conn.commit()
    conn.close()
    return jsonify({'ok': True})

def seed_certifications():
    """Seed the certifications and domains tables with initial IT certification data."""
    conn = get_db()
    c = conn.cursor()

    # Check if already seeded
    c.execute('SELECT COUNT(*) as cnt FROM certifications')
    if c.fetchone()['cnt'] > 0:
        conn.close()
        return

    certs = [
        {
            'code': 'comptia-a-core1-221-1101',
            'name': 'CompTIA A+ Core 1 (220-1101)',
            'vendor': 'CompTIA',
            'description': 'Hardware, networking, mobile devices, virtualization, and troubleshooting',
            'passing_score': 675, 'passing_scale': 'scaled',
            'exam_duration_minutes': 90, 'total_questions': 90,
            'domains': [
                {'name': 'Mobile Devices', 'code': '1.0', 'weight': 0.15},
                {'name': 'Networking', 'code': '2.0', 'weight': 0.20},
                {'name': 'Hardware', 'code': '3.0', 'weight': 0.25},
                {'name': 'Virtualization and Cloud Computing', 'code': '4.0', 'weight': 0.11},
                {'name': 'Hardware and Network Troubleshooting', 'code': '5.0', 'weight': 0.29},
            ]
        },
        {
            'code': 'comptia-a-core2-221-1102',
            'name': 'CompTIA A+ Core 2 (220-1102)',
            'vendor': 'CompTIA',
            'description': 'Operating systems, security, software troubleshooting, and operational procedures',
            'passing_score': 700, 'passing_scale': 'scaled',
            'exam_duration_minutes': 90, 'total_questions': 90,
            'domains': [
                {'name': 'Operating Systems', 'code': '1.0', 'weight': 0.31},
                {'name': 'Security', 'code': '2.0', 'weight': 0.25},
                {'name': 'Software Troubleshooting', 'code': '3.0', 'weight': 0.22},
                {'name': 'Operational Procedures', 'code': '4.0', 'weight': 0.22},
            ]
        },
        {
            'code': 'comptia-net-n10-009',
            'name': 'CompTIA Network+ (N10-009)',
            'vendor': 'CompTIA',
            'description': 'Networking concepts, infrastructure, security, and troubleshooting',
            'passing_score': 720, 'passing_scale': 'scaled',
            'exam_duration_minutes': 90, 'total_questions': 90,
            'domains': [
                {'name': 'Networking Concepts', 'code': '1.0', 'weight': 0.23},
                {'name': 'Network Implementation', 'code': '2.0', 'weight': 0.19},
                {'name': 'Network Operations', 'code': '3.0', 'weight': 0.16},
                {'name': 'Network Security', 'code': '4.0', 'weight': 0.19},
                {'name': 'Network Troubleshooting', 'code': '5.0', 'weight': 0.23},
            ]
        },
        {
            'code': 'comptia-sec-sy0-701',
            'name': 'CompTIA Security+ (SY0-701)',
            'vendor': 'CompTIA',
            'description': 'Security concepts, threats, architecture, operations, and program management',
            'passing_score': 750, 'passing_scale': 'scaled',
            'exam_duration_minutes': 90, 'total_questions': 90,
            'domains': [
                {'name': 'General Security Concepts', 'code': '1.0', 'weight': 0.12},
                {'name': 'Threats, Vulnerabilities, and Mitigations', 'code': '2.0', 'weight': 0.22},
                {'name': 'Security Architecture', 'code': '3.0', 'weight': 0.18},
                {'name': 'Security Operations', 'code': '4.0', 'weight': 0.28},
                {'name': 'Security Program Management and Oversight', 'code': '5.0', 'weight': 0.20},
            ]
        },
        {
            'code': 'comptia-cysa-cs0-003',
            'name': 'CompTIA CySA+ (CS0-003)',
            'vendor': 'CompTIA',
            'description': 'Security operations, vulnerability management, incident response, and reporting',
            'passing_score': 750, 'passing_scale': 'scaled',
            'exam_duration_minutes': 165, 'total_questions': 85,
            'domains': [
                {'name': 'Security Operations', 'code': '1.0', 'weight': 0.33},
                {'name': 'Vulnerability Management', 'code': '2.0', 'weight': 0.30},
                {'name': 'Incident Response and Management', 'code': '3.0', 'weight': 0.20},
                {'name': 'Reporting and Communication', 'code': '4.0', 'weight': 0.17},
            ]
        },
        {
            'code': 'comptia-pentest-pt0-003',
            'name': 'CompTIA PenTest+ (PT0-003)',
            'vendor': 'CompTIA',
            'description': 'Penetration testing, vulnerability assessment, and management',
            'passing_score': 750, 'passing_scale': 'scaled',
            'exam_duration_minutes': 165, 'total_questions': 85,
            'domains': [
                {'name': 'Planning and Scoping', 'code': '1.0', 'weight': 0.14},
                {'name': 'Information Gathering and Vulnerability Scanning', 'code': '2.0', 'weight': 0.22},
                {'name': 'Attacks and Exploits', 'code': '3.0', 'weight': 0.30},
                {'name': 'Reporting and Communication', 'code': '4.0', 'weight': 0.18},
                {'name': 'Tools and Code Analysis', 'code': '5.0', 'weight': 0.16},
            ]
        },
        {
            'code': 'comptia-casp-cas-004',
            'name': 'CompTIA CASP+ (CAS-004)',
            'vendor': 'CompTIA',
            'description': 'Advanced security architecture, operations, engineering, and governance',
            'passing_score': None, 'passing_scale': 'pass_fail',
            'exam_duration_minutes': 165, 'total_questions': 90,
            'domains': [
                {'name': 'Security Architecture', 'code': '1.0', 'weight': 0.29},
                {'name': 'Security Operations', 'code': '2.0', 'weight': 0.30},
                {'name': 'Security Engineering and Cryptography', 'code': '3.0', 'weight': 0.26},
                {'name': 'Governance, Risk, and Compliance', 'code': '4.0', 'weight': 0.15},
            ]
        },
        {
            'code': 'cisco-ccna-200-301',
            'name': 'Cisco CCNA (200-301)',
            'vendor': 'Cisco',
            'description': 'Network fundamentals, access, IP connectivity, services, security, and automation',
            'passing_score': 825, 'passing_scale': 'scaled',
            'exam_duration_minutes': 120, 'total_questions': 100,
            'domains': [
                {'name': 'Network Fundamentals', 'code': '1.0', 'weight': 0.20},
                {'name': 'Network Access', 'code': '2.0', 'weight': 0.20},
                {'name': 'IP Connectivity', 'code': '3.0', 'weight': 0.25},
                {'name': 'IP Services', 'code': '4.0', 'weight': 0.10},
                {'name': 'Security Fundamentals', 'code': '5.0', 'weight': 0.15},
                {'name': 'Automation and Programmability', 'code': '6.0', 'weight': 0.10},
            ]
        },
        {
            'code': 'aws-saa-c03',
            'name': 'AWS Solutions Architect Associate (SAA-C03)',
            'vendor': 'AWS',
            'description': 'Designing resilient, performant, secure, and cost-optimized architectures',
            'passing_score': 720, 'passing_scale': 'scaled',
            'exam_duration_minutes': 130, 'total_questions': 65,
            'domains': [
                {'name': 'Design Secure Architectures', 'code': '1.0', 'weight': 0.30},
                {'name': 'Design Resilient Architectures', 'code': '2.0', 'weight': 0.26},
                {'name': 'Design High-Performing Architectures', 'code': '3.0', 'weight': 0.24},
                {'name': 'Design Cost-Optimized Architectures', 'code': '4.0', 'weight': 0.20},
            ]
        },
        {
            'code': 'aws-dva-c02',
            'name': 'AWS Developer Associate (DVA-C02)',
            'vendor': 'AWS',
            'description': 'Development, deployment, security, and troubleshooting of AWS applications',
            'passing_score': 720, 'passing_scale': 'scaled',
            'exam_duration_minutes': 130, 'total_questions': 65,
            'domains': [
                {'name': 'Development with AWS Services', 'code': '1.0', 'weight': 0.32},
                {'name': 'Security', 'code': '2.0', 'weight': 0.26},
                {'name': 'Deployment', 'code': '3.0', 'weight': 0.24},
                {'name': 'Troubleshooting and Optimization', 'code': '4.0', 'weight': 0.18},
            ]
        },
        {
            'code': 'aws-soa-c02',
            'name': 'AWS SysOps Administrator Associate (SOA-C02)',
            'vendor': 'AWS',
            'description': 'Deployment, management, networking, security, and automation on AWS',
            'passing_score': 720, 'passing_scale': 'scaled',
            'exam_duration_minutes': 130, 'total_questions': 65,
            'domains': [
                {'name': 'Monitoring, Logging, and Remediation', 'code': '1.0', 'weight': 0.20},
                {'name': 'Reliability and Business Continuity', 'code': '2.0', 'weight': 0.16},
                {'name': 'Deployment, Provisioning, and Automation', 'code': '3.0', 'weight': 0.18},
                {'name': 'Security and Compliance', 'code': '4.0', 'weight': 0.16},
                {'name': 'Networking and Content Delivery', 'code': '5.0', 'weight': 0.18},
                {'name': 'Cost and Performance Optimization', 'code': '6.0', 'weight': 0.12},
            ]
        },
        {
            'code': 'aws-clf-c02',
            'name': 'AWS Cloud Practitioner (CLF-C02)',
            'vendor': 'AWS',
            'description': 'Cloud concepts, security, technology, and billing',
            'passing_score': 700, 'passing_scale': 'scaled',
            'exam_duration_minutes': 90, 'total_questions': 65,
            'domains': [
                {'name': 'Cloud Concepts', 'code': '1.0', 'weight': 0.24},
                {'name': 'Security and Compliance', 'code': '2.0', 'weight': 0.30},
                {'name': 'Cloud Technology and Services', 'code': '3.0', 'weight': 0.34},
                {'name': 'Billing, Pricing, and Support', 'code': '4.0', 'weight': 0.12},
            ]
        },
        {
            'code': 'az-900',
            'name': 'Microsoft Azure Fundamentals (AZ-900)',
            'vendor': 'Microsoft',
            'description': 'Cloud concepts, Azure architecture, management, and governance',
            'passing_score': 700, 'passing_scale': 'scaled',
            'exam_duration_minutes': 65, 'total_questions': 50,
            'domains': [
                {'name': 'Cloud Concepts', 'code': '1.0', 'weight': 0.25},
                {'name': 'Azure Architecture and Services', 'code': '2.0', 'weight': 0.35},
                {'name': 'Azure Management and Governance', 'code': '3.0', 'weight': 0.30},
            ]
        },
        {
            'code': 'az-104',
            'name': 'Microsoft Azure Administrator (AZ-104)',
            'vendor': 'Microsoft',
            'description': 'Managing Azure identities, governance, storage, compute, and networks',
            'passing_score': 700, 'passing_scale': 'scaled',
            'exam_duration_minutes': 100, 'total_questions': 55,
            'domains': [
                {'name': 'Manage Azure Identities and Governance', 'code': '1.0', 'weight': 0.20},
                {'name': 'Implement and Manage Storage', 'code': '2.0', 'weight': 0.15},
                {'name': 'Deploy and Manage Azure Compute Resources', 'code': '3.0', 'weight': 0.20},
                {'name': 'Implement and Manage Virtual Networking', 'code': '4.0', 'weight': 0.15},
                {'name': 'Monitor and Maintain Azure Resources', 'code': '5.0', 'weight': 0.10},
            ]
        },
        {
            'code': 'gcp-ace',
            'name': 'Google Cloud Associate Cloud Engineer',
            'vendor': 'Google Cloud',
            'description': 'Deploying, monitoring, and managing solutions on Google Cloud',
            'passing_score': 70, 'passing_scale': 'percentage',
            'exam_duration_minutes': 120, 'total_questions': 50,
            'domains': [
                {'name': 'Setting Up a Cloud Solution Environment', 'code': '1.0', 'weight': 0.20},
                {'name': 'Planning and Configuring a Cloud Solution', 'code': '2.0', 'weight': 0.20},
                {'name': 'Deploying and Implementing a Cloud Solution', 'code': '3.0', 'weight': 0.25},
                {'name': 'Ensuring Successful Operation of a Cloud Solution', 'code': '4.0', 'weight': 0.20},
                {'name': 'Configuring Access and Security', 'code': '5.0', 'weight': 0.15},
            ]
        },
    ]

    for cert_data in certs:
        domains = cert_data.pop('domains')
        try:
            c.execute('''INSERT INTO certifications (code, name, vendor, description, passing_score, passing_scale, exam_duration_minutes, total_questions)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
                (cert_data['code'], cert_data['name'], cert_data['vendor'], cert_data['description'],
                 cert_data['passing_score'], cert_data['passing_scale'],
                 cert_data['exam_duration_minutes'], cert_data['total_questions']))
            cert_id = c.lastrowid
            for i, d in enumerate(domains):
                c.execute('''INSERT INTO domains (certification_id, name, code, weight, sort_order)
                    VALUES (?, ?, ?, ?, ?)''',
                    (cert_id, d['name'], d['code'], d['weight'], i))
        except sqlite3.IntegrityError:
            continue  # Already exists

    conn.commit()
    conn.close()


def seed_sub_objectives():
    """Seed sub-objectives (child domains) for each certification's top-level domains."""
    conn = get_db()
    c = conn.cursor()

    # Check if sub-objectives already exist
    c.execute('SELECT COUNT(*) as cnt FROM domains WHERE parent_domain_id IS NOT NULL')
    if c.fetchone()['cnt'] > 0:
        conn.close()
        return

    # Sub-objectives keyed by certification code, then parent domain code
    sub_objectives = {
        # ── CompTIA Security+ (SY0-701) ──
        'comptia-sec-sy0-701': {
            '1.0': [
                ('1.1', 'Security controls'),
                ('1.2', 'Fundamental security concepts'),
                ('1.3', 'Change management'),
                ('1.4', 'Cryptography'),
            ],
            '2.0': [
                ('2.1', 'Threat actors'),
                ('2.2', 'Attack surfaces'),
                ('2.3', 'Vulnerability types'),
                ('2.4', 'Indicators of malicious activity'),
                ('2.5', 'Mitigation techniques'),
            ],
            '3.0': [
                ('3.1', 'Architecture models'),
                ('3.2', 'Infrastructure considerations'),
                ('3.3', 'Secure data'),
                ('3.4', 'Resiliency and recovery'),
            ],
            '4.0': [
                ('4.1', 'Apply security techniques'),
                ('4.2', 'Asset management'),
                ('4.3', 'Vulnerability management'),
                ('4.4', 'Security alerting and monitoring'),
                ('4.5', 'Modify enterprise capabilities'),
                ('4.6', 'Automation and orchestration'),
                ('4.7', 'Incident response'),
                ('4.8', 'Digital forensics'),
                ('4.9', 'Data sources'),
            ],
            '5.0': [
                ('5.1', 'Security governance'),
                ('5.2', 'Risk management'),
                ('5.3', 'Third-party risk'),
                ('5.4', 'Compliance'),
            ],
        },

        # ── CompTIA Network+ (N10-009) ──
        'comptia-net-n10-009': {
            '1.0': [
                ('1.1', 'OSI model'),
                ('1.2', 'Network topologies'),
                ('1.3', 'Connectors and cabling'),
                ('1.4', 'IP subnetting'),
                ('1.5', 'Ports and protocols'),
                ('1.6', 'Network services'),
                ('1.7', 'Corporate network architecture'),
                ('1.8', 'Cloud concepts'),
            ],
            '2.0': [
                ('2.1', 'Routing technologies'),
                ('2.2', 'Switch features'),
                ('2.3', 'Wireless standards'),
                ('2.4', 'Configuration concepts'),
            ],
            '3.0': [
                ('3.1', 'Network monitoring'),
                ('3.2', 'Documentation and diagrams'),
                ('3.3', 'High availability and disaster recovery'),
                ('3.4', 'Network hardening'),
            ],
            '4.0': [
                ('4.1', 'Security concepts'),
                ('4.2', 'Attacks'),
                ('4.3', 'Security technologies'),
                ('4.4', 'Remote access methods'),
            ],
            '5.0': [
                ('5.1', 'Methodology'),
                ('5.2', 'Cable connectivity'),
                ('5.3', 'Network issues'),
                ('5.4', 'Wireless issues'),
                ('5.5', 'General networking issues'),
            ],
        },

        # ── CompTIA A+ Core 1 (220-1101) ──
        'comptia-a-core1-221-1101': {
            '1.0': [
                ('1.1', 'Laptop hardware'),
                ('1.2', 'Display components'),
                ('1.3', 'Mobile device accessories'),
                ('1.4', 'Cellular standards'),
            ],
            '2.0': [
                ('2.1', 'TCP/IP'),
                ('2.2', 'Network hardware'),
                ('2.3', 'Wireless protocols'),
                ('2.4', 'DNS configuration'),
                ('2.5', 'Network types'),
                ('2.6', 'Cable types'),
                ('2.7', 'Networking tools'),
                ('2.8', 'Copper and fiber'),
            ],
            '3.0': [
                ('3.1', 'Cable types'),
                ('3.2', 'RAM types'),
                ('3.3', 'Storage devices'),
                ('3.4', 'Motherboard components'),
                ('3.5', 'Power supplies'),
                ('3.6', 'Multifunction devices'),
                ('3.7', 'Print technologies'),
            ],
            '4.0': [
                ('4.1', 'Cloud computing concepts'),
                ('4.2', 'Client-side virtualization'),
            ],
            '5.0': [
                ('5.1', 'Methodology'),
                ('5.2', 'Troubleshoot hardware'),
                ('5.3', 'Troubleshoot storage'),
                ('5.4', 'Troubleshoot video/display'),
                ('5.5', 'Troubleshoot mobile'),
                ('5.6', 'Troubleshoot printers'),
                ('5.7', 'Troubleshoot networking'),
            ],
        },

        # ── CompTIA A+ Core 2 (220-1102) ──
        'comptia-a-core2-221-1102': {
            '1.0': [
                ('1.1', 'Windows editions'),
                ('1.2', 'Command-line tools'),
                ('1.3', 'OS features and tools'),
                ('1.4', 'OS management'),
                ('1.5', 'Install Windows'),
                ('1.6', 'macOS features'),
                ('1.7', 'Linux features'),
                ('1.8', 'OS installations'),
            ],
            '2.0': [
                ('2.1', 'Physical security'),
                ('2.2', 'Wireless security'),
                ('2.3', 'Malware detection'),
                ('2.4', 'Social engineering'),
                ('2.5', 'Windows security'),
                ('2.6', 'Security settings'),
                ('2.7', 'Workstation security'),
                ('2.8', 'Mobile device security'),
                ('2.9', 'Data destruction'),
                ('2.10', 'SOHO network security'),
            ],
            '3.0': [
                ('3.1', 'Troubleshoot Windows'),
                ('3.2', 'PC security issues'),
                ('3.3', 'Malware removal'),
                ('3.4', 'Mobile OS issues'),
                ('3.5', 'Mobile security issues'),
            ],
            '4.0': [
                ('4.1', 'Documentation'),
                ('4.2', 'Change management'),
                ('4.3', 'Disaster recovery'),
                ('4.4', 'Safety procedures'),
                ('4.5', 'Environmental impacts'),
                ('4.6', 'Privacy and licensing'),
                ('4.7', 'Communication and professionalism'),
                ('4.8', 'Scripting basics'),
                ('4.9', 'Remote access technologies'),
            ],
        },

        # ── Cisco CCNA (200-301) ──
        'cisco-ccna-200-301': {
            '1.0': [
                ('1.1', 'Network components'),
                ('1.2', 'Network topology'),
                ('1.3', 'Physical interfaces'),
                ('1.4', 'TCP vs UDP'),
                ('1.5', 'IPv4 addressing'),
                ('1.6', 'IPv6 addressing'),
                ('1.7', 'Wireless principles'),
                ('1.8', 'Switching concepts'),
            ],
            '2.0': [
                ('2.1', 'VLANs'),
                ('2.2', 'Interswitch connectivity'),
                ('2.3', 'Layer 2 discovery protocols'),
                ('2.4', 'EtherChannel'),
                ('2.5', 'Spanning tree'),
                ('2.6', 'Wireless architectures'),
                ('2.7', 'WLAN components'),
                ('2.8', 'AP modes'),
                ('2.9', 'Physical access connections'),
            ],
            '3.0': [
                ('3.1', 'Routing concepts'),
                ('3.2', 'Static routing'),
                ('3.3', 'OSPF'),
                ('3.4', 'First hop redundancy'),
            ],
            '4.0': [
                ('4.1', 'NAT'),
                ('4.2', 'NTP'),
                ('4.3', 'DHCP/DNS'),
                ('4.4', 'SNMP'),
                ('4.5', 'Syslog'),
                ('4.6', 'SSH/TFTP/FTP'),
            ],
            '5.0': [
                ('5.1', 'Security concepts'),
                ('5.2', 'Device access security'),
                ('5.3', 'Security program elements'),
                ('5.4', 'Access control lists'),
                ('5.5', 'Wireless security'),
                ('5.6', 'Layer 2 security'),
            ],
            '6.0': [
                ('6.1', 'APIs'),
                ('6.2', 'Configuration management'),
                ('6.3', 'JSON data'),
            ],
        },

        # ── AWS Solutions Architect Associate (SAA-C03) ──
        'aws-saa-c03': {
            '1.0': [
                ('1.1', 'Secure access'),
                ('1.2', 'Secure workloads'),
                ('1.3', 'Encryption strategies'),
            ],
            '2.0': [
                ('2.1', 'Scalable architectures'),
                ('2.2', 'Highly available architectures'),
                ('2.3', 'Decoupling mechanisms'),
            ],
            '3.0': [
                ('3.1', 'Storage solutions'),
                ('3.2', 'Compute solutions'),
                ('3.3', 'Database solutions'),
                ('3.4', 'Network architectures'),
            ],
            '4.0': [
                ('4.1', 'Cost-effective storage'),
                ('4.2', 'Cost-effective compute'),
                ('4.3', 'Cost-effective databases'),
            ],
        },

        # ── CompTIA CySA+ (CS0-003) ──
        'comptia-cysa-cs0-003': {
            '1.0': [
                ('1.1', 'Security monitoring tools'),
                ('1.2', 'Log analysis and SIEM'),
                ('1.3', 'Threat intelligence'),
            ],
            '2.0': [
                ('2.1', 'Vulnerability scanning'),
                ('2.2', 'Vulnerability assessment tools'),
                ('2.3', 'Remediation and mitigation'),
            ],
            '3.0': [
                ('3.1', 'Incident response procedures'),
                ('3.2', 'Forensic analysis'),
                ('3.3', 'Incident recovery'),
            ],
            '4.0': [
                ('4.1', 'Reporting metrics and KPIs'),
                ('4.2', 'Stakeholder communication'),
            ],
        },

        # ── CompTIA PenTest+ (PT0-003) ──
        'comptia-pentest-pt0-003': {
            '1.0': [
                ('1.1', 'Engagement planning'),
                ('1.2', 'Scoping and authorization'),
                ('1.3', 'Compliance requirements'),
            ],
            '2.0': [
                ('2.1', 'Reconnaissance techniques'),
                ('2.2', 'Vulnerability scanning tools'),
                ('2.3', 'Enumeration methods'),
            ],
            '3.0': [
                ('3.1', 'Network attacks'),
                ('3.2', 'Application attacks'),
                ('3.3', 'Wireless and social engineering attacks'),
                ('3.4', 'Post-exploitation techniques'),
            ],
            '4.0': [
                ('4.1', 'Written reports'),
                ('4.2', 'Findings and remediation'),
            ],
            '5.0': [
                ('5.1', 'Penetration testing tools'),
                ('5.2', 'Scripting and code analysis'),
            ],
        },

        # ── CompTIA CASP+ (CAS-004) ──
        'comptia-casp-cas-004': {
            '1.0': [
                ('1.1', 'Security requirements analysis'),
                ('1.2', 'Enterprise architecture design'),
                ('1.3', 'Secure integration techniques'),
            ],
            '2.0': [
                ('2.1', 'Security monitoring activities'),
                ('2.2', 'Incident response management'),
                ('2.3', 'Vulnerability management programs'),
            ],
            '3.0': [
                ('3.1', 'Cryptographic techniques'),
                ('3.2', 'PKI and certificate management'),
                ('3.3', 'Secure protocols and services'),
            ],
            '4.0': [
                ('4.1', 'Risk management frameworks'),
                ('4.2', 'Compliance and policy development'),
            ],
        },

        # ── AWS Developer Associate (DVA-C02) ──
        'aws-dva-c02': {
            '1.0': [
                ('1.1', 'Lambda and serverless'),
                ('1.2', 'API Gateway'),
                ('1.3', 'DynamoDB and data stores'),
            ],
            '2.0': [
                ('2.1', 'IAM policies and roles'),
                ('2.2', 'Encryption and secrets management'),
            ],
            '3.0': [
                ('3.1', 'CI/CD pipelines'),
                ('3.2', 'Elastic Beanstalk and containers'),
                ('3.3', 'CloudFormation and SAM'),
            ],
            '4.0': [
                ('4.1', 'CloudWatch monitoring'),
                ('4.2', 'X-Ray tracing and debugging'),
            ],
        },

        # ── AWS SysOps Administrator Associate (SOA-C02) ──
        'aws-soa-c02': {
            '1.0': [
                ('1.1', 'CloudWatch metrics and alarms'),
                ('1.2', 'Log aggregation and analysis'),
                ('1.3', 'Automated remediation'),
            ],
            '2.0': [
                ('2.1', 'Backup and restore strategies'),
                ('2.2', 'High availability configurations'),
            ],
            '3.0': [
                ('3.1', 'AMI and instance provisioning'),
                ('3.2', 'Infrastructure as code'),
                ('3.3', 'Auto Scaling and load balancing'),
            ],
            '4.0': [
                ('4.1', 'IAM and access management'),
                ('4.2', 'Security groups and NACLs'),
                ('4.3', 'Compliance auditing'),
            ],
            '5.0': [
                ('5.1', 'VPC design and peering'),
                ('5.2', 'Route 53 and DNS'),
                ('5.3', 'CloudFront distributions'),
            ],
            '6.0': [
                ('6.1', 'Cost Explorer and budgets'),
                ('6.2', 'Performance optimization tools'),
            ],
        },

        # ── AWS Cloud Practitioner (CLF-C02) ──
        'aws-clf-c02': {
            '1.0': [
                ('1.1', 'Cloud value proposition'),
                ('1.2', 'Cloud architecture principles'),
                ('1.3', 'Cloud migration strategies'),
            ],
            '2.0': [
                ('2.1', 'Shared responsibility model'),
                ('2.2', 'IAM fundamentals'),
                ('2.3', 'Security services overview'),
            ],
            '3.0': [
                ('3.1', 'Compute services'),
                ('3.2', 'Storage services'),
                ('3.3', 'Database services'),
                ('3.4', 'Networking services'),
            ],
            '4.0': [
                ('4.1', 'Pricing models'),
                ('4.2', 'Support plans'),
            ],
        },

        # ── Microsoft Azure Fundamentals (AZ-900) ──
        'az-900': {
            '1.0': [
                ('1.1', 'Benefits of cloud computing'),
                ('1.2', 'Cloud service types'),
                ('1.3', 'Cloud deployment models'),
            ],
            '2.0': [
                ('2.1', 'Core Azure services'),
                ('2.2', 'Compute and networking'),
                ('2.3', 'Storage services'),
            ],
            '3.0': [
                ('3.1', 'Cost management tools'),
                ('3.2', 'Governance and compliance'),
                ('3.3', 'Resource management tools'),
            ],
        },

        # ── Microsoft Azure Administrator (AZ-104) ──
        'az-104': {
            '1.0': [
                ('1.1', 'Azure Active Directory'),
                ('1.2', 'Role-based access control'),
                ('1.3', 'Subscriptions and governance'),
            ],
            '2.0': [
                ('2.1', 'Storage accounts'),
                ('2.2', 'Blob and file storage'),
                ('2.3', 'Storage security'),
            ],
            '3.0': [
                ('3.1', 'Virtual machines'),
                ('3.2', 'App services'),
                ('3.3', 'Container instances'),
            ],
            '4.0': [
                ('4.1', 'Virtual networks'),
                ('4.2', 'Network security groups'),
                ('4.3', 'Load balancers and DNS'),
            ],
            '5.0': [
                ('5.1', 'Azure Monitor'),
                ('5.2', 'Backup and recovery'),
            ],
        },

        # ── Google Cloud Associate Cloud Engineer ──
        'gcp-ace': {
            '1.0': [
                ('1.1', 'Projects and accounts'),
                ('1.2', 'Billing management'),
                ('1.3', 'CLI and SDK installation'),
            ],
            '2.0': [
                ('2.1', 'Resource planning'),
                ('2.2', 'Pricing calculator'),
                ('2.3', 'Compute resource configuration'),
            ],
            '3.0': [
                ('3.1', 'Compute Engine deployment'),
                ('3.2', 'Kubernetes Engine'),
                ('3.3', 'Cloud Functions and App Engine'),
                ('3.4', 'Data solutions deployment'),
            ],
            '4.0': [
                ('4.1', 'Compute resource management'),
                ('4.2', 'Monitoring and logging'),
                ('4.3', 'Networking resources'),
            ],
            '5.0': [
                ('5.1', 'IAM configuration'),
                ('5.2', 'Service accounts'),
                ('5.3', 'Audit and security tools'),
            ],
        },
    }

    for cert_code, domains_map in sub_objectives.items():
        # Look up the certification ID
        c.execute('SELECT id FROM certifications WHERE code = ?', (cert_code,))
        cert_row = c.fetchone()
        if not cert_row:
            continue
        cert_id = cert_row['id']

        for parent_domain_code, children in domains_map.items():
            # Look up the parent domain ID
            c.execute(
                'SELECT id FROM domains WHERE certification_id = ? AND code = ? AND parent_domain_id IS NULL',
                (cert_id, parent_domain_code)
            )
            parent_row = c.fetchone()
            if not parent_row:
                continue
            parent_id = parent_row['id']

            for sort_idx, (child_code, child_name) in enumerate(children):
                c.execute(
                    '''INSERT INTO domains (certification_id, name, code, weight, parent_domain_id, sort_order)
                       VALUES (?, ?, ?, NULL, ?, ?)''',
                    (cert_id, child_name, child_code, parent_id, sort_idx)
                )

    conn.commit()
    conn.close()


def seed_security_plus_questions():
    """Seed a starter Security+ (SY0-701) question bank with 75 exam-aligned questions.

    Creates a system-owned public quiz with questions distributed across all 5 domains,
    automatically tagged via the domain system. This gives new users immediate access
    to readiness scoring, simulations, and weak-point tracking.
    """
    conn = get_db()
    c = conn.cursor()

    # Check if already seeded (look for our specific system quiz)
    c.execute("SELECT id FROM quizzes WHERE title = 'CompTIA Security+ (SY0-701) — Starter Question Bank'")
    if c.fetchone():
        conn.close()
        return

    # Get the Security+ certification ID
    c.execute("SELECT id FROM certifications WHERE code = 'comptia-sec-sy0-701'")
    cert_row = c.fetchone()
    if not cert_row:
        conn.close()
        return
    cert_id = cert_row['id']

    # Get domain IDs
    c.execute('''SELECT id, code, name FROM domains
        WHERE certification_id = ? AND parent_domain_id IS NULL
        ORDER BY sort_order''', (cert_id,))
    domains = {r['code']: r['id'] for r in c.fetchall()}

    # === QUESTIONS BY DOMAIN ===
    # Domain 1.0: General Security Concepts (12%)
    d1_questions = [
        {
            'question': 'Which security control type is designed to discourage a threat actor from performing an attack?',
            'type': 'choice',
            'options': ['Detective', 'Corrective', 'Deterrent', 'Compensating'],
            'correct': 2,
            'explanation': 'Deterrent controls discourage threat actors from attempting an attack. Examples include warning banners, security cameras, and security guards.'
        },
        {
            'question': 'What is the primary purpose of the CIA triad in information security?',
            'type': 'choice',
            'options': [
                'To classify data based on sensitivity levels',
                'To define the three fundamental goals of information security',
                'To establish a chain of custody for digital evidence',
                'To create a risk assessment framework'
            ],
            'correct': 1,
            'explanation': 'The CIA triad — Confidentiality, Integrity, and Availability — defines the three fundamental goals that information security programs must address.'
        },
        {
            'question': 'A company implements a policy requiring all changes to production systems to go through a formal review board. Which security concept does this represent?',
            'type': 'choice',
            'options': ['Incident response', 'Change management', 'Risk transfer', 'Data classification'],
            'correct': 1,
            'explanation': 'Change management ensures that all changes to IT systems are evaluated, approved, and documented to minimize disruption and security risks.'
        },
        {
            'question': 'Which cryptographic concept ensures that a sender cannot deny having sent a message?',
            'type': 'choice',
            'options': ['Confidentiality', 'Integrity', 'Non-repudiation', 'Authentication'],
            'correct': 2,
            'explanation': 'Non-repudiation ensures that the sender of a message cannot later deny having sent it. Digital signatures provide non-repudiation by using the sender\'s private key.'
        },
        {
            'question': 'What is the key difference between symmetric and asymmetric encryption?',
            'type': 'choice',
            'options': [
                'Symmetric uses one key; asymmetric uses a key pair',
                'Symmetric is slower than asymmetric',
                'Asymmetric cannot be used for digital signatures',
                'Symmetric requires a certificate authority'
            ],
            'correct': 0,
            'explanation': 'Symmetric encryption uses a single shared key for both encryption and decryption. Asymmetric encryption uses a key pair (public and private) — one encrypts, the other decrypts.'
        },
        {
            'question': 'A company wants to ensure that sensitive files have not been tampered with during transfer. Which technology should they use?',
            'type': 'choice',
            'options': ['AES-256 encryption', 'SHA-256 hashing', 'RSA key exchange', 'TLS handshake'],
            'correct': 1,
            'explanation': 'Hashing algorithms like SHA-256 create a fixed-length digest of data. Comparing hashes before and after transfer verifies that the file has not been modified (integrity).'
        },
        {
            'question': 'Which of the following is an example of a technical security control?',
            'type': 'choice',
            'options': ['Security awareness training', 'Firewall rules', 'Acceptable use policy', 'Background checks'],
            'correct': 1,
            'explanation': 'Technical (logical) controls are implemented through technology. Firewalls, IDS/IPS, encryption, and access control lists are all technical controls. Training and policies are administrative controls.'
        },
        {
            'question': 'What does the zero trust security model assume?',
            'type': 'choice',
            'options': [
                'Internal networks are inherently secure',
                'All users and devices must be verified regardless of location',
                'Firewalls provide sufficient perimeter protection',
                'VPN connections are always encrypted'
            ],
            'correct': 1,
            'explanation': 'Zero trust assumes no implicit trust is granted to users, devices, or networks — even those inside the corporate perimeter. Every access request must be verified, validated, and authorized.'
        },
        {
            'question': 'Which key stretching algorithm is specifically designed for password hashing?',
            'type': 'choice',
            'options': ['AES', 'bcrypt', 'RSA', 'SHA-256'],
            'correct': 1,
            'explanation': 'bcrypt is a key stretching algorithm designed specifically for password hashing. It incorporates a salt and is intentionally slow to resist brute-force attacks. PBKDF2 and Argon2 are also key stretching algorithms.'
        },
        {
            'question': 'What is the primary purpose of a gap analysis in security?',
            'type': 'choice',
            'options': [
                'To identify the difference between current and desired security posture',
                'To detect vulnerabilities in network infrastructure',
                'To monitor real-time security events',
                'To encrypt data at rest and in transit'
            ],
            'correct': 0,
            'explanation': 'A gap analysis compares an organization\'s current security posture against a desired state (such as a framework or standard) to identify areas that need improvement.'
        },
        {
            'question': 'Which of the following BEST describes the concept of least privilege?',
            'type': 'choice',
            'options': [
                'Users should have minimal password complexity requirements',
                'Users should only have the minimum access rights needed to perform their job',
                'Administrators should have access to all systems',
                'Guest accounts should be disabled on all systems'
            ],
            'correct': 1,
            'explanation': 'The principle of least privilege states that users, processes, and systems should be granted only the minimum permissions necessary to perform their required tasks — nothing more.'
        },
        {
            'question': 'A block cipher encrypts data in fixed-size blocks. Which of the following is a block cipher?',
            'type': 'choice',
            'options': ['RC4', 'AES', 'ChaCha20', 'Salsa20'],
            'correct': 1,
            'explanation': 'AES (Advanced Encryption Standard) is a block cipher that encrypts data in 128-bit blocks. RC4, ChaCha20, and Salsa20 are stream ciphers that encrypt data one bit or byte at a time.'
        },
        {
            'question': 'What type of certificate is used to establish the root of trust in a PKI hierarchy?',
            'type': 'choice',
            'options': ['Wildcard certificate', 'Self-signed root CA certificate', 'Extended validation certificate', 'Code signing certificate'],
            'correct': 1,
            'explanation': 'A self-signed root CA certificate establishes the root of trust in a PKI hierarchy. All other certificates in the chain derive their trust from this root certificate.'
        },
        {
            'question': 'Which of the following describes defense in depth?',
            'type': 'choice',
            'options': [
                'Using a single powerful firewall to protect the network',
                'Implementing multiple overlapping layers of security controls',
                'Encrypting all data at rest with AES-256',
                'Requiring biometric authentication for all users'
            ],
            'correct': 1,
            'explanation': 'Defense in depth is a security strategy that uses multiple, overlapping layers of security controls. If one layer fails, additional layers continue to provide protection.'
        },
        {
            'question': 'Which of the following is an example of something you are in multi-factor authentication?',
            'type': 'choice',
            'options': ['A password', 'A smart card', 'A fingerprint scan', 'A one-time PIN from an app'],
            'correct': 2,
            'explanation': 'Multi-factor authentication uses at least two of: something you know (password), something you have (smart card, token), and something you are (biometrics like fingerprints, iris scans). A fingerprint is a biometric — something you are.'
        },
    ]

    # Domain 2.0: Threats, Vulnerabilities, and Mitigations (22%)
    d2_questions = [
        {
            'question': 'An attacker sends a carefully crafted email that appears to come from the CEO, asking the CFO to wire money to an external account. What type of attack is this?',
            'type': 'choice',
            'options': ['Phishing', 'Whaling', 'Vishing', 'Smishing'],
            'correct': 1,
            'explanation': 'Whaling is a highly targeted form of phishing (spear phishing) directed at senior executives or high-value targets. The attack impersonates another executive to manipulate the target into taking action.'
        },
        {
            'question': 'Which type of malware encrypts files and demands payment for the decryption key?',
            'type': 'choice',
            'options': ['Trojan', 'Worm', 'Ransomware', 'Rootkit'],
            'correct': 2,
            'explanation': 'Ransomware encrypts the victim\'s files or locks system access, then demands a ransom payment (usually in cryptocurrency) in exchange for the decryption key.'
        },
        {
            'question': 'An attacker positions themselves between two communicating parties to intercept and potentially modify traffic. What is this attack called?',
            'type': 'choice',
            'options': ['SQL injection', 'On-path attack (MITM)', 'Cross-site scripting', 'Buffer overflow'],
            'correct': 1,
            'explanation': 'An on-path attack (man-in-the-middle / MITM) occurs when an attacker intercepts communication between two parties, potentially reading or modifying the data in transit.'
        },
        {
            'question': 'What is the primary difference between a vulnerability scan and a penetration test?',
            'type': 'choice',
            'options': [
                'Vulnerability scans are automated; penetration tests are always manual',
                'Vulnerability scans identify weaknesses; penetration tests actively exploit them',
                'Vulnerability scans require physical access; penetration tests are remote only',
                'There is no difference; the terms are interchangeable'
            ],
            'correct': 1,
            'explanation': 'Vulnerability scans identify and report potential weaknesses but do not exploit them. Penetration tests go further by actively attempting to exploit vulnerabilities to determine real-world impact.'
        },
        {
            'question': 'Which threat actor type is MOST likely to have advanced resources and target critical infrastructure?',
            'type': 'choice',
            'options': ['Script kiddie', 'Hacktivist', 'Nation-state', 'Insider threat'],
            'correct': 2,
            'explanation': 'Nation-state actors are government-sponsored threat actors with significant resources, advanced capabilities, and often target critical infrastructure, intellectual property, and government systems.'
        },
        {
            'question': 'What is a zero-day vulnerability?',
            'type': 'choice',
            'options': [
                'A vulnerability that has been patched within 24 hours',
                'A vulnerability that is unknown to the vendor and has no available patch',
                'A vulnerability in day-zero backup configurations',
                'A vulnerability that only exists on newly installed systems'
            ],
            'correct': 1,
            'explanation': 'A zero-day vulnerability is a previously unknown software flaw that the vendor has not yet patched. Attackers who discover zero-days can exploit them before any defense is available.'
        },
        {
            'question': 'An employee plugs in a USB drive they found in the parking lot. The drive installs malware when connected. What type of attack is this?',
            'type': 'choice',
            'options': ['Tailgating', 'Baiting', 'Pretexting', 'Shoulder surfing'],
            'correct': 1,
            'explanation': 'Baiting involves leaving infected physical media (USB drives, CDs) in a location where a target will find and use them. The attacker relies on human curiosity to deliver the malware payload.'
        },
        {
            'question': 'Which attack targets the web application layer by inserting malicious SQL statements into input fields?',
            'type': 'choice',
            'options': ['XSS', 'CSRF', 'SQL injection', 'Directory traversal'],
            'correct': 2,
            'explanation': 'SQL injection attacks insert malicious SQL code into application input fields that interact with a database, potentially allowing the attacker to read, modify, or delete data.'
        },
        {
            'question': 'What is the MITRE ATT&CK framework used for?',
            'type': 'choice',
            'options': [
                'Managing software development lifecycles',
                'Cataloguing adversary tactics, techniques, and procedures',
                'Encrypting sensitive data in transit',
                'Configuring network firewall rules'
            ],
            'correct': 1,
            'explanation': 'MITRE ATT&CK is a knowledge base of adversary tactics, techniques, and procedures (TTPs) based on real-world observations. It helps security teams understand, detect, and respond to threats.'
        },
        {
            'question': 'A company discovers that a former employee\'s credentials were used to access systems two weeks after termination. What control failure does this represent?',
            'type': 'choice',
            'options': [
                'Lack of encryption',
                'Insufficient offboarding procedures',
                'Missing intrusion detection',
                'Inadequate physical security'
            ],
            'correct': 1,
            'explanation': 'Proper offboarding procedures require immediately disabling or revoking access credentials when an employee leaves the organization to prevent unauthorized access.'
        },
        {
            'question': 'Which type of attack floods a target with traffic from many compromised systems simultaneously?',
            'type': 'choice',
            'options': ['DoS', 'DDoS', 'ARP spoofing', 'DNS poisoning'],
            'correct': 1,
            'explanation': 'A Distributed Denial of Service (DDoS) attack uses multiple compromised systems (a botnet) to overwhelm a target with traffic, making it unavailable to legitimate users.'
        },
        {
            'question': 'What is credential stuffing?',
            'type': 'choice',
            'options': [
                'Guessing passwords using dictionary words',
                'Using previously breached username/password pairs to access other services',
                'Capturing credentials via a keylogger',
                'Creating fake login pages to steal credentials'
            ],
            'correct': 1,
            'explanation': 'Credential stuffing uses stolen username/password combinations from data breaches to attempt login on other services, exploiting the common habit of password reuse across sites.'
        },
        {
            'question': 'Which vulnerability allows an attacker to inject client-side scripts into web pages viewed by other users?',
            'type': 'choice',
            'options': ['SQL injection', 'Cross-site scripting (XSS)', 'Buffer overflow', 'Race condition'],
            'correct': 1,
            'explanation': 'Cross-site scripting (XSS) allows attackers to inject malicious scripts into web applications. When other users view the affected page, the script executes in their browser, potentially stealing session cookies or redirecting to malicious sites.'
        },
        {
            'question': 'An organization wants to understand the likelihood and impact of various threats. What should they perform?',
            'type': 'choice',
            'options': ['Penetration test', 'Threat assessment', 'Risk analysis', 'Vulnerability scan'],
            'correct': 2,
            'explanation': 'Risk analysis evaluates both the likelihood of threats occurring and their potential impact on the organization, helping prioritize security investments and mitigation strategies.'
        },
        {
            'question': 'Which indicator of compromise (IoC) would MOST likely suggest a system has been compromised by a rootkit?',
            'type': 'choice',
            'options': [
                'Increased outbound traffic on port 443',
                'Discrepancies between OS-reported and raw disk file listings',
                'Multiple failed login attempts from a single IP',
                'Unexpected entries in the DNS cache'
            ],
            'correct': 1,
            'explanation': 'Rootkits hide their presence by modifying OS functions. Discrepancies between what the OS reports and what exists at the raw disk level indicate the OS is being manipulated — a classic rootkit indicator.'
        },
    ]

    # Domain 3.0: Security Architecture (18%)
    d3_questions = [
        {
            'question': 'Which network device operates at Layer 7 of the OSI model and can make routing decisions based on application data?',
            'type': 'choice',
            'options': ['Switch', 'Router', 'Application-layer firewall', 'Hub'],
            'correct': 2,
            'explanation': 'Application-layer (Layer 7) firewalls can inspect, filter, and make decisions based on application-specific data such as HTTP headers, URLs, and payload content.'
        },
        {
            'question': 'What is the primary purpose of a DMZ in network architecture?',
            'type': 'choice',
            'options': [
                'To connect branch offices via VPN tunnels',
                'To provide a buffer zone between the internet and the internal network for public-facing services',
                'To encrypt all internal network traffic',
                'To monitor network traffic for intrusion detection'
            ],
            'correct': 1,
            'explanation': 'A DMZ (demilitarized zone) is a network segment that sits between the internet and the internal network. Public-facing servers (web, email, DNS) are placed in the DMZ to limit exposure of the internal network.'
        },
        {
            'question': 'A company needs to ensure that their cloud provider meets specific security requirements. What document should they review?',
            'type': 'choice',
            'options': ['NDA', 'SLA', 'MOU', 'AUP'],
            'correct': 1,
            'explanation': 'A Service Level Agreement (SLA) defines the expected level of service, including security requirements, uptime guarantees, and remedies for failures between a service provider and customer.'
        },
        {
            'question': 'Which cloud deployment model provides dedicated infrastructure for a single organization?',
            'type': 'choice',
            'options': ['Public cloud', 'Private cloud', 'Community cloud', 'Hybrid cloud'],
            'correct': 1,
            'explanation': 'A private cloud provides infrastructure dedicated exclusively to a single organization, offering greater control over security, compliance, and customization compared to public cloud deployments.'
        },
        {
            'question': 'What is microsegmentation in a data center environment?',
            'type': 'choice',
            'options': [
                'Dividing storage into small partitions',
                'Creating granular security zones with individual workload-level policies',
                'Splitting DNS zones into subdomains',
                'Breaking large databases into smaller tables'
            ],
            'correct': 1,
            'explanation': 'Microsegmentation creates fine-grained security zones around individual workloads or applications, applying specific security policies at the workload level rather than just at the network perimeter.'
        },
        {
            'question': 'Which technology provides site-to-site encrypted connectivity over the public internet?',
            'type': 'choice',
            'options': ['NAT', 'IPSec VPN', 'VLAN', 'Port mirroring'],
            'correct': 1,
            'explanation': 'IPSec VPN creates encrypted tunnels over the public internet between two sites, providing confidentiality and integrity for data in transit between networks.'
        },
        {
            'question': 'In the shared responsibility model for IaaS cloud services, who is responsible for patching the guest operating system?',
            'type': 'choice',
            'options': [
                'The cloud provider exclusively',
                'The customer',
                'Both the cloud provider and the customer equally',
                'Neither — it is automated'
            ],
            'correct': 1,
            'explanation': 'In IaaS, the customer is responsible for everything from the OS up: patching, application security, data, and access control. The cloud provider manages the underlying infrastructure (hypervisor, network, physical security).'
        },
        {
            'question': 'What does a load balancer provide in terms of security and availability?',
            'type': 'choice',
            'options': [
                'Encrypts data at rest across multiple servers',
                'Distributes traffic across multiple servers to prevent overload and single points of failure',
                'Scans incoming traffic for malware signatures',
                'Provides NAT translation for internal servers'
            ],
            'correct': 1,
            'explanation': 'Load balancers distribute incoming traffic across multiple backend servers, improving availability (no single point of failure) and providing some DDoS mitigation by absorbing excess traffic.'
        },
        {
            'question': 'Which design principle involves running applications with the fewest features and services necessary?',
            'type': 'choice',
            'options': ['Defense in depth', 'Least functionality', 'Separation of duties', 'Fail-open'],
            'correct': 1,
            'explanation': 'Least functionality means configuring systems to provide only essential capabilities, disabling or removing unnecessary services, ports, and protocols to reduce the attack surface.'
        },
        {
            'question': 'What is the purpose of an air-gapped network?',
            'type': 'choice',
            'options': [
                'To improve wireless signal strength',
                'To physically isolate a network from all external connections',
                'To create redundant network paths',
                'To separate voice and data traffic'
            ],
            'correct': 1,
            'explanation': 'An air-gapped network is physically isolated from any external network, including the internet. It is used for highly sensitive systems where any external connectivity would pose an unacceptable risk.'
        },
        {
            'question': 'Which infrastructure as code concept ensures that deploying the same configuration always produces the same result?',
            'type': 'choice',
            'options': ['Elasticity', 'Idempotence', 'Scalability', 'Orchestration'],
            'correct': 1,
            'explanation': 'Idempotence means that applying the same operation multiple times produces the same result as applying it once. In infrastructure as code, this ensures consistent, repeatable deployments.'
        },
        {
            'question': 'A company deploys a web application firewall (WAF). At which layer does a WAF primarily operate?',
            'type': 'choice',
            'options': ['Layer 2 — Data Link', 'Layer 3 — Network', 'Layer 4 — Transport', 'Layer 7 — Application'],
            'correct': 3,
            'explanation': 'A WAF operates at Layer 7 (Application layer), inspecting HTTP/HTTPS traffic for attacks like SQL injection, XSS, and other web application exploits.'
        },
        {
            'question': 'What is the primary security benefit of containerization compared to traditional virtual machines?',
            'type': 'choice',
            'options': [
                'Containers provide stronger isolation than VMs',
                'Containers reduce the attack surface by sharing the host kernel without a full OS per instance',
                'Containers automatically encrypt all network traffic',
                'Containers eliminate the need for access controls'
            ],
            'correct': 1,
            'explanation': 'Containers are lightweight, sharing the host OS kernel. They reduce the attack surface compared to full VMs by eliminating the need for a complete OS per instance, though they provide less isolation than VMs.'
        },
        {
            'question': 'Which secure protocol should replace Telnet for remote command-line access?',
            'type': 'choice',
            'options': ['FTP', 'SSH', 'SNMP', 'RDP'],
            'correct': 1,
            'explanation': 'SSH (Secure Shell) provides encrypted remote command-line access and should always be used instead of Telnet, which transmits all data (including credentials) in cleartext.'
        },
        {
            'question': 'What type of redundancy involves maintaining a secondary site with current data that can take over operations within minutes?',
            'type': 'choice',
            'options': ['Cold site', 'Warm site', 'Hot site', 'Mobile site'],
            'correct': 2,
            'explanation': 'A hot site is a fully operational duplicate of the primary site with real-time data replication. It can assume operations within minutes (or instantly) after a failure, providing the fastest recovery time.'
        },
    ]

    # Domain 4.0: Security Operations (28%)
    d4_questions = [
        {
            'question': 'During an incident response, what is the FIRST step after detection?',
            'type': 'choice',
            'options': ['Eradication', 'Containment', 'Analysis', 'Recovery'],
            'correct': 2,
            'explanation': 'After detection, the next step is analysis (identification) — confirming and understanding the incident, its scope, and impact before moving to containment. The full order is: Preparation → Detection → Analysis → Containment → Eradication → Recovery → Lessons Learned.'
        },
        {
            'question': 'Which log source would BEST help identify unauthorized access attempts to a web application?',
            'type': 'choice',
            'options': ['Firewall logs', 'Web server access logs', 'DHCP logs', 'DNS query logs'],
            'correct': 1,
            'explanation': 'Web server access logs record all HTTP requests including URLs, status codes, client IPs, and user agents. They are the best source for identifying unauthorized access attempts, brute-force attacks, and exploitation against web applications.'
        },
        {
            'question': 'What is the primary purpose of a SIEM system?',
            'type': 'choice',
            'options': [
                'To encrypt sensitive data at rest',
                'To aggregate, correlate, and analyze security events from multiple sources',
                'To prevent malware from executing on endpoints',
                'To manage user identities and access rights'
            ],
            'correct': 1,
            'explanation': 'A SIEM (Security Information and Event Management) system collects and aggregates log data from multiple sources, correlates events, generates alerts, and provides dashboards for real-time security monitoring and incident investigation.'
        },
        {
            'question': 'Which of the following is a benefit of automated patch management?',
            'type': 'choice',
            'options': [
                'It eliminates all vulnerabilities',
                'It reduces the window of exposure by applying patches consistently and quickly',
                'It removes the need for change management',
                'It guarantees zero downtime during updates'
            ],
            'correct': 1,
            'explanation': 'Automated patch management reduces the window between patch release and deployment, ensuring consistent and timely application of security updates across the organization.'
        },
        {
            'question': 'An organization wants to test its incident response procedures without affecting production systems. What should they conduct?',
            'type': 'choice',
            'options': ['Full-scale disaster recovery', 'Tabletop exercise', 'Red team engagement', 'Production penetration test'],
            'correct': 1,
            'explanation': 'A tabletop exercise is a discussion-based session where team members walk through an incident scenario step by step. It tests the incident response plan without impacting production systems.'
        },
        {
            'question': 'What is the order of volatility in digital forensics?',
            'type': 'choice',
            'options': [
                'Hard drive → RAM → CPU cache → network traffic',
                'CPU registers → RAM → disk → archival media',
                'Archival media → disk → RAM → CPU registers',
                'RAM → disk → CPU registers → archival media'
            ],
            'correct': 1,
            'explanation': 'The order of volatility (most volatile first): CPU registers/cache → RAM → disk (temporary files) → disk (permanent) → remote logs → archival media. During forensics, collect the most volatile evidence first.'
        },
        {
            'question': 'Which identity management concept allows a user to authenticate once and access multiple applications?',
            'type': 'choice',
            'options': ['MFA', 'SSO', 'RBAC', 'PAM'],
            'correct': 1,
            'explanation': 'Single Sign-On (SSO) allows users to authenticate once and then access multiple applications and systems without re-entering credentials for each one.'
        },
        {
            'question': 'What is the primary purpose of a data loss prevention (DLP) solution?',
            'type': 'choice',
            'options': [
                'To backup data to prevent loss from hardware failure',
                'To detect and prevent unauthorized transmission of sensitive data',
                'To encrypt all data on endpoint devices',
                'To manage user passwords and access tokens'
            ],
            'correct': 1,
            'explanation': 'DLP solutions monitor, detect, and block the unauthorized transfer of sensitive data (PII, financial data, intellectual property) through email, web uploads, USB drives, and other channels.'
        },
        {
            'question': 'Which access control model assigns permissions based on a user\'s job function?',
            'type': 'choice',
            'options': ['DAC', 'MAC', 'RBAC', 'ABAC'],
            'correct': 2,
            'explanation': 'Role-Based Access Control (RBAC) assigns permissions to roles (e.g., "HR Manager", "Network Admin") rather than to individual users. Users inherit permissions when assigned to a role matching their job function.'
        },
        {
            'question': 'During a forensic investigation, what must be maintained to ensure evidence is admissible in court?',
            'type': 'choice',
            'options': ['Encryption key', 'Chain of custody', 'Service level agreement', 'Change management log'],
            'correct': 1,
            'explanation': 'Chain of custody documents who collected the evidence, when it was collected, how it was stored, and who has had access to it. Without proper chain of custody, evidence may be inadmissible in legal proceedings.'
        },
        {
            'question': 'What type of security tool uses signature-based and behavioral analysis to detect threats on individual devices?',
            'type': 'choice',
            'options': ['SIEM', 'EDR', 'NAC', 'SOAR'],
            'correct': 1,
            'explanation': 'Endpoint Detection and Response (EDR) tools monitor endpoint devices for suspicious activity using signatures, behavioral analysis, and machine learning. They provide visibility, detection, and response capabilities at the endpoint level.'
        },
        {
            'question': 'What is the purpose of network access control (NAC)?',
            'type': 'choice',
            'options': [
                'To encrypt wireless network traffic',
                'To ensure only compliant devices can connect to the network',
                'To route traffic between network segments',
                'To cache frequently accessed web content'
            ],
            'correct': 1,
            'explanation': 'NAC solutions verify that devices meet security requirements (antivirus updated, patches applied, compliant configuration) before allowing them to connect to the network, and can quarantine non-compliant devices.'
        },
        {
            'question': 'Which of the following is the MOST effective way to protect against password-based attacks?',
            'type': 'choice',
            'options': [
                'Requiring passwords to be changed every 30 days',
                'Implementing multi-factor authentication',
                'Increasing minimum password length to 10 characters',
                'Using a password complexity requirement'
            ],
            'correct': 1,
            'explanation': 'MFA adds additional authentication factors beyond just a password. Even if a password is compromised through phishing, brute force, or credential stuffing, the attacker still needs the additional factor(s) to gain access.'
        },
        {
            'question': 'A security analyst notices unusual outbound DNS queries to a suspicious domain. What might this indicate?',
            'type': 'choice',
            'options': [
                'A misconfigured DNS server',
                'DNS tunneling used for command-and-control or data exfiltration',
                'A standard software update process',
                'Normal web browsing activity'
            ],
            'correct': 1,
            'explanation': 'DNS tunneling encodes data within DNS queries and responses to establish covert communication channels. Attackers use this technique for command-and-control (C2) communication and data exfiltration, often bypassing traditional security controls.'
        },
        {
            'question': 'What is the key difference between a vulnerability assessment and a threat hunt?',
            'type': 'choice',
            'options': [
                'Vulnerability assessments are automated; threat hunts are always manual',
                'Vulnerability assessments look for weaknesses; threat hunts proactively look for active compromises',
                'Threat hunts only use open-source tools',
                'Vulnerability assessments require physical access to systems'
            ],
            'correct': 1,
            'explanation': 'Vulnerability assessments identify potential weaknesses in systems. Threat hunting is a proactive process where analysts actively search for signs of compromise or ongoing attacks that may have evaded automated detection.'
        },
    ]

    # Domain 5.0: Security Program Management and Oversight (20%)
    d5_questions = [
        {
            'question': 'Which risk management strategy involves purchasing insurance to cover potential losses?',
            'type': 'choice',
            'options': ['Risk avoidance', 'Risk transfer', 'Risk mitigation', 'Risk acceptance'],
            'correct': 1,
            'explanation': 'Risk transfer shifts the financial impact of a risk to a third party. Purchasing cyber insurance is a common example — the insurance company assumes the financial burden if a covered incident occurs.'
        },
        {
            'question': 'What is the purpose of a Business Impact Analysis (BIA)?',
            'type': 'choice',
            'options': [
                'To test backup restoration procedures',
                'To identify critical business functions and the impact of their disruption',
                'To scan systems for vulnerabilities',
                'To document network topology'
            ],
            'correct': 1,
            'explanation': 'A BIA identifies critical business functions, assesses the impact of their disruption over time, and determines recovery priorities and resource requirements. It informs business continuity and disaster recovery planning.'
        },
        {
            'question': 'Which compliance framework is specifically designed for organizations that process credit card payments?',
            'type': 'choice',
            'options': ['HIPAA', 'SOX', 'PCI DSS', 'FERPA'],
            'correct': 2,
            'explanation': 'PCI DSS (Payment Card Industry Data Security Standard) is a set of security requirements for organizations that store, process, or transmit cardholder data. It is maintained by the PCI Security Standards Council.'
        },
        {
            'question': 'What defines the maximum acceptable amount of data loss measured in time during a disaster?',
            'type': 'choice',
            'options': ['RTO', 'RPO', 'MTTR', 'MTBF'],
            'correct': 1,
            'explanation': 'Recovery Point Objective (RPO) defines the maximum acceptable amount of data loss measured in time. An RPO of 4 hours means the organization can tolerate losing up to 4 hours of data.'
        },
        {
            'question': 'An organization implements security awareness training for all employees. What type of control is this?',
            'type': 'choice',
            'options': ['Technical', 'Administrative', 'Physical', 'Compensating'],
            'correct': 1,
            'explanation': 'Security awareness training is an administrative (managerial) control. Administrative controls include policies, procedures, training, and governance mechanisms that direct how security is managed.'
        },
        {
            'question': 'Which regulation specifically protects the personal data and privacy of EU citizens?',
            'type': 'choice',
            'options': ['HIPAA', 'SOX', 'GDPR', 'CCPA'],
            'correct': 2,
            'explanation': 'The General Data Protection Regulation (GDPR) protects personal data and privacy of individuals in the European Union, imposing strict requirements on data collection, processing, storage, and transfer.'
        },
        {
            'question': 'What is the purpose of a risk register?',
            'type': 'choice',
            'options': [
                'To log all security incidents as they occur',
                'To document identified risks, their assessment, and mitigation plans',
                'To track employee access permissions',
                'To store encryption keys securely'
            ],
            'correct': 1,
            'explanation': 'A risk register is a centralized document that tracks all identified risks, their likelihood, impact, risk rating, owners, mitigation strategies, and current status. It is a core artifact in risk management programs.'
        },
        {
            'question': 'What does RTO (Recovery Time Objective) define?',
            'type': 'choice',
            'options': [
                'The maximum acceptable data loss in time',
                'The maximum acceptable time to restore a system after a failure',
                'The expected frequency of system failures',
                'The total cost of a security incident'
            ],
            'correct': 1,
            'explanation': 'Recovery Time Objective (RTO) defines the maximum acceptable time between a system failure and its restoration to operational status. A shorter RTO requires more investment in recovery capabilities.'
        },
        {
            'question': 'A company wants to evaluate the effectiveness of its security controls against a recognized standard. What should they conduct?',
            'type': 'choice',
            'options': ['Vulnerability scan', 'Penetration test', 'Security audit', 'Threat hunt'],
            'correct': 2,
            'explanation': 'A security audit is a systematic evaluation of an organization\'s security posture against a set of criteria, standards, or regulatory requirements (such as ISO 27001, NIST, or PCI DSS).'
        },
        {
            'question': 'Which of the following is an example of data sovereignty?',
            'type': 'choice',
            'options': [
                'Encrypting data before storing it in the cloud',
                'A regulation requiring citizen data to be stored within the country\'s borders',
                'Implementing role-based access controls on databases',
                'Using a content delivery network for global distribution'
            ],
            'correct': 1,
            'explanation': 'Data sovereignty refers to laws and regulations that require data to be stored and processed within the borders of a specific country or jurisdiction, reflecting the principle that data is subject to the laws of the country where it resides.'
        },
        {
            'question': 'What is the primary purpose of a security policy?',
            'type': 'choice',
            'options': [
                'To configure firewall rules for the network',
                'To establish management direction and expectations for protecting organizational assets',
                'To detail step-by-step procedures for system administration',
                'To define technical specifications for security tools'
            ],
            'correct': 1,
            'explanation': 'A security policy is a high-level document that establishes management\'s intent, direction, and expectations for how the organization will protect its information assets. Procedures and standards provide the implementation details.'
        },
        {
            'question': 'Which of the following BEST describes the concept of due diligence in security?',
            'type': 'choice',
            'options': [
                'Installing antivirus software on all endpoints',
                'Continuously researching and understanding threats and risks relevant to the organization',
                'Responding to incidents within the defined SLA',
                'Requiring all employees to use complex passwords'
            ],
            'correct': 1,
            'explanation': 'Due diligence is the ongoing practice of researching, understanding, and staying informed about threats, vulnerabilities, and risks. Due care is the implementation of security controls based on that understanding.'
        },
        {
            'question': 'An organization classifies its data as Public, Internal, Confidential, and Restricted. What security concept does this represent?',
            'type': 'choice',
            'options': ['Access control lists', 'Data classification', 'Network segmentation', 'Role-based access'],
            'correct': 1,
            'explanation': 'Data classification assigns sensitivity labels to data based on its value and the impact of unauthorized disclosure. It guides the appropriate level of security controls and handling procedures for each category.'
        },
        {
            'question': 'What is the role of a data processor under GDPR?',
            'type': 'choice',
            'options': [
                'The individual whose data is being collected',
                'The entity that determines the purposes and means of processing personal data',
                'The entity that processes personal data on behalf of the data controller',
                'The regulatory authority that enforces GDPR compliance'
            ],
            'correct': 2,
            'explanation': 'Under GDPR, the data processor processes personal data on behalf of the data controller (who determines the purpose and means of processing). Cloud service providers often act as data processors for their customers.'
        },
        {
            'question': 'Which type of agreement defines the terms and conditions for sharing data between two organizations?',
            'type': 'choice',
            'options': ['SLA', 'NDA', 'ISA (Interconnection Security Agreement)', 'BPA'],
            'correct': 2,
            'explanation': 'An Interconnection Security Agreement (ISA) defines the security requirements and responsibilities for connecting information systems between two organizations, including data sharing terms and security controls.'
        },
    ]

    # Build the full question list with domain info
    all_questions = []
    domain_map = [
        ('1.0', d1_questions),
        ('2.0', d2_questions),
        ('3.0', d3_questions),
        ('4.0', d4_questions),
        ('5.0', d5_questions),
    ]

    for domain_code, questions in domain_map:
        for q in questions:
            q['_domain_code'] = domain_code
            all_questions.append(q)

    # Create the system quiz (user_id=0 as system-owned)
    # First, ensure user_id 0 doesn't cause FK issues — use the first user, or create without FK
    questions_json = json.dumps([{
        'question': q['question'],
        'type': q['type'],
        'options': q.get('options'),
        'correct': q.get('correct'),
        'explanation': q.get('explanation'),
    } for q in all_questions])

    # Use user_id 0 for system content (not FK-enforced in SQLite by default)
    c.execute('''INSERT INTO quizzes
        (user_id, title, description, questions, color, is_public, is_migrated)
        VALUES (0, ?, ?, ?, ?, 1, 1)''',
        ('CompTIA Security+ (SY0-701) — Starter Question Bank',
         '75 exam-aligned practice questions covering all 5 Security+ domains. '
         'Use for readiness scoring, exam simulation, and identifying weak areas.',
         questions_json, '#8b5cf6'))
    quiz_id = c.lastrowid

    # Insert normalized questions and tag domains
    for idx, q in enumerate(all_questions):
        c.execute('''INSERT INTO questions
            (quiz_id, question_index, question_text, type, options, correct,
             pairs, code, code_language, image, image_alt, explanation, difficulty)
            VALUES (?, ?, ?, ?, ?, ?, NULL, NULL, NULL, NULL, NULL, ?, 0)''',
            (quiz_id, idx, q['question'], q['type'],
             json.dumps(q.get('options')),
             json.dumps(q.get('correct')),
             q.get('explanation')))
        question_id = c.lastrowid

        # Tag to the specific domain
        domain_code = q['_domain_code']
        domain_id = domains.get(domain_code)
        if domain_id:
            c.execute('INSERT OR IGNORE INTO question_domains (question_id, domain_id) VALUES (?, ?)',
                      (question_id, domain_id))

    conn.commit()
    conn.close()
    print(f"[STARTUP] Seeded Security+ question bank: {len(all_questions)} questions across 5 domains", flush=True)


def seed_study_resources():
    """Seed study resources for major certifications."""
    conn = get_db()
    c = conn.cursor()

    c.execute('SELECT COUNT(*) as cnt FROM study_resources')
    if c.fetchone()['cnt'] > 0:
        conn.close()
        return

    # Build cert code -> id map
    c.execute('SELECT id, code FROM certifications')
    cert_map = {r['code']: r['id'] for r in c.fetchall()}

    resources = [
        # CompTIA Network+ (N10-009)
        ('comptia-net-n10-009', 'Professor Messer Network+ Course', 'https://www.professormesser.com/network-plus/n10-009/n10-009-video/n10-009-training-course/', 'video', 'Professor Messer', True),
        ('comptia-net-n10-009', 'CompTIA Network+ Official Study Guide', 'https://www.comptia.org/training/books/network-n10-009-study-guide', 'book', 'CompTIA', False),
        ('comptia-net-n10-009', 'CompTIA Network+ Exam Objectives', 'https://www.comptia.org/certifications/network', 'article', 'CompTIA', True),
        # CompTIA Security+ (SY0-701)
        ('comptia-sec-sy0-701', 'Professor Messer Security+ Course', 'https://www.professormesser.com/security-plus/sy0-701/sy0-701-video/sy0-701-comptia-security-plus-course/', 'video', 'Professor Messer', True),
        ('comptia-sec-sy0-701', 'CompTIA Security+ Official Study Guide', 'https://www.comptia.org/training/books/security-sy0-701-study-guide', 'book', 'CompTIA', False),
        ('comptia-sec-sy0-701', 'CompTIA Security+ Exam Objectives', 'https://www.comptia.org/certifications/security', 'article', 'CompTIA', True),
        # CompTIA A+ Core 1
        ('comptia-a-core1-221-1101', 'Professor Messer A+ Core 1 Course', 'https://www.professormesser.com/a-plus/220-1101/220-1101-video/220-1101-training-course/', 'video', 'Professor Messer', True),
        ('comptia-a-core1-221-1101', 'CompTIA A+ Exam Objectives', 'https://www.comptia.org/certifications/a', 'article', 'CompTIA', True),
        # CompTIA A+ Core 2
        ('comptia-a-core2-221-1102', 'Professor Messer A+ Core 2 Course', 'https://www.professormesser.com/a-plus/220-1102/220-1102-video/220-1102-training-course/', 'video', 'Professor Messer', True),
        # Cisco CCNA
        ('cisco-ccna-200-301', 'Jeremy\'s IT Lab CCNA Course', 'https://www.youtube.com/playlist?list=PLxbwE86jKRgMpuZuLBivzlM8s2Dk5lXBQ', 'video', 'Jeremy\'s IT Lab', True),
        ('cisco-ccna-200-301', 'Cisco CCNA Official Cert Guide', 'https://www.ciscopress.com/store/ccna-200-301-official-cert-guide-library-9781587147142', 'book', 'Cisco Press', False),
        ('cisco-ccna-200-301', 'Cisco CCNA Exam Topics', 'https://www.cisco.com/c/en/us/training-events/training-certifications/exams/current-list/ccna-200-301.html', 'article', 'Cisco', True),
        # AWS Solutions Architect Associate
        ('aws-saa-c03', 'Stephane Maarek AWS SAA Course', 'https://www.udemy.com/course/aws-certified-solutions-architect-associate-saa-c03/', 'video', 'Udemy', False),
        ('aws-saa-c03', 'AWS SAA Exam Guide', 'https://aws.amazon.com/certification/certified-solutions-architect-associate/', 'article', 'AWS', True),
        ('aws-saa-c03', 'AWS Well-Architected Framework', 'https://docs.aws.amazon.com/wellarchitected/latest/framework/welcome.html', 'article', 'AWS', True),
        # AWS Cloud Practitioner
        ('aws-clf-c02', 'AWS Cloud Practitioner Essentials', 'https://aws.amazon.com/training/digital/aws-cloud-practitioner-essentials/', 'video', 'AWS', True),
        ('aws-clf-c02', 'AWS CLF Exam Guide', 'https://aws.amazon.com/certification/certified-cloud-practitioner/', 'article', 'AWS', True),
        # Azure Fundamentals
        ('az-900', 'Microsoft Learn AZ-900 Path', 'https://learn.microsoft.com/en-us/certifications/azure-fundamentals/', 'article', 'Microsoft', True),
        ('az-900', 'John Savill AZ-900 Study Cram', 'https://www.youtube.com/watch?v=8n-kWJetQRk', 'video', 'YouTube', True),
    ]

    for i, (cert_code, title, url, rtype, provider, is_free) in enumerate(resources):
        cert_id = cert_map.get(cert_code)
        if not cert_id:
            continue
        c.execute('''INSERT INTO study_resources (certification_id, title, url, resource_type, provider, is_free, sort_order)
            VALUES (?, ?, ?, ?, ?, ?, ?)''',
            (cert_id, title, url, rtype, provider, 1 if is_free else 0, i))

    conn.commit()
    conn.close()
    print(f"[STARTUP] Seeded {len(resources)} study resources", flush=True)


# Initialize database tables on module load (for WSGI)
try:
    init_db()
    seed_certifications()
    seed_sub_objectives()
    seed_security_plus_questions()
    seed_study_resources()
except Exception as e:
    print(f"[STARTUP] DB init error: {e}", flush=True)

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)